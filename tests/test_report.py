import base64
import os
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
import pytest
from sqlmodel import select

from assetmap.config import AiConfig, AppConfig, DatabaseConfig
from assetmap.db import create_db_and_engine, get_session
from assetmap.models import (
    Company,
    CompanyAssetLink,
    CompanyEdge,
    AiAnalysis,
    DnsRecord,
    InternetAsset,
    NmapPort,
    NmapScanTask,
    ReportGenerationTask,
    ScanTask,
    ServiceAsset,
    SourceRawRecord,
    WebEntrypoint,
)
from assetmap.services.operations.gap_template import GapTemplateService
from assetmap.services.operations.improvement_plan import ImprovementPlanService
from assetmap.services.delivery.exporter import ExportService
from assetmap.services.delivery.package import DeliveryPackageService, DeliveryPackageVerifier
from assetmap.services.delivery.quality import DeliveryQualityService
from assetmap.services.delivery.report import ReportService
from assetmap.services.operations.review_import import ReviewImportService
from assetmap.services.operations.review_workorder import ReviewWorkOrderService


def test_report_generates_docx_and_two_workbooks(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    config.ai.enabled = False
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    task = ScanTask(id=1, target="示例集团有限公司", status="completed")
    company = Company(id=1, name="示例集团有限公司", normalized_name="示例集团有限公司")
    domain = InternetAsset(
        id=1,
        asset_type="icp_domain",
        normalized_identifier="example.cn",
        display_name="example.cn",
        raw_payload={},
    )
    screenshot_dir = tmp_path / "data" / "screenshots" / "task_1"
    screenshot_dir.mkdir(parents=True)
    screenshot_path = screenshot_dir / "1_example.png"
    screenshot_path.write_bytes(
        base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
        )
    )
    session.add(task)
    session.add(company)
    session.add(domain)
    session.add(CompanyAssetLink(task_id=1, company_id=1, asset_id=1, source_tool="manual", raw_payload={}))
    session.add(DnsRecord(scan_task_id=1, fqdn="www.example.cn", root_domain="example.cn", record_type="A", value="203.0.113.10"))
    session.add(DnsRecord(scan_task_id=1, fqdn="expired.example.cn", root_domain="example.cn", record_type="A", value="203.0.113.11"))
    session.add(NmapScanTask(scan_task_id=1, status="completed", targets=["203.0.113.10"]))
    session.add(
        NmapPort(
            scan_task_id=1,
            target_ip="203.0.113.10",
            protocol="tcp",
            port=443,
            state="open",
            service="https",
            raw_payload={
                "sources": ["nmap", "fofa"],
                "fofa": {"host": "https://www.example.cn/", "title": "示例门户"},
            },
        )
    )
    session.add(
        ServiceAsset(
            id=1,
            scan_task_id=1,
            target_ip="203.0.113.10",
            protocol="tcp",
            port=443,
            asset_kind="web",
            host_mode="virtual_host",
            representative_url="https://www.example.cn/",
            domains=["www.example.cn"],
            title="示例门户",
        )
    )
    session.add(
        WebEntrypoint(
            scan_task_id=1,
            service_asset_id=1,
            target_ip="203.0.113.10",
            port=443,
            host="www.example.cn",
            url="https://www.example.cn/",
            normalized_url="https://www.example.cn/",
            http_status=200,
            title="示例门户",
            evidence={
                "visual_analysis": {
                    "system_name": "示例门户",
                    "site_purpose": "对外服务入口",
                    "screenshot_path": str(screenshot_path),
                }
            },
        )
    )
    session.commit()

    result = ReportService(session, config).run(1, output_dir=tmp_path / "reports")

    assert result.report_path.exists()
    assert result.asset_workbook_path.exists()
    assert result.web_workbook_path.exists()
    generation = session.exec(select(ReportGenerationTask).where(ReportGenerationTask.scan_task_id == 1)).one()
    assert generation.status == "completed"
    assert generation.report_path == str(result.report_path)
    ai_audit = __import__("json").loads(Path("data/report/task_1/report_ai_audit.json").read_text(encoding="utf-8"))
    assert ai_audit["section_count"] == 4
    assert ai_audit["status_counts"] == {"skipped": 4}
    assert {section["analysis_type"] for section in ai_audit["sections"]} == {
        "report_dns",
        "report_ports",
        "report_web",
        "report_summary",
    }
    required_audit_fields = {
        "mode",
        "input_fingerprint",
        "data_shape",
        "response_keys",
        "response_id",
        "response_model",
        "usage_keys",
        "prompt_tokens",
        "completion_tokens",
        "total_tokens",
    }
    assert all(required_audit_fields <= set(section) for section in ai_audit["sections"])
    doc = Document(result.report_path)
    assert "互联网数字资产暴露面测绘报告" in "\n".join(p.text for p in doc.paragraphs)
    assert "示例集团有限公司" in doc.sections[0].header.paragraphs[0].text
    assert "assetmap 自动生成" in doc.sections[0].footer.paragraphs[0].text

    assert len(doc.tables) >= 10
    asset_wb = load_workbook(result.asset_workbook_path)
    asset_wb_sheets = asset_wb.sheetnames
    assert asset_wb_sheets[0] == "阅读导航"
    assert "管理驾驶舱" in asset_wb_sheets
    assert "报告概览" in asset_wb_sheets
    assert "AI分析审计" in asset_wb_sheets
    assert "风险统计" in asset_wb_sheets
    assert "重点资产视图" in asset_wb_sheets
    assert "单位覆盖台账" in asset_wb_sheets
    assert "开放端口" in asset_wb_sheets
    assert "覆盖缺口" in asset_wb_sheets
    assert "风险清单" in asset_wb_sheets
    assert "整改矩阵" in asset_wb_sheets
    assert "DNS复核清单" in asset_wb_sheets
    assert "端口目标台账" in asset_wb_sheets
    assert "服务识别台账" in asset_wb_sheets
    assert "URL入口覆盖" in asset_wb_sheets
    assert "交付审计文件" in asset_wb_sheets
    web_wb = load_workbook(result.web_workbook_path)
    assert web_wb.sheetnames[0] == "阅读导航"
    assert "重点Web资产" in web_wb.sheetnames
    assert "截图证据" in web_wb.sheetnames
    assert "Web资产详情" in web_wb.sheetnames
    assert "视觉复核清单" in web_wb.sheetnames
    nav = asset_wb["阅读导航"]
    assert nav["A2"].value == "管理驾驶舱"
    assert nav["A2"].hyperlink.target == "#'管理驾驶舱'!A1"
    dashboard = asset_wb["管理驾驶舱"]
    dashboard_headers = [cell.value for cell in next(dashboard.iter_rows(min_row=1, max_row=1))]
    assert dashboard_headers[:4] == ["分组", "指标", "数值", "说明"]
    assert len(getattr(dashboard, "_charts", [])) >= 2
    ai_audit_sheet = asset_wb["AI分析审计"]
    ai_audit_headers = [cell.value for cell in next(ai_audit_sheet.iter_rows(min_row=1, max_row=1))]
    assert {"分析分块", "状态", "模型", "输入指纹", "输入规模", "响应ID", "响应模型", "Usage字段", "总Token"} <= set(ai_audit_headers)
    assert ai_audit_sheet.max_row == 5
    assert asset_wb["覆盖缺口"].sheet_properties.tabColor.rgb.endswith("ED7D31")
    assert asset_wb["覆盖缺口"]["A1"].fill.fgColor.rgb.endswith("1F4E78")
    unit_sheet = asset_wb["单位覆盖台账"]
    unit_headers = [cell.value for cell in next(unit_sheet.iter_rows(min_row=1, max_row=1))]
    assert "覆盖状态" in unit_headers
    assert "建议动作" in unit_headers
    assert "股权路径" in unit_headers
    assert unit_sheet.max_row >= 2
    port_sheet = load_workbook(result.asset_workbook_path, read_only=True)["开放端口"]
    headers = [cell.value for cell in next(port_sheet.iter_rows(min_row=1, max_row=1))]
    assert "FOFA Host" in headers
    assert "证据类型" in headers
    assert "主动扫描确认" in headers
    assert "被动FOFA证据" in headers
    target_sheet = load_workbook(result.asset_workbook_path, read_only=True)["端口目标台账"]
    target_headers = [cell.value for cell in next(target_sheet.iter_rows(min_row=1, max_row=1))]
    assert "目标来源" in target_headers
    assert "是否进入扫描目标" in target_headers
    assert "端口证据类型" in target_headers
    service_sheet = load_workbook(result.asset_workbook_path, read_only=True)["服务识别台账"]
    service_headers = [cell.value for cell in next(service_sheet.iter_rows(min_row=1, max_row=1))]
    assert "分类依据" in service_headers
    assert "URL入口数量" in service_headers
    url_coverage_sheet = load_workbook(result.asset_workbook_path, read_only=True)["URL入口覆盖"]
    url_coverage_headers = [cell.value for cell in next(url_coverage_sheet.iter_rows(min_row=1, max_row=1))]
    assert "覆盖结论" in url_coverage_headers
    key_asset_sheet = load_workbook(result.asset_workbook_path, read_only=True)["重点资产视图"]
    key_asset_headers = [cell.value for cell in next(key_asset_sheet.iter_rows(min_row=1, max_row=1))]
    assert "优先级" in key_asset_headers
    assert "风险分值" in key_asset_headers
    assert "处置建议" in key_asset_headers
    assert "责任建议" in key_asset_headers
    web_sheet = load_workbook(result.web_workbook_path, read_only=True)["Web资产详情"]
    web_headers = [cell.value for cell in next(web_sheet.iter_rows(min_row=1, max_row=1))]
    assert "识别方式" in web_headers
    assert "识别置信度" in web_headers
    key_web_sheet = load_workbook(result.web_workbook_path, read_only=True)["重点Web资产"]
    key_web_headers = [cell.value for cell in next(key_web_sheet.iter_rows(min_row=1, max_row=1))]
    assert "URL" in key_web_headers
    assert "AI识别系统" in key_web_headers
    evidence_sheet = load_workbook(result.web_workbook_path, read_only=True)["截图证据"]
    evidence_headers = [cell.value for cell in next(evidence_sheet.iter_rows(min_row=1, max_row=1))]
    assert "缩略图" in evidence_headers
    assert "截图文件" in evidence_headers
    assert "截图状态" in evidence_headers
    review_sheet = load_workbook(result.web_workbook_path, read_only=True)["视觉复核清单"]
    assert review_sheet["A1"].value in {"无数据", "复核优先级"}
    overview = load_workbook(result.asset_workbook_path, read_only=True)["报告概览"]
    overview_headers = [cell.value for cell in next(overview.iter_rows(min_row=1, max_row=1))]
    assert overview_headers == ["指标", "结果", "说明"]
    coverage_sheet = load_workbook(result.asset_workbook_path, read_only=True)["覆盖缺口"]
    coverage_headers = [cell.value for cell in next(coverage_sheet.iter_rows(min_row=1, max_row=1))]
    assert coverage_headers == ["环节", "指标", "结果", "缺口等级", "缺口样例", "建议动作"]
    coverage_rows = list(coverage_sheet.iter_rows(min_row=2, values_only=True))
    port_coverage = next(row for row in coverage_rows if row[0] == "端口发现")
    assert port_coverage[1] == "候选 IP 端口发现覆盖"
    assert port_coverage[3] == "无"
    dns_review_sheet = load_workbook(result.asset_workbook_path, read_only=True)["DNS复核清单"]
    dns_review_headers = [cell.value for cell in next(dns_review_sheet.iter_rows(min_row=1, max_row=1))]
    assert "复核优先级" in dns_review_headers
    assert "建议动作" in dns_review_headers
    audit_sheet = load_workbook(result.asset_workbook_path, read_only=True)["交付审计文件"]
    audit_values = list(audit_sheet.iter_rows(min_row=2, values_only=True))
    assert any("报告AI分析审计" in row[0] for row in audit_values)
    doc_text = "\n".join(p.text for p in Document(result.report_path).paragraphs)
    assert "覆盖缺口分析" in doc_text
    assert "报告信息" in doc_text
    assert "报告目录" in doc_text
    assert "风险概览" in doc_text
    assert "重点单位覆盖矩阵" in doc_text
    assert "复核与质量门禁计划" in doc_text
    assert "质量门禁" in doc_text
    assert "服务识别台账" in doc_text
    assert "交付审计文件" in doc_text
    doc_cell_text = "\n".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)
    assert "对应材料" in doc_cell_text
    assert "单位资产补充" in doc_cell_text
    assert "高/中优先级待补充" in doc_cell_text
    assert not any(p.text.strip().startswith("#") for p in Document(result.report_path).paragraphs)

    quality = DeliveryQualityService(session, config).check(1, output_dir=tmp_path / "reports")
    assert quality.status == "PASS"
    assert any("Quality: PASS" in line for line in quality.lines)
    assert any("Word报告: ok" in line for line in quality.lines)
    assert not any("缺少字段" in item for item in quality.failures)
    assert not any("图表数量偏少" in item for item in quality.warnings)

    package = DeliveryPackageService(session, config).package(
        1,
        reports_dir=tmp_path / "reports",
        output_dir=tmp_path / "deliveries",
    )
    assert package.zip_path.exists()
    assert package.manifest_path.exists()
    assert len(package.packaged_files) == len(package.files) + 1
    manifest = __import__("json").loads(package.manifest_path.read_text(encoding="utf-8"))
    assert manifest["quality_status"] == "PASS"
    assert {item["path"] for item in manifest["files"]} >= {
        "task_1_互联网资产暴露面测绘报告.docx",
        "task_1_资产汇总.xlsx",
        "task_1_Web资产详情.xlsx",
        "quality_summary.txt",
        "task_1_待补充资产模板.yaml",
        "task_1_复核工作单.yaml",
        "task_1_补全计划.json",
        "task_1_补全计划.txt",
        "task_1_报告AI分析审计.json",
        "交付说明.txt",
    }
    assert any(item["path"].startswith("screenshots/") and item["path"].endswith("1_example.png") for item in manifest["files"])
    assert (package.package_dir / "screenshots" / "1_example.png").exists()
    screenshot_manifest_path = package.package_dir / "task_1_截图证据清单.json"
    assert any(item["path"] == "task_1_截图证据清单.json" for item in manifest["files"])
    screenshot_manifest = __import__("json").loads(screenshot_manifest_path.read_text(encoding="utf-8"))
    assert screenshot_manifest["screenshot_count"] == 1
    assert screenshot_manifest["file_count"] == 1
    screenshot_record = screenshot_manifest["screenshots"][0]
    assert screenshot_record["url"] == "https://www.example.cn/"
    assert screenshot_record["host"] == "www.example.cn"
    assert screenshot_record["package_path"] == "screenshots/1_example.png"
    assert screenshot_record["source_path"].endswith("1_example.png")
    assert screenshot_record["system_name"] == "示例门户"
    screenshot_manifest_record = next(item for item in manifest["files"] if item["path"] == "screenshots/1_example.png")
    assert screenshot_record["package_size"] == screenshot_manifest_record["size"]
    assert screenshot_record["package_sha256"] == screenshot_manifest_record["sha256"]
    assert DeliveryPackageVerifier().verify(package.package_dir).status == "PASS"
    assert DeliveryPackageVerifier().verify(package.zip_path).status == "PASS"

    screenshot_manifest["screenshots"][0]["package_sha256"] = "0" * 64
    screenshot_manifest_path.write_text(__import__("json").dumps(screenshot_manifest, ensure_ascii=False), encoding="utf-8")
    screenshot_broken = DeliveryPackageVerifier().verify(package.package_dir)
    assert screenshot_broken.status == "FAIL"
    assert any("截图证据清单哈希" in item for item in screenshot_broken.failures)
    screenshot_manifest["screenshots"][0]["package_sha256"] = screenshot_manifest_record["sha256"]
    screenshot_manifest_path.write_text(__import__("json").dumps(screenshot_manifest, ensure_ascii=False), encoding="utf-8")

    report_ai_audit_path = package.package_dir / "task_1_报告AI分析审计.json"
    report_ai_audit_original = report_ai_audit_path.read_bytes()
    report_ai_audit = __import__("json").loads(report_ai_audit_original.decode("utf-8"))
    report_ai_audit["sections"][0].pop("usage_keys", None)
    report_ai_audit_path.write_text(__import__("json").dumps(report_ai_audit, ensure_ascii=False), encoding="utf-8")
    report_ai_broken = DeliveryPackageVerifier().verify(package.package_dir)
    assert report_ai_broken.status == "FAIL"
    assert any("报告AI分析审计存在不完整分块记录" in item for item in report_ai_broken.failures)
    report_ai_audit_path.write_bytes(report_ai_audit_original)

    asset_package_path = package.package_dir / "task_1_资产汇总.xlsx"
    broken_wb = load_workbook(asset_package_path)
    del broken_wb["管理驾驶舱"]
    broken_wb.save(asset_package_path)
    structural_broken = DeliveryPackageVerifier().verify(package.package_dir)
    assert structural_broken.status == "FAIL"
    assert any("管理驾驶舱" in item for item in structural_broken.failures)

    manifest["files"] = [item for item in manifest["files"] if "补全计划" not in item["path"]]
    broken_dir = tmp_path / "broken_delivery"
    broken_dir.mkdir()
    (broken_dir / "manifest.json").write_text(__import__("json").dumps(manifest, ensure_ascii=False), encoding="utf-8")
    broken = DeliveryPackageVerifier().verify(broken_dir)
    assert broken.status == "FAIL"
    assert any("补全计划" in item for item in broken.failures)


def test_report_generation_marks_artifact_write_failure(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    config.ai.enabled = False
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    session.add(ScanTask(id=1, target="示例集团有限公司", status="completed"))
    session.commit()
    service = ReportService(session, config)
    service._write_asset_workbook = lambda path, context: (_ for _ in ()).throw(OSError("disk full"))  # type: ignore[method-assign]

    with pytest.raises(OSError, match="disk full"):
        service.run(1, output_dir=tmp_path / "reports")

    generation = session.exec(select(ReportGenerationTask).where(ReportGenerationTask.scan_task_id == 1)).one()
    assert generation.status == "failed"
    assert generation.error_message == "disk full"


def test_package_build_failure_keeps_previous_delivery(tmp_path: Path, monkeypatch):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    task = ScanTask(id=1, target="示例集团有限公司", status="completed")
    session.add(task)
    session.commit()
    output_dir = tmp_path / "deliveries"
    package_dir = output_dir / "task_1_示例集团有限公司"
    package_dir.mkdir(parents=True)
    (package_dir / "previous.txt").write_text("keep", encoding="utf-8")
    zip_path = package_dir.with_suffix(".zip")
    zip_path.write_bytes(b"previous zip")

    class FakeQuality:
        def __init__(self, *args, **kwargs):
            pass

        def check(self, *args, **kwargs):
            return type("Quality", (), {"failures": [], "warnings": [], "lines": ["ok"], "status": "PASS"})()

        def _report_paths(self, *args, **kwargs):
            return {"Word报告": tmp_path / "source.docx"}

    monkeypatch.setattr("assetmap.services.delivery.package.DeliveryQualityService", FakeQuality)
    service = DeliveryPackageService(session, config)
    service._copy = lambda source, destination: (_ for _ in ()).throw(OSError("copy failed"))  # type: ignore[method-assign]

    with pytest.raises(OSError, match="copy failed"):
        service.package(1, reports_dir=tmp_path / "reports", output_dir=output_dir)

    assert (package_dir / "previous.txt").read_text(encoding="utf-8") == "keep"
    assert zip_path.read_bytes() == b"previous zip"
    assert not list(output_dir.glob(".task_1_示例集团有限公司.building-*"))


def test_workbook_sanitizes_tool_control_characters(tmp_path: Path):
    service = ReportService(None, AppConfig())  # type: ignore[arg-type]
    workbook_path = tmp_path / "report.xlsx"

    result = service._write_workbook(
        workbook_path,
        {
            "交付审计文件": [
                {
                    "工具失败": "dnsx=\x1b[[1;31mFatal\x1b[0m] DNS 查询超时\x00\x08",
                }
            ]
        },
    )

    sheet = load_workbook(result)["交付审计文件"]
    assert sheet["A2"].value == "dnsx=Fatal DNS 查询超时"


def test_quality_uses_latest_fallback_report_artifacts(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    service = DeliveryQualityService(None, config)  # type: ignore[arg-type]
    task = ScanTask(id=9, target="示例集团有限公司", status="completed")
    root = tmp_path / "reports" / "task_9_示例集团有限公司"
    root.mkdir(parents=True)
    preferred = root / "task_9_互联网资产暴露面测绘报告.docx"
    fallback = root / "task_9_互联网资产暴露面测绘报告_20260601_120000_000000.docx"
    preferred.write_bytes(b"old")
    fallback.write_bytes(b"new")
    os.utime(preferred, (1000, 1000))
    os.utime(fallback, (2000, 2000))

    paths = service._report_paths(task, tmp_path / "reports")

    assert paths["Word报告"] == fallback


def test_package_verifier_accepts_timestamped_report_artifacts():
    verifier = DeliveryPackageVerifier()
    records = [
        {"path": "task_9_互联网资产暴露面测绘报告_20260601_120000_000000.docx"},
        {"path": "task_9_资产汇总_20260601_120000_000000.xlsx"},
        {"path": "task_9_Web资产详情_20260601_120000_000000.xlsx"},
        {"path": "quality_summary.txt"},
        {"path": "task_9_待补充资产模板.yaml"},
        {"path": "task_9_复核工作单.yaml"},
        {"path": "task_9_补全计划.json"},
        {"path": "task_9_补全计划.txt"},
        {"path": "交付说明.txt"},
    ]
    failures: list[str] = []

    verifier._check_required_delivery_files({"task_id": 9}, records, failures)

    assert failures == []
    assert verifier._find_record_path(records, "互联网资产暴露面测绘报告") == records[0]["path"]


def test_workbook_headers_include_late_review_columns(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    service = ReportService(None, config)  # type: ignore[arg-type]
    output = tmp_path / "late_columns.xlsx"

    service._write_workbook(
        output,
        {
            "DNS复核清单": [
                {"根域名": "a.example.cn", "复核优先级": "中"},
                {"根域名": "b.example.cn", "复核优先级": "无", "人工复核状态": "confirmed"},
            ]
        },
    )

    headers = [cell.value for cell in next(load_workbook(output).active.iter_rows(min_row=1, max_row=1))]

    assert "人工复核状态" in headers


def test_report_markdown_tables_render_as_word_tables(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)
    doc = Document()

    service._add_paragraphs(
        doc,
        """
### 风险表
| 风险等级 | 资产 | 建议 |
| :--- | --- | ---: |
| 高 | `1.1.1.1:3306` | **关闭公网访问** |
| 中 | 门户系统 | 增加强认证 |
| 低 | 截断行 | 仍应进入表格

后续说明
""",
    )

    assert len(doc.tables) == 1
    assert [cell.text for cell in doc.tables[0].rows[0].cells] == ["风险等级", "资产", "建议"]
    assert doc.tables[0].rows[1].cells[1].text == "1.1.1.1:3306"
    assert doc.tables[0].rows[3].cells[2].text == "仍应进入表格"
    assert not any(paragraph.text.strip().startswith("|") for paragraph in doc.paragraphs)


def test_report_ai_analysis_cache_refreshes_when_payload_changes(tmp_path: Path, monkeypatch):
    config = AppConfig(
        database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"),
        ai=AiConfig(enabled=True, model="mock-model"),
    )
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)
    old_prompt = service._analysis_prompt("端口与服务暴露分析", {"stats": {"资产数量": 1}})
    session.add(
        AiAnalysis(
            scan_task_id=1,
            analysis_type="report_ports",
            status="completed",
            model="mock-model",
            prompt_json={**old_prompt, "_cache_key": service._analysis_cache_key(old_prompt)},
            response_json={"choices": [{"message": {"content": "旧结论"}}]},
            summary="旧结论",
        )
    )
    session.commit()
    calls = []

    def fake_chat_completion(*args, **kwargs):
        calls.append((args, kwargs))
        return {"choices": [{"message": {"content": "新结论"}}]}

    monkeypatch.setattr("assetmap.services.delivery.report.chat_completion", fake_chat_completion)

    refreshed = service._analysis(1, "report_ports", "端口与服务暴露分析", {"stats": {"资产数量": 2}}, rerun_ai=False)
    cached = service._analysis(1, "report_ports", "端口与服务暴露分析", {"stats": {"资产数量": 2}}, rerun_ai=False)
    row = session.get(AiAnalysis, 1)

    assert refreshed == "新结论"
    assert cached == "新结论"
    assert len(calls) == 1
    assert row is not None
    assert row.prompt_json["_cache_key"] == service._analysis_cache_key(service._analysis_prompt("端口与服务暴露分析", {"stats": {"资产数量": 2}}))


def test_report_ai_analysis_retries_when_output_truncated(tmp_path: Path, monkeypatch):
    config = AppConfig(
        database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"),
        ai=AiConfig(enabled=True, model="mock-model"),
    )
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)
    calls = []

    def fake_chat_completion(*args, **kwargs):
        calls.append(kwargs)
        if len(calls) == 1:
            return {"choices": [{"finish_reason": "length", "message": {"content": "未完成"}}]}
        return {"choices": [{"finish_reason": "stop", "message": {"content": "完整结论"}}]}

    monkeypatch.setattr("assetmap.services.delivery.report.chat_completion", fake_chat_completion)

    summary = service._analysis(1, "report_ports", "端口与服务暴露分析", {"stats": {"资产数量": 2}}, rerun_ai=False)
    row = session.exec(select(AiAnalysis).where(AiAnalysis.analysis_type == "report_ports")).one()

    assert summary == "完整结论"
    assert [call["max_completion_tokens"] for call in calls] == [1800, 4096]
    assert row.response_json["assetmap_retry"]["reason"] == "finish_reason=length"


def test_improvement_plan_exports_actionable_next_steps(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    task = ScanTask(id=1, target="示例集团有限公司", status="completed")
    covered = Company(id=1, name="示例集团有限公司", normalized_name="示例集团有限公司")
    uncovered = Company(id=2, name="示例子公司", normalized_name="示例子公司")
    domain = InternetAsset(
        id=1,
        asset_type="icp_domain",
        normalized_identifier="example.cn",
        display_name="example.cn",
        raw_payload={},
    )
    session.add(task)
    session.add(covered)
    session.add(uncovered)
    session.add(domain)
    session.add(
        CompanyEdge(
            task_id=1,
            parent_company_id=1,
            child_company_id=2,
            direct_holding_ratio=1,
            cumulative_holding_ratio=1,
            depth=1,
            path="示例集团有限公司 > 示例子公司",
        )
    )
    session.add(CompanyAssetLink(task_id=1, company_id=1, asset_id=1, source_tool="manual", raw_payload={}))
    session.commit()

    result = ImprovementPlanService(session, config).write(
        1,
        output_dir=tmp_path / "improve",
        reports_dir=tmp_path / "reports",
    )

    assert result.json_path.exists()
    assert result.text_path.exists()
    payload = __import__("json").loads(result.json_path.read_text(encoding="utf-8"))
    assert payload["quality"]["status"] in {"FAIL", "WARN"}
    asset_action = next(action for action in payload["actions"] if action["phase"] == "企业/备案资产")
    assert "高/中优先级 1 家" in asset_action["reason"]
    assert asset_action["samples"][0] == "示例子公司[高]"
    assert any("asset-gap-template 1 --priority high-medium" in action["command"] for action in payload["actions"])
    assert "下一轮补全计划" in result.text_path.read_text(encoding="utf-8")


def test_improvement_plan_visual_retry_policy(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ImprovementPlanService(session, config)

    assert not service._visual_retry_needed([{"识别方式": "http_probe_fallback", "分析错误": ""}])
    assert service._visual_retry_needed([{"识别方式": "", "分析错误": ""}])
    assert service._visual_retry_needed([{"识别方式": "screenshot_ai", "分析错误": "timeout"}])


def test_improvement_plan_splits_manual_dns_and_port_active_validation(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ImprovementPlanService(session, config)

    assert not service._dns_rerun_needed(
        [{"环节": "子域名/DNS", "指标": "DNS 解析复核质量"}],
        [{"复核类型": "no_public_ip+third_party_cname"}],
    )
    assert service._dns_rerun_needed(
        [{"环节": "子域名/DNS", "指标": "DNS 解析复核质量"}],
        [{"复核类型": "tool_failure+no_public_ip"}],
    )
    assert service._dns_rerun_needed(
        [{"环节": "子域名/DNS", "指标": "子域名枚举质量"}],
        [],
    )
    assert service._port_active_validation_needed(
        [{"环节": "端口发现", "指标": "端口证据来源质量"}]
    )


def test_package_copies_port_audit_files(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    source_dir = tmp_path / "data" / "nmap" / "task_1"
    source_dir.mkdir(parents=True)
    (source_dir / "target_sources.json").write_text('{"merged_count": 1}', encoding="utf-8")
    (source_dir / "fofa_errors.json").write_text('{"error_count": 1}', encoding="utf-8")
    package_dir = tmp_path / "deliveries" / "task_1"
    package_dir.mkdir(parents=True)

    copied = DeliveryPackageService(session, config)._copy_port_audit_files(1, package_dir)

    assert {path.name for path in copied} == {"task_1_端口目标来源.json", "task_1_FOFA失败记录.json"}
    assert (package_dir / "task_1_端口目标来源.json").exists()
    assert (package_dir / "task_1_FOFA失败记录.json").exists()


def test_package_copies_subdomain_audit_files(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    source_dir = tmp_path / "data" / "subdomains" / "task_1"
    source_dir.mkdir(parents=True)
    (source_dir / "subdomain_audit.json").write_text('{"subdomain_count": 1}', encoding="utf-8")
    package_dir = tmp_path / "deliveries" / "task_1"
    package_dir.mkdir(parents=True)

    copied = DeliveryPackageService(session, config)._copy_subdomain_audit_files(1, package_dir)

    assert {path.name for path in copied} == {"task_1_子域名DNS审计.json"}
    assert (package_dir / "task_1_子域名DNS审计.json").exists()


def test_package_copies_classification_audit_files(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    source_dir = tmp_path / "data" / "classify" / "task_1"
    source_dir.mkdir(parents=True)
    (source_dir / "web_probe_audit.json").write_text('{"total": 1}', encoding="utf-8")
    (source_dir / "service_classification_audit.json").write_text('{"total": 1}', encoding="utf-8")
    package_dir = tmp_path / "deliveries" / "task_1"
    package_dir.mkdir(parents=True)

    copied = DeliveryPackageService(session, config)._copy_classification_audit_files(1, package_dir)

    assert {path.name for path in copied} == {"task_1_HTTP探测审计.json", "task_1_服务分类审计.json"}
    assert (package_dir / "task_1_HTTP探测审计.json").exists()
    assert (package_dir / "task_1_服务分类审计.json").exists()


def test_package_copies_visual_audit_files(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    source_dir = tmp_path / "data" / "url_discovery" / "task_1"
    source_dir.mkdir(parents=True)
    (source_dir / "visual_analysis_audit.json").write_text('{"total": 1}', encoding="utf-8")
    package_dir = tmp_path / "deliveries" / "task_1"
    package_dir.mkdir(parents=True)

    copied = DeliveryPackageService(session, config)._copy_visual_audit_files(1, package_dir)

    assert {path.name for path in copied} == {"task_1_视觉识别审计.json"}
    assert (package_dir / "task_1_视觉识别审计.json").exists()


def test_package_copies_report_ai_audit_files(tmp_path: Path, monkeypatch):
    monkeypatch.chdir(tmp_path)
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    source_dir = tmp_path / "data" / "report" / "task_1"
    source_dir.mkdir(parents=True)
    (source_dir / "report_ai_audit.json").write_text('{"section_count": 4}', encoding="utf-8")
    package_dir = tmp_path / "deliveries" / "task_1"
    package_dir.mkdir(parents=True)

    copied = DeliveryPackageService(session, config)._copy_report_audit_files(1, package_dir)

    assert {path.name for path in copied} == {"task_1_报告AI分析审计.json"}
    assert (package_dir / "task_1_报告AI分析审计.json").exists()


def test_report_risk_rows_identify_sensitive_services(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._risk_rows(
        [
            {
                "单位": "示例集团有限公司",
                "IP": "8.8.8.8",
                "端口": 3389,
                "服务": "ms-wbt-server",
                "产品": "",
                "版本": "",
                "Web标题": "",
                "Web URL": "",
            }
        ],
        [
            {
                "单位": "示例集团有限公司",
                "URL": "https://vpn.example.cn/",
                "AI识别系统": "EASY CONNECT",
                "HTML标题": "",
                "网站用途": "VPN 远程接入",
                "登录特征": "登录页",
            }
        ],
    )

    assert [row["风险等级"] for row in rows] == ["高", "高"]
    assert {row["风险类型"] for row in rows} == {"敏感服务暴露", "远程接入 Web 入口"}


def test_visual_review_rows_flag_fallback_and_low_confidence(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._visual_review_rows(
        [
            {
                "单位": "示例集团有限公司",
                "URL": "https://www.example.cn/",
                "IP": "203.0.113.10",
                "端口": 443,
                "HTTP状态": 200,
                "HTML标题": "示例门户",
                "识别方式": "screenshot_ai",
                "识别置信度": 0.9,
                "AI识别系统": "示例门户",
                "网站用途": "对外门户",
                "截图": "ok.png",
                "分析错误": "",
            },
            {
                "单位": "示例集团有限公司",
                "URL": "https://slow.example.cn/",
                "IP": "203.0.113.11",
                "端口": 443,
                "HTTP状态": 200,
                "HTML标题": "慢页面",
                "识别方式": "http_probe_fallback",
                "识别置信度": "",
                "AI识别系统": "",
                "网站用途": "HTTP探测摘要",
                "截图": "",
                "分析错误": "",
            },
            {
                "单位": "示例集团有限公司",
                "URL": "https://unclear.example.cn/",
                "IP": "203.0.113.12",
                "端口": 443,
                "HTTP状态": 200,
                "HTML标题": "低置信度页面",
                "识别方式": "screenshot_ai",
                "识别置信度": 0.3,
                "AI识别系统": "",
                "网站用途": "",
                "截图": "unclear.png",
                "分析错误": "",
            },
        ]
    )

    assert [row["URL"] for row in rows] == ["https://unclear.example.cn/", "https://slow.example.cn/"]
    assert rows[0]["复核优先级"] == "中"
    assert rows[0]["复核类型"] == "manual_low_confidence_review"
    assert rows[1]["复核优先级"] == "低"
    assert rows[1]["复核类型"] == "manual_http_fallback_review"
    assert "HTTP探测信息降级识别" in rows[1]["复核原因"]


def test_visual_review_rows_lower_priority_for_error_pages(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._visual_review_rows(
        [
            {
                "单位": "示例集团有限公司",
                "URL": "https://forbidden.example.cn/",
                "IP": "203.0.113.10",
                "端口": 443,
                "HTTP状态": 403,
                "HTML标题": "403 Forbidden",
                "识别方式": "screenshot_ai",
                "识别置信度": 0.0,
                "AI识别系统": "403 Forbidden",
                "网站用途": "",
                "截图": "403.png",
                "分析错误": "",
            }
        ]
    )

    assert rows[0]["复核优先级"] == "低"
    assert rows[0]["复核类型"] == "manual_low_value_page_review"
    assert "低价值错误/拦截页面" in rows[0]["复核原因"]
    assert "错误页、默认页、拦截页或停放页" in rows[0]["建议动作"]


def test_visual_review_rows_lower_priority_for_named_http_fallback(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._visual_review_rows(
        [
            {
                "单位": "示例集团有限公司",
                "URL": "https://mail.example.cn/",
                "IP": "203.0.113.10",
                "端口": 443,
                "HTTP状态": 200,
                "HTML标题": "example.cn - 邮箱用户登录",
                "识别方式": "http_probe_fallback",
                "识别置信度": 0.35,
                "AI识别系统": "example.cn - 邮箱用户登录",
                "网站用途": "HTTP探测摘要",
                "截图": "",
                "分析错误": "",
            }
        ]
    )

    assert rows[0]["复核优先级"] == "低"
    assert rows[0]["复核类型"] == "manual_http_fallback_review"


def test_web_rows_enrich_empty_visual_result_from_title(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._web_rows(
        {
            "service_assets": [],
            "web_entrypoints": [
                {
                    "normalized_url": "https://portal.example.cn/",
                    "host": "portal.example.cn",
                    "target_ip": "203.0.113.10",
                    "port": 443,
                    "http_status": 200,
                    "title": "统一门户",
                    "evidence": {"visual_analysis": {"analysis_method": "screenshot_ai", "confidence": 0}},
                }
            ],
        },
        {"portal.example.cn": "示例集团有限公司"},
        {},
        {},
    )

    assert rows[0]["AI识别系统"] == "统一门户"
    assert rows[0]["识别置信度"] == 0.55


def test_web_rows_enrich_empty_fallback_from_service_product_and_reason(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._web_rows(
        {
            "service_assets": [
                {
                    "id": 7,
                    "target_ip": "203.0.113.10",
                    "port": 9080,
                    "asset_kind": "web",
                    "host_mode": "passive_fofa",
                    "product": "WebSphere Application Server/6.1",
                    "representative_url": "http://203.0.113.10:9080/",
                }
            ],
            "web_entrypoints": [
                {
                    "service_asset_id": 7,
                    "normalized_url": "http://203.0.113.10:9080/",
                    "host": "203.0.113.10",
                    "target_ip": "203.0.113.10",
                    "port": 9080,
                    "evidence": {
                        "source": "service_asset",
                        "visual_analysis": {
                            "analysis_method": "http_probe_fallback",
                            "confidence": 0.35,
                            "screenshot_error": "Page.goto: net::ERR_EMPTY_RESPONSE",
                        },
                    },
                }
            ],
        },
        {},
        {"203.0.113.10": "示例集团有限公司"},
        {},
    )

    assert rows[0]["AI识别系统"] == "WebSphere Application Server/6.1"
    assert rows[0]["降级类型"] == "被动FOFA证据"
    assert "ERR_EMPTY_RESPONSE" in rows[0]["降级原因"]


def test_screenshot_evidence_rows_keep_all_visual_results(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)
    web_rows = [
        {
            "单位": "示例集团有限公司",
            "URL": f"https://site{index}.example.cn/",
            "AI识别系统": f"系统{index}",
            "HTML标题": f"系统{index}",
            "网站用途": "业务系统",
            "识别方式": "screenshot_ai",
            "识别置信度": 0.9,
            "截图": f"shot{index}.png",
            "分析错误": "",
        }
        for index in range(60)
    ]

    rows = service._screenshot_evidence_rows(web_rows, [])

    assert len(rows) == 60
    assert {row["URL"] for row in rows} == {row["URL"] for row in web_rows}


def test_workbook_embeds_all_screenshot_thumbnails(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)
    png_bytes = base64.b64decode(
        "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO+/p9sAAAAASUVORK5CYII="
    )
    rows = []
    for index in range(25):
        screenshot = tmp_path / f"shot_{index}.png"
        screenshot.write_bytes(png_bytes)
        rows.append({"缩略图": "", "URL": f"https://site{index}.example.cn/", "截图文件": str(screenshot)})

    workbook_path = service._write_workbook(tmp_path / "web.xlsx", {"截图证据": rows})
    sheet = load_workbook(workbook_path)["截图证据"]

    assert sheet.max_row == 26
    assert len(getattr(sheet, "_images", [])) == 25


def test_service_audit_rows_include_visual_summary(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._service_audit_rows(
        {
            "service_assets": [
                {
                    "id": 1,
                    "target_ip": "203.0.113.10",
                    "protocol": "tcp",
                    "port": 443,
                    "asset_kind": "web",
                    "host_mode": "virtual_host",
                    "representative_url": "https://portal.example.cn/",
                    "domains": ["portal.example.cn"],
                }
            ],
            "web_probe_results": [],
            "web_entrypoints": [
                {
                    "service_asset_id": 1,
                    "normalized_url": "https://portal.example.cn/",
                    "target_ip": "203.0.113.10",
                    "port": 443,
                },
                {
                    "service_asset_id": 1,
                    "normalized_url": "https://portal.example.cn/admin",
                    "target_ip": "203.0.113.10",
                    "port": 443,
                },
            ],
        },
        {"portal.example.cn": "示例集团有限公司"},
        {},
        {},
        web_rows=[
            {
                "IP": "203.0.113.10",
                "端口": 443,
                "AI识别系统": "统一门户",
                "网站用途": "业务入口",
                "截图": "portal.png",
            },
            {
                "IP": "203.0.113.10",
                "端口": 443,
                "AI识别系统": "后台管理",
                "网站用途": "管理入口",
                "截图": "admin.png",
            },
        ],
    )

    assert rows[0]["URL入口数量"] == 2
    assert rows[0]["视觉识别数量"] == 2
    assert rows[0]["截图数量"] == 2
    assert rows[0]["AI识别系统"] == "统一门户；后台管理"
    assert rows[0]["网站用途"] == "业务入口；管理入口"


def test_visual_review_rows_explain_fallback_category_and_mojibake(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._visual_review_rows(
        [
            {
                "单位": "示例集团有限公司",
                "URL": "https://203.0.113.10/",
                "IP": "203.0.113.10",
                "端口": 443,
                "HTTP状态": 200,
                "HTML标题": "���ݾ�������Ƽ�",
                "识别方式": "http_probe_fallback",
                "识别置信度": 0.35,
                "AI识别系统": "ݾƼ",
                "网站用途": "HTTP探测摘要",
                "降级类型": "截图超时",
                "降级原因": "screenshot hard timeout after 60s",
                "截图": "",
                "分析错误": "",
            }
        ]
    )

    assert "HTTP探测信息降级识别(截图超时)" in rows[0]["复核原因"]
    assert "疑似编码乱码" in rows[0]["复核原因"]
    assert rows[0]["降级类型"] == "截图超时"


def test_dns_quality_rows_flag_tool_failures_and_third_party_cname(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._dns_quality_rows(
        {
            "assets": [{"asset_type": "icp_domain", "normalized_identifier": "example.cn"}],
            "subdomain_task": {"status": "completed"},
            "subdomains": [],
            "dns_records": [
                {"fqdn": "www.example.cn", "root_domain": "example.cn", "record_type": "CNAME", "value": "expired.hichina.com"},
            ],
            "subdomain_tool_runs": [
                {"root_domain": "example.cn", "tool_name": "dnsx", "status": "failed", "error_message": "timeout"}
            ],
        },
        {"example.cn": "示例集团有限公司"},
    )

    assert rows[0]["复核优先级"] == "中"
    assert rows[0]["复核类型"] == "no_subdomain+tool_failure+no_public_ip+third_party_cname"
    assert rows[0]["归属单位"] == "示例集团有限公司"
    assert "未发现子域名" in rows[0]["复核原因"]
    assert "子域名工具失败" in rows[0]["复核原因"]
    assert "expired.hichina.com" in rows[0]["第三方CNAME线索"]


def test_coverage_rows_include_dns_review_quality_gate(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._coverage_rows(
        {
            "companies": [],
            "assets": [],
            "subdomains": [],
            "dns_records": [],
            "subdomain_tool_runs": [],
            "service_assets": [],
            "nmap_ports": [],
        },
        [],
        [
            {
                "复核优先级": "中",
                "根域名": "no-a.example.cn",
                "复核原因": "未获得公网A/AAAA记录",
            },
            {
                "复核优先级": "低",
                "根域名": "cdn.example.cn",
                "复核原因": "存在第三方/CDN/停放CNAME",
            },
        ],
        [],
        [],
        [],
        [],
        [],
        [],
    )

    dns_review = next(row for row in rows if row["指标"] == "DNS 解析复核质量")
    assert dns_review["环节"] == "子域名/DNS"
    assert dns_review["缺口等级"] == "中"
    assert "中优先级 1 个" in dns_review["结果"]
    assert "no-a.example.cn" in dns_review["缺口样例"]


def test_port_target_rows_merge_sources_and_open_ports(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._port_target_rows(
        {
            "nmap_task": {"targets": ["8.8.8.8", "1.1.1.1"]},
            "dns_records": [
                {"fqdn": "www.example.cn", "root_domain": "example.cn", "record_type": "A", "value": "8.8.8.8", "raw_payload": {}},
                {"fqdn": "manual", "root_domain": "manual", "record_type": "A", "value": "1.1.1.1", "raw_payload": {"kind": "manual_ip"}},
            ],
            "assets": [{"asset_type": "ip", "normalized_identifier": "1.1.1.1"}],
            "ai_analyses": [{"analysis_type": "dns_inference", "status": "completed", "summary": "NMAP_TARGET_IPS: 8.8.8.8"}],
        },
        {"example.cn": "示例集团有限公司"},
        {"1.1.1.1": "手工单位"},
        {"8.8.8.8": "示例集团有限公司"},
        [{"单位": "示例集团有限公司", "IP": "8.8.8.8", "端口": 443}],
        [{"IP": "8.8.8.8", "URL": "https://www.example.cn/"}],
        [{"资产": "8.8.8.8:443", "风险等级": "高"}],
    )

    by_ip = {row["IP"]: row for row in rows}
    assert by_ip["8.8.8.8"]["目标来源"] == "ai, dns_public"
    assert by_ip["8.8.8.8"]["开放端口"] == "443"
    assert by_ip["8.8.8.8"]["高风险数量"] == 1
    assert by_ip["1.1.1.1"]["目标来源"] == "manual"
    assert by_ip["1.1.1.1"]["开放端口数量"] == 0


def test_port_rows_classify_active_passive_and_merged_evidence(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._port_rows(
        {
            "service_assets": [],
            "dns_records": [],
            "nmap_ports": [
                {
                    "target_ip": "8.8.8.8",
                    "protocol": "tcp",
                    "port": 22,
                    "state": "open",
                    "service": "ssh",
                    "product": "",
                    "version": "",
                    "raw_payload": {},
                },
                {
                    "target_ip": "8.8.8.8",
                    "protocol": "tcp",
                    "port": 443,
                    "state": "open",
                    "service": "https",
                    "product": "",
                    "version": "",
                    "raw_payload": {"source": "fofa", "host": "https://portal.example.cn/"},
                },
                {
                    "target_ip": "8.8.8.8",
                    "protocol": "tcp",
                    "port": 8443,
                    "state": "open",
                    "service": "https",
                    "product": "",
                    "version": "",
                    "raw_payload": {"sources": ["nmap", "fofa"], "fofa": {"host": "https://portal.example.cn:8443/"}},
                },
            ],
        },
        {},
        {},
        {},
    )

    by_port = {row["端口"]: row for row in rows}
    assert by_port[22]["证据类型"] == "主动Nmap"
    assert by_port[22]["主动扫描确认"] == "是"
    assert by_port[443]["证据类型"] == "被动FOFA"
    assert by_port[443]["被动FOFA证据"] == "是"
    assert by_port[8443]["证据类型"] == "主动+被动"


def test_coverage_rows_flag_passive_only_port_evidence(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._coverage_rows(
        {
            "companies": [],
            "assets": [],
            "subdomains": [],
            "dns_records": [],
            "subdomain_tool_runs": [],
            "service_assets": [],
            "nmap_ports": [],
        },
        [],
        [],
        [
            {
                "IP": "8.8.8.8",
                "端口": 443,
                "证据类型": "被动FOFA",
                "主动扫描确认": "否",
                "被动FOFA证据": "是",
            }
        ],
        [],
        [],
        [],
        [],
        [],
    )

    port_quality = next(row for row in rows if row["指标"] == "端口证据来源质量")
    assert port_quality["缺口等级"] == "低"
    assert "被动FOFA-only 1 个" in port_quality["结果"]
    assert "8.8.8.8:443" in port_quality["缺口样例"]


def test_service_audit_rows_flag_web_like_non_web(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._service_audit_rows(
        {
            "service_assets": [
                {
                    "id": 1,
                    "target_ip": "203.0.113.10",
                    "protocol": "tcp",
                    "port": 9080,
                    "asset_kind": "non_web",
                    "host_mode": "none",
                    "service": "tcp",
                    "product": "WebSphere Application Server",
                    "version": "6.1",
                    "domains": [],
                }
            ],
            "web_probe_results": [
                {"target_ip": "203.0.113.10", "port": 9080, "status": "failed", "error_message": "timeout"}
            ],
            "web_entrypoints": [],
        },
        {},
        {"203.0.113.10": "示例集团有限公司"},
        {},
    )

    assert rows[0]["复核优先级"] == "中"
    assert rows[0]["资产类型"] == "non_web"
    assert "疑似Web" in rows[0]["分类依据"]


def test_service_audit_rows_flag_passive_fofa_web_for_review(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._service_audit_rows(
        {
            "service_assets": [
                {
                    "id": 1,
                    "target_ip": "203.0.113.10",
                    "protocol": "tcp",
                    "port": 9080,
                    "asset_kind": "web",
                    "host_mode": "passive_fofa",
                    "service": "https",
                    "product": "nginx",
                    "representative_url": "https://portal.example.cn:9080/",
                    "domains": ["portal.example.cn"],
                    "title": "示例门户",
                }
            ],
            "web_probe_results": [
                {"target_ip": "203.0.113.10", "port": 9080, "status": "failed", "error_message": "timeout"}
            ],
            "web_entrypoints": [
                {
                    "service_asset_id": 1,
                    "target_ip": "203.0.113.10",
                    "port": 9080,
                    "normalized_url": "https://portal.example.cn:9080/",
                }
            ],
        },
        {"portal.example.cn": "示例集团有限公司"},
        {},
        {},
    )

    assert rows[0]["复核优先级"] == "低"
    assert rows[0]["复核类型"] == "passive_fofa_review"
    assert rows[0]["主机模式"] == "passive_fofa"
    assert "FOFA被动Web证据" in rows[0]["分类依据"]
    assert "人工确认 FOFA Host" in rows[0]["建议动作"]


def test_url_coverage_rows_flag_orphan_entrypoints(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._url_coverage_rows(
        {
            "service_assets": [
                {
                    "id": 1,
                    "target_ip": "203.0.113.10",
                    "port": 443,
                    "asset_kind": "web",
                    "host_mode": "ip_site",
                    "representative_url": "https://203.0.113.10/",
                }
            ],
            "web_entrypoints": [
                {"service_asset_id": 999, "target_ip": "203.0.113.11", "port": 443, "normalized_url": "https://orphan.example.cn/"}
            ],
        },
        [{"IP": "203.0.113.10", "端口": 443, "单位": "示例集团有限公司"}],
    )

    assert {row["覆盖结论"] for row in rows} == {"Web服务未生成URL入口", "URL入口未关联到服务资产"}


def test_url_coverage_rows_reuse_existing_representative_url_entry(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._url_coverage_rows(
        {
            "service_assets": [
                {
                    "id": 1,
                    "target_ip": "203.0.113.10",
                    "port": 80,
                    "asset_kind": "web",
                    "host_mode": "ip_site",
                    "representative_url": "https://portal.example.cn/",
                },
                {
                    "id": 2,
                    "target_ip": "203.0.113.10",
                    "port": 443,
                    "asset_kind": "web",
                    "host_mode": "ip_site",
                    "representative_url": "https://portal.example.cn/",
                },
            ],
            "web_entrypoints": [
                {"service_asset_id": 2, "target_ip": "203.0.113.10", "port": 443, "normalized_url": "https://portal.example.cn/"}
            ],
        },
        [
            {"IP": "203.0.113.10", "端口": 80, "单位": "示例集团有限公司"},
            {"IP": "203.0.113.10", "端口": 443, "单位": "示例集团有限公司"},
        ],
    )

    by_port = {row["端口"]: row for row in rows}
    assert by_port[80]["复核优先级"] == "无"
    assert by_port[80]["覆盖结论"] == "代表URL已由其他服务入口覆盖"
    assert by_port[80]["来源"] == "service_asset/web_probe/representative_url_reuse"


def test_service_audit_rows_count_reused_representative_url_entry(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._service_audit_rows(
        {
            "service_assets": [
                {
                    "id": 1,
                    "target_ip": "203.0.113.10",
                    "protocol": "tcp",
                    "port": 80,
                    "asset_kind": "web",
                    "host_mode": "ip_site",
                    "representative_url": "https://portal.example.cn/",
                },
                {
                    "id": 2,
                    "target_ip": "203.0.113.10",
                    "protocol": "tcp",
                    "port": 443,
                    "asset_kind": "web",
                    "host_mode": "ip_site",
                    "representative_url": "https://portal.example.cn/",
                },
            ],
            "web_probe_results": [],
            "web_entrypoints": [
                {"service_asset_id": 2, "target_ip": "203.0.113.10", "port": 443, "normalized_url": "https://portal.example.cn/"}
            ],
        },
        {},
        {"203.0.113.10": "示例集团有限公司"},
        {},
    )

    by_port = {row["端口"]: row for row in rows}
    assert by_port[80]["复核优先级"] == "无"
    assert by_port[80]["URL入口数量"] == 1
    assert by_port[80]["URL入口覆盖方式"] == "代表URL复用覆盖"


def test_plain_markdown_removes_report_markers(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    assert service._plain_markdown("**高风险** `https://example.cn`") == "高风险 https://example.cn"
    assert service._plain_markdown("`http://example.cn") == "http://example.cn"


def test_web_rows_use_manual_url_owner_when_no_service_link(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)
    rows = service._web_rows(
        {
            "service_assets": [],
            "web_entrypoints": [
                {
                    "normalized_url": "https://portal.example.cn/",
                    "host": "portal.example.cn",
                    "evidence": {
                        "manual_import": {"unit": "示例集团有限公司"},
                        "visual_analysis": {"system_name": "统一门户", "analysis_method": "manual"},
                    },
                }
            ],
        },
        {},
        {},
        {},
    )

    assert rows[0]["单位"] == "示例集团有限公司"
    assert rows[0]["AI识别系统"] == "统一门户"
    assert rows[0]["识别方式"] == "manual"


def test_review_workorder_exports_quality_review_items(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    task = ScanTask(id=1, target="示例集团有限公司", status="completed")
    company = Company(id=1, name="示例集团有限公司", normalized_name="示例集团有限公司")
    child = Company(id=2, name="核心子公司", normalized_name="核心子公司")
    asset = InternetAsset(
        id=1,
        asset_type="icp_domain",
        normalized_identifier="example.cn",
        display_name="example.cn",
        raw_payload={},
    )
    session.add(task)
    session.add(company)
    session.add(child)
    session.add(asset)
    session.add(CompanyAssetLink(task_id=1, company_id=1, asset_id=1, source_tool="manual", raw_payload={}))
    session.add(
        CompanyEdge(
            task_id=1,
            parent_company_id=1,
            child_company_id=2,
            direct_holding_ratio=1,
            cumulative_holding_ratio=1,
            depth=1,
            path="示例集团有限公司 > 核心子公司",
        )
    )
    session.add(DnsRecord(scan_task_id=1, fqdn="example.cn", root_domain="example.cn", record_type="CNAME", value="expired.hichina.com"))
    session.commit()

    output = tmp_path / "review.yaml"
    result = ReviewWorkOrderService(session, config).write(1, output)
    content = output.read_text(encoding="utf-8")

    assert result.total_items >= 1
    assert result.asset_items == 2
    assert result.dns_items >= 1
    assert "review_items:" in content
    assert "asset_supplement:" in content
    assert "核心子公司" in content
    assert "expired.hichina.com" in content
    assert "review_type:" in content
    assert "review_status: pending" in content
    assert "third_party_cname" in content
    assert "assetmap asset-gap-template 1 --priority high-medium" in content
    assert "assetmap import-review 1 --file data/review_workorder.task_1.yaml" in content
    assert "assetmap deliver 1" in content
    assert "assetmap run 1 --from-stage subdomains --rerun-subdomain-tools" not in content
    service = ReviewWorkOrderService(session, config)
    assert service._manual_url_candidates({"IP": "203.0.113.10", "端口": 8443}) == ["https://203.0.113.10:8443/"]
    assert service._manual_url_candidates({"IP": "203.0.113.10", "端口": 9000}) == [
        "http://203.0.113.10:9000/",
        "https://203.0.113.10:9000/",
    ]


def test_import_review_attestations_close_dns_review_items(tmp_path: Path):
    import yaml

    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    task = ScanTask(id=1, target="示例集团有限公司", status="completed")
    company = Company(id=1, name="示例集团有限公司", normalized_name="示例集团有限公司")
    asset = InternetAsset(
        id=1,
        asset_type="icp_domain",
        normalized_identifier="example.cn",
        display_name="example.cn",
        raw_payload={},
    )
    session.add(task)
    session.add(company)
    session.add(asset)
    session.add(CompanyAssetLink(task_id=1, company_id=1, asset_id=1, source_tool="manual", raw_payload={}))
    session.add(DnsRecord(scan_task_id=1, fqdn="example.cn", root_domain="example.cn", record_type="CNAME", value="expired.hichina.com"))
    session.commit()

    output = tmp_path / "review.yaml"
    ReviewWorkOrderService(session, config).write(1, output)
    payload = yaml.safe_load(output.read_text(encoding="utf-8"))
    payload["review_items"]["dns"][0]["review_status"] = "confirmed_no_business"
    payload["review_items"]["dns"][0]["review_notes"] = "人工确认该域名为停放页，无独立业务。"
    output.write_text(yaml.safe_dump(payload, sort_keys=False, allow_unicode=True), encoding="utf-8")

    result = ReviewImportService(session).run(1, output)
    context = ReportService(session, config)._context(ExportService(session)._bundle(1))
    dns_row = next(row for row in context["dns_quality_rows"] if row["根域名"] == "example.cn")

    assert result.imported == 1
    assert dns_row["复核优先级"] == "无"
    assert dns_row["人工复核状态"] == "confirmed_no_business"
    assert "人工复核" in dns_row["复核原因"]


def test_import_review_asset_supplement_no_assets_updates_unit_coverage(tmp_path: Path):
    import yaml

    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    task = ScanTask(id=1, target="Root Co", status="completed")
    root = Company(id=1, name="Root Co", normalized_name="root co")
    empty = Company(id=2, name="Empty Co", normalized_name="empty co")
    asset = InternetAsset(id=1, asset_type="icp_domain", normalized_identifier="root.example", display_name="root.example", raw_payload={})
    session.add(task)
    session.add(root)
    session.add(empty)
    session.add(asset)
    session.add(CompanyAssetLink(task_id=1, company_id=1, asset_id=1, source_tool="manual", raw_payload={}))
    session.add(
        CompanyEdge(
            task_id=1,
            parent_company_id=1,
            child_company_id=2,
            direct_holding_ratio=1,
            cumulative_holding_ratio=1,
            depth=1,
            path="Root Co > Empty Co",
        )
    )
    session.commit()

    output = tmp_path / "review.yaml"
    ReviewWorkOrderService(session, config).write(1, output)
    payload = yaml.safe_load(output.read_text(encoding="utf-8"))
    item = payload["review_items"]["asset_supplement"][0]
    item["review_status"] = "no_assets_found"
    item["review_notes"] = "备案、官网、公众号、小程序和应用商店均未发现独立互联网资产。"
    item["source_urls"] = ["https://beian.miit.gov.cn/"]
    output.write_text(yaml.safe_dump(payload, sort_keys=False, allow_unicode=True), encoding="utf-8")

    result = ReviewImportService(session).run(1, output)
    context = ReportService(session, config)._context(ExportService(session)._bundle(1))
    by_unit = {row["单位"]: row for row in context["unit_coverage_rows"]}

    assert result.imported == 1
    assert result.categories["asset_supplement"] == 1
    assert by_unit["Empty Co"]["覆盖状态"] == "人工确认无独立互联网资产"
    assert by_unit["Empty Co"]["复核优先级"] == "无"
    assert "备案、官网" in by_unit["Empty Co"]["人工复核依据"]


def test_import_visual_review_updates_web_asset_detail(tmp_path: Path):
    import yaml

    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    task = ScanTask(id=1, target="示例集团有限公司", status="completed")
    company = Company(id=1, name="示例集团有限公司", normalized_name="示例集团有限公司")
    service_asset = ServiceAsset(
        id=1,
        scan_task_id=1,
        target_ip="203.0.113.10",
        protocol="tcp",
        port=443,
        asset_kind="web",
        host_mode="ip",
        representative_url="https://203.0.113.10/",
    )
    session.add(task)
    session.add(company)
    session.add(service_asset)
    session.add(
        WebEntrypoint(
            scan_task_id=1,
            service_asset_id=1,
            target_ip="203.0.113.10",
            port=443,
            host="203.0.113.10",
            url="https://203.0.113.10/",
            normalized_url="https://203.0.113.10/",
            http_status=200,
            title="乱码标题",
            evidence={
                "visual_analysis": {
                    "analysis_method": "http_probe_fallback",
                    "confidence": 0.35,
                    "website_title": "乱码标题",
                }
            },
        )
    )
    session.commit()

    output = tmp_path / "review.yaml"
    ReviewWorkOrderService(session, config).write(1, output)
    payload = yaml.safe_load(output.read_text(encoding="utf-8"))
    item = payload["review_items"]["visual_identification"][0]
    item["review_status"] = "confirmed"
    item["confirmed_system_name"] = "统一身份认证平台"
    item["confirmed_site_purpose"] = "集团统一登录入口"
    item["confirmed_owner_unit"] = "示例集团有限公司"
    item["review_notes"] = "截图人工确认"
    output.write_text(yaml.safe_dump(payload, sort_keys=False, allow_unicode=True), encoding="utf-8")

    result = ReviewImportService(session).run(1, output)
    context = ReportService(session, config)._context(ExportService(session)._bundle(1))
    web_row = context["web_rows"][0]

    assert result.imported == 1
    assert web_row["AI识别系统"] == "统一身份认证平台"
    assert web_row["网站用途"] == "集团统一登录入口"
    assert web_row["识别方式"] == "manual_review"
    assert web_row["人工复核状态"] == "confirmed"
    assert context["visual_review_rows"] == []


def test_review_workorder_only_retries_missing_or_failed_visual_items(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReviewWorkOrderService(session, config)
    coverage_rows = [{"环节": "URL视觉识别", "缺口等级": "低"}]

    fallback_only = service._next_commands(
        1,
        coverage_rows,
        asset_items=[],
        dns_items=[],
        service_items=[],
        url_items=[],
        visual_items=[{"priority": "中", "identify_method": "http_probe_fallback", "analysis_error": ""}],
    )
    missing_visual = service._next_commands(
        1,
        coverage_rows,
        asset_items=[],
        dns_items=[],
        service_items=[],
        url_items=[],
        visual_items=[{"priority": "高", "identify_method": "", "analysis_error": ""}],
    )

    assert not any("url-discover" in command for command in fallback_only)
    assert any("url-discover 1 --retry-failed" in command for command in missing_visual)


def test_review_workorder_only_reruns_dns_for_rerunnable_gaps(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReviewWorkOrderService(session, config)

    manual_dns_only = service._next_commands(
        1,
        [{"环节": "子域名/DNS", "指标": "DNS 解析复核质量", "缺口等级": "中"}],
        asset_items=[],
        dns_items=[],
        service_items=[],
        url_items=[],
        visual_items=[],
    )
    rerunnable_dns = service._next_commands(
        1,
        [{"环节": "子域名/DNS", "指标": "子域名枚举质量", "缺口等级": "低"}],
        asset_items=[],
        dns_items=[],
        service_items=[],
        url_items=[],
        visual_items=[],
    )

    assert not any("--rerun-subdomain-tools" in command for command in manual_dns_only)
    assert any("--rerun-subdomain-tools" in command for command in rerunnable_dns)


def test_review_workorder_only_reruns_classify_for_rerunnable_service_gaps(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReviewWorkOrderService(session, config)

    passive_fofa_only = service._next_commands(
        1,
        [{"环节": "服务识别/URL", "指标": "服务分类复核", "缺口等级": "低"}],
        asset_items=[],
        dns_items=[],
        service_items=[{"review_type": "passive_fofa_review"}],
        url_items=[],
        visual_items=[],
    )
    missing_url = service._next_commands(
        1,
        [{"环节": "服务识别/URL", "指标": "服务分类复核", "缺口等级": "高"}],
        asset_items=[],
        dns_items=[],
        service_items=[{"review_type": "missing_url_entry"}],
        url_items=[],
        visual_items=[],
    )
    entry_coverage_gap = service._next_commands(
        1,
        [{"环节": "服务识别/URL", "指标": "Web 服务入口覆盖", "缺口等级": "高"}],
        asset_items=[],
        dns_items=[],
        service_items=[],
        url_items=[],
        visual_items=[],
    )

    assert not any("--rerun-classify" in command for command in passive_fofa_only)
    assert any("--rerun-classify" in command for command in missing_url)
    assert any("--rerun-classify" in command for command in entry_coverage_gap)


def test_review_workorder_visual_items_are_actionable(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReviewWorkOrderService(session, config)

    items = service._visual_items(
        [
            {
                "复核优先级": "中",
                "复核类型": "manual_http_fallback_review",
                "单位": "示例集团有限公司",
                "URL": "https://slow.example.cn/",
                "识别方式": "http_probe_fallback",
                "识别置信度": 0.35,
                "复核原因": "HTTP探测信息降级识别(截图超时)",
                "建议动作": "人工核对截图、标题和业务用途。",
                "截图": r"data\screenshots\task_1\1_slow.png",
                "分析错误": "",
            }
        ]
    )

    assert items[0]["review_type"] == "manual_http_fallback_review"
    assert items[0]["package_screenshot_path"] == "screenshots/1_slow.png"
    assert "confirmed_system_name" in items[0]["manual_result_fields"]

    low_value_items = service._visual_items(
        [
            {
                "复核优先级": "低",
                "单位": "示例集团有限公司",
                "URL": "https://forbidden.example.cn/",
                "识别方式": "screenshot_ai",
                "识别置信度": 0.45,
                "复核原因": "低价值错误/拦截页面置信度偏低(0.45)",
                "建议动作": "确认是否为错误页、默认页、拦截页或停放页。",
                "截图": r"data\screenshots\task_1\forbidden.png",
                "分析错误": "",
            }
        ]
    )
    assert low_value_items[0]["review_type"] == "manual_low_value_page_review"


def test_report_risk_rows_identify_easy_connect_with_space(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._risk_rows(
        [],
        [
            {
                "单位": "示例集团有限公司",
                "URL": "https://vpn.example.cn/",
                "AI识别系统": "EASY CONNECT",
                "HTML标题": "",
                "网站用途": "远程接入",
                "登录特征": "登录页",
            }
        ],
    )

    assert rows[0]["风险等级"] == "高"
    assert rows[0]["风险分值"] >= 90
    assert rows[0]["风险类型"] == "远程接入 Web 入口"
    assert "身份" in rows[0]["责任建议"]


def test_report_remediation_rows_add_sla_and_acceptance(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._remediation_rows(
        [
            {
                "风险等级": "高",
                "风险类型": "敏感服务暴露",
                "单位": "示例集团有限公司",
                "资产": "8.8.8.8:3389",
                "处置建议": "收敛访问源",
            }
        ]
    )

    assert rows[0]["优先级"] == "高"
    assert rows[0]["风险分值"] == 96
    assert "7 天" in rows[0]["建议时限"]
    assert rows[0]["当前状态"] == "待确认"
    assert "验收证据" in rows[0]


def test_unit_coverage_rows_count_unique_ips(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._unit_coverage_rows(
        {1: "示例集团有限公司"},
        {"示例集团有限公司": {"股权层级": 0, "直接持股": "100.00%", "累计持股": "100.00%", "股权路径": "示例集团有限公司"}},
        [
            {"单位": "示例集团有限公司", "资产类型": "ip", "资产标识": "8.8.8.8", "资产名称": ""},
        ],
        [
            {"单位": "示例集团有限公司", "主机名": "a.example.cn", "记录类型": "A", "记录值": "8.8.8.8"},
            {"单位": "示例集团有限公司", "主机名": "b.example.cn", "记录类型": "A", "记录值": "8.8.8.8"},
            {"单位": "示例集团有限公司", "主机名": "c.example.cn", "记录类型": "A", "记录值": "8.8.4.4"},
        ],
        [],
        [],
        [],
    )

    assert rows[0]["IP数量"] == 2
    assert rows[0]["子域名/DNS主机数量"] == 3
    assert rows[0]["复核优先级"] == "低"


def test_unit_coverage_rows_prioritize_missing_core_units(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    service = ReportService(session, config)

    rows = service._unit_coverage_rows(
        {1: "根集团", 2: "一级平台", 3: "二级控股", 4: "三级项目"},
        {
            "根集团": {"股权层级": 0, "直接持股": "100.00%", "累计持股": "100.00%", "股权路径": "根集团"},
            "一级平台": {"股权层级": 1, "直接持股": "100.00%", "累计持股": "100.00%", "股权路径": "根集团 > 一级平台"},
            "二级控股": {"股权层级": 2, "直接持股": "60.00%", "累计持股": "60.00%", "股权路径": "根集团 > 一级平台 > 二级控股", "子公司数量": 1},
            "三级项目": {"股权层级": 3, "直接持股": "100.00%", "累计持股": "60.00%", "股权路径": "根集团 > 一级平台 > 二级控股 > 三级项目"},
        },
        [],
        [],
        [],
        [],
        [],
    )

    by_unit = {row["单位"]: row for row in rows}
    assert by_unit["根集团"]["复核优先级"] == "高"
    assert by_unit["一级平台"]["复核优先级"] == "高"
    assert by_unit["二级控股"]["复核优先级"] == "中"
    assert by_unit["三级项目"]["复核优先级"] == "低"


def test_gap_template_lists_units_without_asset_lines(tmp_path: Path):
    config = AppConfig(database=DatabaseConfig(url=f"sqlite:///{tmp_path / 'assetmap.db'}"))
    engine = create_db_and_engine(config.database.url)
    session = get_session(engine)
    task = ScanTask(id=1, target="示例集团有限公司", status="completed")
    with_asset = Company(id=1, name="已有资产有限公司", normalized_name="已有资产有限公司")
    missing = Company(id=2, name="待补充资产有限公司", normalized_name="待补充资产有限公司")
    no_asset = Company(id=3, name="确认无资产有限公司", normalized_name="确认无资产有限公司")
    asset = InternetAsset(
        id=1,
        asset_type="icp_domain",
        normalized_identifier="example.cn",
        display_name="example.cn",
        raw_payload={},
    )
    session.add(task)
    session.add(with_asset)
    session.add(missing)
    session.add(no_asset)
    session.add(asset)
    session.add(CompanyAssetLink(task_id=1, company_id=1, asset_id=1, source_tool="manual", raw_payload={}))
    session.add(
        CompanyEdge(
            task_id=1,
            parent_company_id=1,
            child_company_id=2,
                direct_holding_ratio=1.0,
                cumulative_holding_ratio=1.0,
                depth=1,
                    path="已有资产有限公司 > 待补充资产有限公司",
                )
            )
    session.add(
        CompanyEdge(
            task_id=1,
            parent_company_id=1,
            child_company_id=3,
            direct_holding_ratio=1.0,
            cumulative_holding_ratio=1.0,
            depth=1,
            path="已有资产有限公司 > 确认无资产有限公司",
        )
    )
    session.add(
        SourceRawRecord(
            task_id=1,
            source="manual_import",
            action="no_assets_found",
            parameter_hash="confirmed-empty",
            request_payload={"unit": "确认无资产有限公司"},
            response_json={
                "unit": "确认无资产有限公司",
                "review_status": "no_assets_found",
                "source_urls": ["https://example.com/review"],
                "notes": "人工复核未发现独立互联网资产",
            },
        )
    )
    session.commit()

    output = tmp_path / "manual_assets.gaps.yaml"
    result = GapTemplateService(session, config).write(1, output)
    content = output.read_text(encoding="utf-8")

    assert result.units == 1
    assert "待补充资产有限公司" in content
    assert 'unit: "已有资产有限公司"' not in content
    assert "确认无资产有限公司" not in content
    assert "已有资产有限公司 > 待补充资产有限公司" in content
    assert "复核优先级" in content
    assert "缺口原因" in content
    assert "domains: []" in content
    assert "urls: []" in content
    assert "source_urls: []" in content
    assert "search_keywords:" in content
    assert '"待补充资产有限公司 官网"' in content
    assert "system_name: 示例门户" in content
    assert "review_status: pending" in content
    assert "assetmap run 1 --manual-file <本文件>" in content
    assert "assetmap import-assets 1 --file <本文件> --continue" in content
    assert "minimum_required" in content
    assert "review_checklist" in content
    assert "no_assets_found" in content
    assert "工信部ICP备案" in content
    assert "APP" in content
    assert "邮箱" in content

    output_partial = tmp_path / "manual_assets.partial.gaps.yaml"
    GapTemplateService(session, config).write(1, output_partial, include_partial=True)
    partial_content = output_partial.read_text(encoding="utf-8")
    assert partial_content.index('unit: "待补充资产有限公司"') < partial_content.index('unit: "已有资产有限公司"')
    assert "确认无资产有限公司" not in partial_content

    output_high = tmp_path / "manual_assets.high.gaps.yaml"
    result_high = GapTemplateService(session, config).write(1, output_high, priority_filter="high-medium")
    high_content = output_high.read_text(encoding="utf-8")
    assert result_high.units == 1
    assert "# 优先级过滤: high-medium" in high_content

    output_high_partial = tmp_path / "manual_assets.high.partial.gaps.yaml"
    result_high_partial = GapTemplateService(session, config).write(
        1,
        output_high_partial,
        include_partial=True,
        priority_filter="high-medium",
    )
    high_partial_content = output_high_partial.read_text(encoding="utf-8")
    assert result_high_partial.units == 2
    assert 'unit: "待补充资产有限公司"' in high_partial_content
    assert 'unit: "已有资产有限公司"' in high_partial_content
