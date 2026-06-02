from __future__ import annotations

import hashlib
import json
import shutil
import zipfile
from io import BytesIO
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from sqlmodel import Session, select

from assetmap.config import AppConfig
from assetmap.models import ScanTask, ServiceAsset, WebEntrypoint
from assetmap.services.gap_template import GapTemplateService
from assetmap.services.improvement_plan import ImprovementPlanService
from assetmap.services.quality import DeliveryQualityService, EXPECTED_ASSET_SHEETS, EXPECTED_WEB_SHEETS
from assetmap.services.report import _safe_name
from assetmap.services.review_workorder import ReviewWorkOrderService
from assetmap.services.asset_classifier import AssetClassifierService


@dataclass
class DeliveryPackageResult:
    package_dir: Path
    zip_path: Path
    manifest_path: Path
    quality_status: str
    files: list[dict]
    packaged_files: list[dict]


@dataclass
class PackageVerificationResult:
    status: str
    failures: list[str]
    warnings: list[str]
    lines: list[str]


class DeliveryPackageService:
    def __init__(self, session: Session, config: AppConfig) -> None:
        self.session = session
        self.config = config

    def package(
        self,
        task_id: int,
        *,
        reports_dir: Path | str = "reports",
        output_dir: Path | str = "deliveries",
        include_gap_template: bool = True,
        include_review_workorder: bool = True,
        include_partial_gaps: bool = True,
        strict: bool = False,
    ) -> DeliveryPackageResult:
        task = self.session.get(ScanTask, task_id)
        if not task:
            raise ValueError(f"Task not found: {task_id}")
        quality = DeliveryQualityService(self.session, self.config).check(task_id, output_dir=reports_dir)
        if quality.failures:
            raise ValueError("交付质量检查失败，不能打包: " + "; ".join(quality.failures))
        if strict and quality.warnings:
            raise ValueError("严格模式下存在质量警告，不能打包: " + "; ".join(quality.warnings))

        package_dir = Path(output_dir) / f"task_{task.id}_{_safe_name(task.target)}"
        if package_dir.exists():
            shutil.rmtree(package_dir)
        package_dir.mkdir(parents=True, exist_ok=True)

        report_paths = DeliveryQualityService(self.session, self.config)._report_paths(task, reports_dir)
        copied: list[Path] = []
        for source in report_paths.values():
            copied.append(self._copy(source, package_dir / source.name))

        quality_path = package_dir / "quality_summary.txt"
        quality_path.write_text("\n".join(quality.lines) + "\n", encoding="utf-8")
        copied.append(quality_path)

        if include_gap_template:
            gap_path = package_dir / f"task_{task.id}_待补充资产模板.yaml"
            GapTemplateService(self.session, self.config).write(
                task.id,
                gap_path,
                include_partial=include_partial_gaps,
                priority_filter="high-medium",
                force=True,
            )
            copied.append(gap_path)
        if include_review_workorder:
            review_path = package_dir / f"task_{task.id}_复核工作单.yaml"
            ReviewWorkOrderService(self.session, self.config).write(task.id, review_path, force=True)
            copied.append(review_path)
        plan = ImprovementPlanService(self.session, self.config).write(
            task.id,
            output_dir=package_dir,
            reports_dir=reports_dir,
        )
        plan_json = plan.json_path.rename(package_dir / f"task_{task.id}_补全计划.json")
        plan_text = plan.text_path.rename(package_dir / f"task_{task.id}_补全计划.txt")
        copied.extend([plan_json, plan_text])
        copied.extend(self._copy_subdomain_audit_files(task.id, package_dir))
        copied.extend(self._copy_port_audit_files(task.id, package_dir))
        copied.extend(self._copy_classification_audit_files(task.id, package_dir))
        copied.extend(self._copy_visual_audit_files(task.id, package_dir))
        copied.extend(self._copy_report_audit_files(task.id, package_dir))
        copied.extend(self._copy_screenshot_evidence(task.id, package_dir))
        readme_path = package_dir / "交付说明.txt"
        readme_path.write_text(self._delivery_readme(task, quality.status, quality.warnings, copied), encoding="utf-8")
        copied.append(readme_path)

        manifest_path = package_dir / "manifest.json"
        files = [self._file_record(path, package_dir) for path in copied]
        manifest = {
            "task_id": task.id,
            "target": task.target,
            "quality_status": quality.status,
            "warnings": quality.warnings,
            "files": files,
        }
        manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        packaged_files = [*files, self._file_record(manifest_path, package_dir)]

        zip_path = package_dir.with_suffix(".zip")
        if zip_path.exists():
            zip_path.unlink()
        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
            for path in [*copied, manifest_path]:
                archive.write(path, path.relative_to(package_dir).as_posix())
        return DeliveryPackageResult(
            package_dir=package_dir,
            zip_path=zip_path,
            manifest_path=manifest_path,
            quality_status=quality.status,
            files=files,
            packaged_files=packaged_files,
        )

    def _copy(self, source: Path, destination: Path) -> Path:
        destination.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(source, destination)
        return destination

    def _copy_subdomain_audit_files(self, task_id: int, package_dir: Path) -> list[Path]:
        source_dir = Path("data") / "subdomains" / f"task_{task_id}"
        candidates = [
            (source_dir / "subdomain_audit.json", f"task_{task_id}_子域名DNS审计.json"),
        ]
        copied = []
        for source, name in candidates:
            if source.exists() and source.stat().st_size > 0:
                copied.append(self._copy(source, package_dir / name))
        return copied

    def _copy_port_audit_files(self, task_id: int, package_dir: Path) -> list[Path]:
        source_dir = Path("data") / "nmap" / f"task_{task_id}"
        candidates = [
            (source_dir / "target_sources.json", f"task_{task_id}_端口目标来源.json"),
            (source_dir / "fofa_errors.json", f"task_{task_id}_FOFA失败记录.json"),
        ]
        copied = []
        for source, name in candidates:
            if source.exists() and source.stat().st_size > 0:
                copied.append(self._copy(source, package_dir / name))
        return copied

    def _copy_classification_audit_files(self, task_id: int, package_dir: Path) -> list[Path]:
        source_dir = Path("data") / "classify" / f"task_{task_id}"
        service_rows = self.session.exec(select(ServiceAsset).where(ServiceAsset.scan_task_id == task_id)).all()
        if service_rows:
            AssetClassifierService(self.session, self.config)._write_service_classification_audit(task_id, service_rows)
        candidates = [
            (source_dir / "web_probe_audit.json", f"task_{task_id}_HTTP探测审计.json"),
            (source_dir / "service_classification_audit.json", f"task_{task_id}_服务分类审计.json"),
        ]
        copied = []
        for source, name in candidates:
            if source.exists() and source.stat().st_size > 0:
                copied.append(self._copy(source, package_dir / name))
        return copied

    def _copy_visual_audit_files(self, task_id: int, package_dir: Path) -> list[Path]:
        source_dir = Path("data") / "url_discovery" / f"task_{task_id}"
        candidates = [
            (source_dir / "visual_analysis_audit.json", f"task_{task_id}_视觉识别审计.json"),
        ]
        copied = []
        for source, name in candidates:
            if source.exists() and source.stat().st_size > 0:
                copied.append(self._copy(source, package_dir / name))
        return copied

    def _copy_report_audit_files(self, task_id: int, package_dir: Path) -> list[Path]:
        source_dir = Path("data") / "report" / f"task_{task_id}"
        candidates = [
            (source_dir / "report_ai_audit.json", f"task_{task_id}_报告AI分析审计.json"),
        ]
        copied = []
        for source, name in candidates:
            if source.exists() and source.stat().st_size > 0:
                copied.append(self._copy(source, package_dir / name))
        return copied

    def _copy_screenshot_evidence(self, task_id: int, package_dir: Path) -> list[Path]:
        copied: list[Path] = []
        records: list[dict] = []
        source_to_package: dict[str, str] = {}
        used_names: set[str] = set()
        rows = self.session.exec(select(WebEntrypoint).where(WebEntrypoint.scan_task_id == task_id)).all()
        for row in rows:
            evidence = row.evidence or {}
            visual = evidence.get("visual_analysis") if isinstance(evidence.get("visual_analysis"), dict) else {}
            candidates = [
                visual.get("screenshot_path") if isinstance(visual, dict) else None,
                evidence.get("visual_analysis_screenshot_path"),
            ]
            source = self._first_existing_screenshot(candidates)
            if not source:
                continue
            source_key = str(source.resolve())
            package_path = source_to_package.get(source_key)
            if not package_path:
                package_path = self._screenshot_package_path(source, used_names)
                copied.append(self._copy(source, package_dir / package_path))
                source_to_package[source_key] = package_path
            records.append(
                {
                    "entry_id": row.id,
                    "url": row.normalized_url,
                    "final_url": row.final_url or visual.get("final_url") or row.normalized_url,
                    "host": row.host,
                    "ip": row.target_ip,
                    "port": row.port,
                    "http_status": row.http_status,
                    "title": row.title or visual.get("website_title"),
                    "system_name": visual.get("system_name"),
                    "site_purpose": visual.get("site_purpose"),
                    "analysis_method": visual.get("analysis_method"),
                    "source_path": str(source),
                    "package_path": package_path,
                    "package_size": source.stat().st_size,
                    "package_sha256": self._sha256(source),
                }
            )
        if records:
            manifest_path = package_dir / f"task_{task_id}_截图证据清单.json"
            manifest_path.write_text(
                json.dumps(
                    {
                        "task_id": task_id,
                        "screenshot_count": len(records),
                        "file_count": len(source_to_package),
                        "screenshots": records,
                    },
                    ensure_ascii=False,
                    indent=2,
                ),
                encoding="utf-8",
            )
            copied.append(manifest_path)
        return copied

    def _copy_screenshot_files(self, task_id: int, package_dir: Path) -> list[Path]:
        return self._copy_screenshot_evidence(task_id, package_dir)

    def _first_existing_screenshot(self, candidates: list[str | None]) -> Path | None:
        for value in candidates:
            if not value:
                continue
            path = Path(str(value))
            if path.exists() and path.is_file() and path.stat().st_size > 0:
                return path
        return None

    def _screenshot_package_path(self, source: Path, used_names: set[str]) -> str:
        name = source.name or f"{hashlib.sha256(str(source).encode('utf-8')).hexdigest()[:12]}.png"
        candidate = name
        index = 2
        while candidate in used_names:
            candidate = f"{source.stem}_{index}{source.suffix}"
            index += 1
        used_names.add(candidate)
        return f"screenshots/{candidate}"

    def _delivery_readme(self, task: ScanTask, quality_status: str, warnings: list[str], files: list[Path]) -> str:
        lines = [
            "互联网数字资产暴露面测绘交付说明",
            "",
            f"任务编号：{task.id}",
            f"测绘对象：{task.target}",
            f"质量状态：{quality_status}",
        ]
        if warnings:
            lines.append("质量警告：")
            lines.extend(f"- {item}" for item in warnings)
        lines.extend(
            [
                "",
                "交付内容：",
            ]
        )
        descriptions = {
            "互联网资产暴露面测绘报告.docx": "Word 主报告，包含摘要、范围、统计、风险、复核计划和整改建议。",
            "资产汇总.xlsx": "资产汇总附件，包含单位覆盖、资产、DNS、端口、风险、审计文件说明等台账。",
            "Web资产详情.xlsx": "Web 资产详情附件，包含重点 Web、截图证据、视觉识别、截图路径和复核清单。",
            "quality_summary.txt": "质量门禁摘要。",
            "待补充资产模板.yaml": "高/中优先级及部分覆盖单位的人工补充资产模板。",
            "复核工作单.yaml": "交付后复核工作单。",
            "补全计划.json": "下一轮补全计划的机器可读版本，包含质量状态、缺口和建议命令。",
            "补全计划.txt": "下一轮补全计划的人读版本，适合交付评审时直接查看。",
            "子域名DNS审计.json": "子域名枚举、DNS 解析、工具失败和覆盖缺口审计。",
            "端口目标来源.json": "端口扫描目标来源审计。",
            "HTTP探测审计.json": "HTTP 探测响应和失败审计。",
            "服务分类审计.json": "服务识别分类审计。",
            "视觉识别审计.json": "截图 AI、复用和降级识别审计。",
            "报告AI分析审计.json": "报告 AI 分块分析审计。",
            "截图证据清单.json": "截图证据映射清单，记录原始截图路径、包内截图路径、截图哈希、URL、主机、端口和视觉识别结论。",
            "screenshots": "原始网页截图证据文件，可与 Web 资产详情附件中的截图路径对应复核。",
        }
        for path in files:
            name = path.name
            if "screenshots" in path.parts:
                lines.append(f"- screenshots/{name}: {descriptions['screenshots']}")
                continue
            description = next((text for marker, text in descriptions.items() if marker in name), "交付证据文件。")
            lines.append(f"- {name}: {description}")
        lines.extend(
            [
                "",
                "校验方式：",
                f"- assetmap verify-package deliveries/task_{task.id}_{_safe_name(task.target)}.zip",
                "",
                "复测建议：",
                f"- assetmap improve {task.id}",
                f"- assetmap improve {task.id} --mode all --include-deliver --execute",
                f"- assetmap quality-check {task.id}",
                f"- assetmap review-workorder {task.id} --output data/review_workorder.task_{task.id}.yaml --force",
                f"- assetmap import-review {task.id} --file data/review_workorder.task_{task.id}.yaml",
            ]
        )
        return "\n".join(lines) + "\n"

    def _file_record(self, path: Path, root: Path) -> dict:
        return {
            "path": path.relative_to(root).as_posix(),
            "size": path.stat().st_size,
            "sha256": self._sha256(path),
        }

    def _sha256(self, path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()


class DeliveryPackageVerifier:
    def verify(self, package_path: Path | str) -> PackageVerificationResult:
        path = Path(package_path)
        failures: list[str] = []
        warnings: list[str] = []
        if not path.exists():
            failures.append(f"交付包不存在: {path}")
            return self._result(path, None, failures, warnings)
        try:
            if path.is_dir():
                manifest, file_records = self._verify_directory(path, failures, warnings)
            else:
                manifest, file_records = self._verify_zip(path, failures, warnings)
        except Exception as exc:
            failures.append(f"交付包无法校验: {str(exc)[:300]}")
            manifest, file_records = None, []
        return self._result(path, manifest, failures, warnings, file_records=file_records)

    def _verify_directory(self, package_dir: Path, failures: list[str], warnings: list[str]) -> tuple[dict | None, list[dict]]:
        manifest_path = package_dir / "manifest.json"
        if not manifest_path.exists():
            failures.append(f"缺少 manifest.json: {manifest_path}")
            return None, []
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        records = self._manifest_files(manifest)
        for record in records:
            path = package_dir / record["path"]
            if not path.exists():
                failures.append(f"缺少文件: {record['path']}")
                continue
            self._check_record(record, path.stat().st_size, self._sha256(path), failures)
        self._check_required_delivery_files(manifest, records, failures)
        self._check_delivery_structures(lambda name: (package_dir / name).read_bytes() if (package_dir / name).exists() else None, records, failures, warnings)
        return manifest, records

    def _verify_zip(self, zip_path: Path, failures: list[str], warnings: list[str]) -> tuple[dict | None, list[dict]]:
        with zipfile.ZipFile(zip_path) as archive:
            names = set(archive.namelist())
            if "manifest.json" not in names:
                failures.append("zip 缺少 manifest.json")
                return None, []
            manifest = json.loads(archive.read("manifest.json").decode("utf-8"))
            records = self._manifest_files(manifest)
            expected_names = {record["path"] for record in records} | {"manifest.json"}
            missing = sorted(expected_names - names)
            if missing:
                failures.append("zip 缺少文件: " + ", ".join(missing))
            extras = sorted(names - expected_names)
            if extras:
                warnings.append("zip 存在 manifest 未记录文件: " + ", ".join(extras))
            for record in records:
                if record["path"] not in names:
                    continue
                data = archive.read(record["path"])
                self._check_record(record, len(data), hashlib.sha256(data).hexdigest(), failures)
            self._check_required_delivery_files(manifest, records, failures)
            self._check_delivery_structures(
                lambda name: archive.read(name) if name in names else None,
                records,
                failures,
                warnings,
            )
            return manifest, records

    def _manifest_files(self, manifest: dict) -> list[dict]:
        files = manifest.get("files")
        if not isinstance(files, list):
            raise ValueError("manifest.json 中 files 字段不是列表")
        return [record for record in files if isinstance(record, dict)]

    def _check_required_delivery_files(self, manifest: dict, records: list[dict], failures: list[str]) -> None:
        task_id = manifest.get("task_id")
        required_markers = [
            "互联网资产暴露面测绘报告",
            "资产汇总",
            "Web资产详情",
            "quality_summary.txt",
            "待补充资产模板.yaml",
            "复核工作单.yaml",
            "补全计划.json",
            "补全计划.txt",
            "交付说明.txt",
        ]
        paths = [str(record.get("path") or "") for record in records]
        missing = [marker for marker in required_markers if not any(marker in path for path in paths)]
        if missing:
            failures.append("manifest 缺少企业级交付文件: " + ", ".join(missing))

    def _check_delivery_structures(self, read_file, records: list[dict], failures: list[str], warnings: list[str]) -> None:
        report_name = self._find_record_path(records, "互联网资产暴露面测绘报告")
        asset_name = self._find_record_path(records, "资产汇总")
        web_name = self._find_record_path(records, "Web资产详情")
        plan_name = self._find_record_path(records, "补全计划.json")
        quality_name = self._find_record_path(records, "quality_summary.txt")
        screenshot_manifest_name = self._find_record_path(records, "截图证据清单.json")
        service_audit_name = self._find_record_path(records, "服务分类审计.json")
        report_ai_audit_name = self._find_record_path(records, "报告AI分析审计.json")

        if report_name:
            self._check_docx_bytes(read_file(report_name), failures)
        if asset_name:
            self._check_asset_workbook_bytes(read_file(asset_name), failures)
        if web_name:
            self._check_web_workbook_bytes(read_file(web_name), failures, warnings)
        if plan_name:
            self._check_improvement_plan_bytes(read_file(plan_name), failures)
        if quality_name:
            data = read_file(quality_name)
            if data and "Quality:" not in data.decode("utf-8", errors="ignore"):
                failures.append("quality_summary.txt 缺少 Quality 状态")
        if screenshot_manifest_name:
            self._check_screenshot_manifest_bytes(read_file(screenshot_manifest_name), records, failures)
        elif any(str(record.get("path") or "").startswith("screenshots/") for record in records):
            warnings.append("交付包包含截图文件但缺少截图证据清单")
        if service_audit_name:
            self._check_service_audit_bytes(read_file(service_audit_name), failures)
        if report_ai_audit_name:
            self._check_report_ai_audit_bytes(read_file(report_ai_audit_name), failures)

    def _find_record_path(self, records: list[dict], marker: str) -> str | None:
        for record in records:
            path = str(record.get("path") or "")
            if marker in path:
                return path
        return None

    def _check_docx_bytes(self, data: bytes | None, failures: list[str]) -> None:
        if not data:
            return
        try:
            doc = Document(BytesIO(data))
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        except Exception as exc:
            failures.append(f"Word报告无法打开: {str(exc)[:200]}")
            return
        if "互联网数字资产暴露面测绘报告" not in text:
            failures.append("Word报告缺少报告标题")
        if len(doc.tables) < 6:
            failures.append(f"Word报告表格数量不足: {len(doc.tables)}")

    def _check_asset_workbook_bytes(self, data: bytes | None, failures: list[str]) -> None:
        if not data:
            return
        try:
            workbook = load_workbook(BytesIO(data), data_only=True)
        except Exception as exc:
            failures.append(f"资产汇总附件无法打开: {str(exc)[:200]}")
            return
        failure_count = len(failures)
        self._require_sheets(workbook, EXPECTED_ASSET_SHEETS, "资产汇总附件", failures)
        if len(failures) > failure_count:
            return
        if len(getattr(workbook["管理驾驶舱"], "_charts", [])) < 2:
            failures.append("资产汇总附件管理驾驶舱图表数量不足")
        self._require_headers(workbook["风险清单"], {"风险分值", "责任建议", "验收证据"}, "资产汇总附件/风险清单", failures)
        self._require_headers(workbook["重点资产视图"], {"风险分值", "责任建议"}, "资产汇总附件/重点资产视图", failures)
        self._require_headers(workbook["整改矩阵"], {"风险分值", "责任建议", "验收证据"}, "资产汇总附件/整改矩阵", failures)

    def _check_web_workbook_bytes(self, data: bytes | None, failures: list[str], warnings: list[str]) -> None:
        if not data:
            return
        try:
            workbook = load_workbook(BytesIO(data), data_only=True)
        except Exception as exc:
            failures.append(f"Web资产详情附件无法打开: {str(exc)[:200]}")
            return
        failure_count = len(failures)
        self._require_sheets(workbook, EXPECTED_WEB_SHEETS, "Web资产详情附件", failures)
        if len(failures) > failure_count:
            return
        sheet = workbook["截图证据"]
        self._require_headers(sheet, {"缩略图", "截图文件", "截图状态"}, "Web资产详情附件/截图证据", failures)
        headers = [cell.value for cell in sheet[1]]
        if "截图状态" in headers:
            status_column = headers.index("截图状态") + 1
            rows_with_screenshot = sum(
                1
                for row_index in range(2, sheet.max_row + 1)
                if sheet.cell(row=row_index, column=status_column).value == "有截图"
            )
            if rows_with_screenshot and not getattr(sheet, "_images", []):
                warnings.append("Web资产详情附件截图证据存在截图路径但未嵌入缩略图")

    def _check_improvement_plan_bytes(self, data: bytes | None, failures: list[str]) -> None:
        if not data:
            return
        try:
            payload = json.loads(data.decode("utf-8"))
        except Exception as exc:
            failures.append(f"补全计划JSON无法解析: {str(exc)[:200]}")
            return
        if not isinstance(payload.get("actions"), list):
            failures.append("补全计划JSON缺少 actions 列表")
        if not isinstance(payload.get("quality"), dict):
            failures.append("补全计划JSON缺少 quality 状态")

    def _check_screenshot_manifest_bytes(self, data: bytes | None, records: list[dict], failures: list[str]) -> None:
        if not data:
            return
        try:
            payload = json.loads(data.decode("utf-8"))
        except Exception as exc:
            failures.append(f"截图证据清单JSON无法解析: {str(exc)[:200]}")
            return
        screenshots = payload.get("screenshots")
        if not isinstance(screenshots, list):
            failures.append("截图证据清单缺少 screenshots 列表")
            return
        manifest_by_path = {str(record.get("path") or ""): record for record in records}
        manifest_paths = set(manifest_by_path)
        missing_paths = sorted(
            {
                str(item.get("package_path") or "")
                for item in screenshots
                if isinstance(item, dict)
                and str(item.get("package_path") or "").startswith("screenshots/")
                and str(item.get("package_path") or "") not in manifest_paths
            }
        )
        if missing_paths:
            failures.append("截图证据清单引用了未打包截图: " + ", ".join(missing_paths))
        hash_mismatches = []
        size_mismatches = []
        for item in screenshots:
            if not isinstance(item, dict):
                continue
            package_path = str(item.get("package_path") or "")
            record = manifest_by_path.get(package_path)
            if not record:
                continue
            if item.get("package_sha256") and item.get("package_sha256") != record.get("sha256"):
                hash_mismatches.append(package_path)
            if item.get("package_size") not in {None, ""} and item.get("package_size") != record.get("size"):
                size_mismatches.append(package_path)
        if hash_mismatches:
            failures.append("截图证据清单哈希与 manifest 不一致: " + ", ".join(sorted(set(hash_mismatches))))
        if size_mismatches:
            failures.append("截图证据清单大小与 manifest 不一致: " + ", ".join(sorted(set(size_mismatches))))
        incomplete_count = sum(
            1
            for item in screenshots
            if (
                not isinstance(item, dict)
                or not item.get("url")
                or not item.get("package_path")
                or not item.get("source_path")
                or not item.get("package_sha256")
                or item.get("package_size") in {None, ""}
            )
        )
        if incomplete_count:
            failures.append(f"截图证据清单存在不完整记录: {incomplete_count}")

    def _check_service_audit_bytes(self, data: bytes | None, failures: list[str]) -> None:
        if not data:
            return
        try:
            payload = json.loads(data.decode("utf-8"))
        except Exception as exc:
            failures.append(f"服务分类审计JSON无法解析: {str(exc)[:200]}")
            return
        if not isinstance(payload.get("kind_counts"), dict):
            failures.append("服务分类审计缺少 kind_counts")
        if not isinstance(payload.get("host_mode_counts"), dict):
            failures.append("服务分类审计缺少 host_mode_counts")
        services = payload.get("services")
        if not isinstance(services, list):
            failures.append("服务分类审计缺少 services 明细")
            return
        if payload.get("total") not in {None, len(services)}:
            failures.append(f"服务分类审计 total 与 services 数量不一致: total={payload.get('total')} services={len(services)}")
        incomplete_count = sum(
            1
            for item in services
            if not isinstance(item, dict)
            or not item.get("endpoint")
            or not item.get("asset_kind")
            or item.get("port") in {None, ""}
        )
        if incomplete_count:
            failures.append(f"服务分类审计存在不完整服务记录: {incomplete_count}")

    def _check_report_ai_audit_bytes(self, data: bytes | None, failures: list[str]) -> None:
        if not data:
            return
        try:
            payload = json.loads(data.decode("utf-8"))
        except Exception as exc:
            failures.append(f"报告AI分析审计JSON无法解析: {str(exc)[:200]}")
            return
        sections = payload.get("sections")
        if not isinstance(sections, list):
            failures.append("报告AI分析审计缺少 sections 列表")
            return
        if payload.get("section_count") != len(sections):
            failures.append(f"报告AI分析审计 section_count 不一致: section_count={payload.get('section_count')} sections={len(sections)}")
        expected = {"report_dns", "report_ports", "report_web", "report_summary"}
        actual = {str(item.get("analysis_type") or "") for item in sections if isinstance(item, dict)}
        missing = sorted(expected - actual)
        if missing:
            failures.append("报告AI分析审计缺少分块: " + ", ".join(missing))
        required_keys = {
            "analysis_type",
            "title",
            "status",
            "mode",
            "model",
            "summary_chars",
            "prompt_chars",
            "input_fingerprint",
            "data_shape",
            "response_keys",
            "response_id",
            "response_model",
            "usage_keys",
            "prompt_tokens",
            "completion_tokens",
            "total_tokens",
        }
        incomplete_count = sum(
            1
            for item in sections
            if not isinstance(item, dict)
            or not required_keys.issubset(item.keys())
            or not item.get("analysis_type")
            or not item.get("status")
            or item.get("prompt_chars") in {None, ""}
            or not isinstance(item.get("data_shape"), dict)
            or not isinstance(item.get("response_keys"), list)
            or not isinstance(item.get("usage_keys"), list)
        )
        if incomplete_count:
            failures.append(f"报告AI分析审计存在不完整分块记录: {incomplete_count}")

    def _require_sheets(self, workbook, required: set[str], label: str, failures: list[str]) -> None:
        missing = sorted(required - set(workbook.sheetnames))
        if missing:
            failures.append(f"{label}缺少Sheet: {', '.join(missing)}")

    def _require_headers(self, sheet, required: set[str], label: str, failures: list[str]) -> None:
        headers = {cell.value for cell in sheet[1] if cell.value}
        missing = sorted(required - headers)
        if missing:
            failures.append(f"{label}缺少字段: {', '.join(missing)}")

    def _check_record(self, record: dict, actual_size: int, actual_sha256: str, failures: list[str]) -> None:
        path = str(record.get("path") or "")
        expected_size = record.get("size")
        expected_sha256 = record.get("sha256")
        if expected_size != actual_size:
            failures.append(f"文件大小不匹配: {path} expected={expected_size} actual={actual_size}")
        if expected_sha256 != actual_sha256:
            failures.append(f"SHA256不匹配: {path}")

    def _result(
        self,
        package_path: Path,
        manifest: dict | None,
        failures: list[str],
        warnings: list[str],
        *,
        file_records: list[dict] | None = None,
    ) -> PackageVerificationResult:
        status = "FAIL" if failures else ("WARN" if warnings else "PASS")
        records = file_records or []
        lines = [
            f"Package verification: {status}",
            f"Package: {package_path}",
        ]
        if manifest:
            lines.extend(
                [
                    f"Task: {manifest.get('task_id')}",
                    f"Target: {manifest.get('target')}",
                    f"Quality: {manifest.get('quality_status')}",
                    f"Manifest files: {len(records)}",
                ]
            )
        if failures:
            lines.extend(["", "Failures:"])
            lines.extend(f"- {item}" for item in failures)
        if warnings:
            lines.extend(["", "Warnings:"])
            lines.extend(f"- {item}" for item in warnings)
        return PackageVerificationResult(status=status, failures=failures, warnings=warnings, lines=lines)

    def _sha256(self, path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()
