from __future__ import annotations

import json
import ipaddress
import hashlib
import re
from collections import Counter
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Callable
from urllib.parse import urlparse, urlunparse

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.chart import BarChart, PieChart, Reference
from openpyxl.drawing.image import Image as XLImage
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from sqlmodel import Session, select

from assetmap.config import AppConfig
from assetmap.models import AiAnalysis, ScanTask
from assetmap.services.ai_client import chat_completion
from assetmap.services.exporter import ExportService


PARKING_CNAME_KEYWORDS = ("expired.", "parking", "parked", "hichina.com")


def _utcnow() -> datetime:
    return datetime.now(timezone.utc)


def _safe_name(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9_.\-\u4e00-\u9fff]+", "_", value).strip("._") or "report"


def _timestamped_path(path: Path) -> Path:
    suffix = _utcnow().strftime("%Y%m%d_%H%M%S_%f")
    return path.with_name(f"{path.stem}_{suffix}{path.suffix}")


def _default_port(scheme: str) -> int:
    return 443 if scheme == "https" else 80


def _normalize_url_for_match(value: str | None) -> str:
    if not value:
        return ""
    parsed = urlparse(value.strip())
    if parsed.scheme not in {"http", "https"} or not parsed.hostname:
        return ""
    netloc = parsed.hostname.lower().rstrip(".")
    if parsed.port and parsed.port != _default_port(parsed.scheme):
        netloc = f"{netloc}:{parsed.port}"
    return urlunparse((parsed.scheme.lower(), netloc, (parsed.path or "/").rstrip("/") or "/", "", parsed.query, ""))


def _is_real_public_ip(value: str) -> bool:
    try:
        ip = ipaddress.ip_address(value)
    except ValueError:
        return False
    return ip.is_global and ip not in ipaddress.ip_network("198.18.0.0/15")


def _is_parking_cname(value: str) -> bool:
    text = value.lower().rstrip(".")
    return any(keyword in text for keyword in PARKING_CNAME_KEYWORDS)


def _text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, (list, tuple, set)):
        return ", ".join(_text(item) for item in value if _text(item))
    if isinstance(value, dict):
        return json.dumps(value, ensure_ascii=False)
    return str(value)


def _short(value: Any, limit: int = 4000) -> str:
    text = _text(value)
    return text if len(text) <= limit else text[:limit] + "..."


def _fofa_payload(port: dict) -> dict:
    payload = port.get("raw_payload") or {}
    if isinstance(payload.get("fofa"), dict):
        return payload["fofa"]
    return payload if payload.get("source") == "fofa" else {}


HEADER_FILL = "1F4E78"
SUBTLE_FILL = "EAF2F8"
HIGH_FILL = "FCE4D6"
MEDIUM_FILL = "FFF2CC"
LOW_FILL = "E2F0D9"
BORDER_COLOR = "D9E2F3"
REPORT_ANALYSIS_TITLES = {
    "report_dns": "DNS与域名解析分析",
    "report_ports": "端口与服务暴露分析",
    "report_web": "Web资产视觉识别分析",
    "report_summary": "总体暴露面结论与处置建议",
}
REPORT_ANALYSIS_CACHE_VERSION = "report-analysis-v2"
MAX_EXCEL_SCREENSHOT_THUMBNAILS = 20
LOW_VALUE_WEB_TITLE_MARKERS = (
    "400",
    "403",
    "404",
    "500",
    "not found",
    "forbidden",
    "bad request",
    "internal server error",
    "non-compliance icp filing",
    "welcome to nginx",
    "nginx",
    "apache tomcat",
)


@dataclass
class ReportResult:
    report_path: Path
    asset_workbook_path: Path
    web_workbook_path: Path
    analysis_count: int


class ReportService:
    def __init__(
        self,
        session: Session,
        config: AppConfig,
        progress: Callable[[str], None] | None = None,
    ) -> None:
        self.session = session
        self.config = config
        self.progress = progress

    def _log(self, message: str) -> None:
        if self.progress:
            try:
                self.progress(message)
            except OSError:
                self.progress = None

    def run(self, task_id: int, output_dir: Path | str = "reports", rerun_ai: bool = False) -> ReportResult:
        task = self.session.get(ScanTask, task_id)
        if not task:
            raise ValueError(f"Task not found: {task_id}")
        bundle = ExportService(self.session)._bundle(task_id)
        context = self._context(bundle)
        root = Path(output_dir) / f"task_{task_id}_{_safe_name(task.target)}"
        root.mkdir(parents=True, exist_ok=True)

        self._log("[report] running chunked AI analysis")
        analyses = self._run_analyses(task_id, bundle, context, rerun_ai=rerun_ai)
        audit_path = self._write_report_ai_audit(task_id)
        self._log(f"[report] AI analysis audit: {audit_path}")
        context["ai_audit_rows"] = self._report_ai_audit_rows(task_id)

        self._log("[report] preparing Excel attachments")
        asset_workbook = root / f"task_{task_id}_资产汇总.xlsx"
        web_workbook = root / f"task_{task_id}_Web资产详情.xlsx"
        asset_workbook = self._write_asset_workbook(asset_workbook, context)
        web_workbook = self._write_web_workbook(web_workbook, context)

        self._log("[report] writing Word report")
        report_path = root / f"task_{task_id}_互联网资产暴露面测绘报告.docx"
        report_path = self._write_docx(report_path, task, context, analyses, asset_workbook, web_workbook)
        return ReportResult(report_path, asset_workbook, web_workbook, len(analyses))

    def _context(self, bundle: dict) -> dict:
        companies = {row["id"]: row for row in bundle["companies"]}
        company_names = {company_id: row["name"] for company_id, row in companies.items()}
        domain_units: dict[str, str] = {}
        ip_units: dict[str, str] = {}
        named_assets: list[dict] = []
        for asset in bundle["assets"]:
            unit = company_names.get(asset["company_id"], "")
            asset_type = asset["asset_type"]
            name = asset.get("display_name") or asset.get("normalized_identifier")
            if asset_type == "icp_domain":
                domain_units[asset["normalized_identifier"].lower().rstrip(".")] = unit
            elif asset_type == "ip":
                ip_units[asset["normalized_identifier"]] = unit
            else:
                named_assets.append(
                    {
                        "单位": unit,
                        "资产类型": asset_type,
                        "资产名称": name,
                        "标识": asset.get("normalized_identifier"),
                        "来源": asset.get("source_tool"),
                    }
                )

        dns_records = bundle["dns_records"]
        company_scope = self._company_scope(bundle, company_names)
        dns_unit_by_ip: dict[str, str] = {}
        for record in dns_records:
            if record["record_type"] not in {"A", "AAAA"}:
                continue
            unit = self._unit_for_host(record["fqdn"], domain_units)
            if unit:
                dns_unit_by_ip.setdefault(record["value"], unit)

        review_attestations = self._review_attestations(bundle)
        port_rows = self._port_rows(bundle, domain_units, ip_units, dns_unit_by_ip)
        dns_rows = self._dns_rows(dns_records, domain_units)
        dns_quality_rows = self._dns_quality_rows(bundle, domain_units, review_attestations)
        unit_asset_rows = self._unit_asset_rows(bundle, company_names)
        web_rows = self._web_rows(bundle, domain_units, ip_units, dns_unit_by_ip, review_attestations)
        service_audit_rows = self._service_audit_rows(bundle, domain_units, ip_units, dns_unit_by_ip, review_attestations)
        url_coverage_rows = self._url_coverage_rows(bundle, service_audit_rows, review_attestations)
        visual_review_rows = self._visual_review_rows(web_rows, review_attestations)
        risk_rows = self._risk_rows(port_rows, web_rows)
        port_target_rows = self._port_target_rows(bundle, domain_units, ip_units, dns_unit_by_ip, port_rows, web_rows, risk_rows)
        manual_no_asset_reviews = self._manual_no_asset_reviews(bundle)
        unit_coverage_rows = self._unit_coverage_rows(
            company_names,
            company_scope,
            unit_asset_rows,
            dns_rows,
            port_rows,
            web_rows,
            risk_rows,
            manual_no_asset_reviews,
        )
        coverage_rows = self._coverage_rows(
            bundle,
            dns_rows,
            dns_quality_rows,
            port_rows,
            web_rows,
            service_audit_rows,
            url_coverage_rows,
            unit_coverage_rows,
            port_target_rows,
        )
        remediation_rows = self._remediation_rows(risk_rows)
        key_asset_rows = self._key_asset_rows(risk_rows, port_target_rows, web_rows)
        key_web_rows = self._top_web_rows(web_rows, risk_rows)[:50]
        screenshot_evidence_rows = self._screenshot_evidence_rows(web_rows, risk_rows)
        stats = self._stats(
            bundle,
            port_rows,
            port_target_rows,
            web_rows,
            dns_quality_rows,
            service_audit_rows,
            url_coverage_rows,
            visual_review_rows,
            risk_rows,
            coverage_rows,
            unit_coverage_rows,
        )
        overview_rows = self._overview_rows(stats)
        risk_summary_rows = self._risk_summary_rows(risk_rows)
        dashboard_rows = self._dashboard_rows(stats, risk_summary_rows, coverage_rows)
        audit_file_rows = self._audit_file_rows(bundle["task"]["id"])
        return {
            "company_names": company_names,
            "company_scope": company_scope,
            "domain_units": domain_units,
            "ip_units": ip_units,
            "named_assets": named_assets,
            "dns_rows": dns_rows,
            "dns_quality_rows": dns_quality_rows,
            "port_rows": port_rows,
            "port_target_rows": port_target_rows,
            "unit_asset_rows": unit_asset_rows,
            "key_asset_rows": key_asset_rows,
            "web_rows": web_rows,
            "key_web_rows": key_web_rows,
            "screenshot_evidence_rows": screenshot_evidence_rows,
            "service_audit_rows": service_audit_rows,
            "url_coverage_rows": url_coverage_rows,
            "visual_review_rows": visual_review_rows,
            "risk_rows": risk_rows,
            "coverage_rows": coverage_rows,
            "unit_coverage_rows": unit_coverage_rows,
            "remediation_rows": remediation_rows,
            "overview_rows": overview_rows,
            "risk_summary_rows": risk_summary_rows,
            "dashboard_rows": dashboard_rows,
            "audit_file_rows": audit_file_rows,
            "stats": stats,
        }

    def _unit_for_host(self, host: str | None, domain_units: dict[str, str]) -> str:
        if not host:
            return ""
        value = host.lower().rstrip(".")
        for domain in sorted(domain_units, key=len, reverse=True):
            if value == domain or value.endswith(f".{domain}"):
                return domain_units[domain]
        return ""

    def _company_scope(self, bundle: dict, company_names: dict[int, str]) -> dict[str, dict[str, Any]]:
        child_ids = {edge["child_company_id"] for edge in bundle["edges"]}
        child_counts: dict[int, int] = {}
        for edge in bundle["edges"]:
            parent_id = edge.get("parent_company_id")
            if parent_id:
                child_counts[parent_id] = child_counts.get(parent_id, 0) + 1
        root_ids = sorted(set(company_names) - child_ids)
        scope = {
            company_names[company_id]: {
                "股权层级": 0,
                "直接持股": "100.00%",
                "累计持股": "100.00%",
                "股权路径": company_names[company_id],
                "子公司数量": child_counts.get(company_id, 0),
            }
            for company_id in root_ids
            if company_id in company_names
        }
        for edge in sorted(
            bundle["edges"],
            key=lambda item: (item.get("depth", 0), -(item.get("cumulative_holding_ratio") or 0), company_names.get(item["child_company_id"], "")),
        ):
            child = company_names.get(edge["child_company_id"])
            if not child:
                continue
            candidate = {
                "股权层级": edge.get("depth", ""),
                "直接持股": f"{(edge.get('direct_holding_ratio') or 0):.2%}",
                "累计持股": f"{(edge.get('cumulative_holding_ratio') or 0):.2%}",
                "股权路径": edge.get("path") or child,
                "子公司数量": child_counts.get(edge["child_company_id"], 0),
            }
            existing = scope.get(child)
            if not existing or self._ratio_value(candidate["累计持股"]) > self._ratio_value(existing.get("累计持股")):
                scope[child] = candidate
        for name in company_names.values():
            scope.setdefault(
                name,
                {
                    "股权层级": "",
                    "直接持股": "",
                    "累计持股": "",
                    "股权路径": name,
                    "子公司数量": 0,
                },
            )
        return scope

    def _ratio_value(self, value: Any) -> float:
        text = _text(value).strip().rstrip("%")
        try:
            return float(text)
        except ValueError:
            return 0.0

    def _port_rows(
        self,
        bundle: dict,
        domain_units: dict[str, str],
        ip_units: dict[str, str],
        dns_unit_by_ip: dict[str, str],
    ) -> list[dict]:
        service_by_key = {
            (row["target_ip"], row["protocol"], row["port"]): row
            for row in bundle["service_assets"]
        }
        dns_by_ip: dict[str, set[str]] = {}
        for record in bundle["dns_records"]:
            if record["record_type"] in {"A", "AAAA"}:
                dns_by_ip.setdefault(record["value"], set()).add(record["fqdn"])
        rows = []
        seen = set()
        for port in sorted(bundle["nmap_ports"], key=lambda item: (item["target_ip"], item["port"])):
            if port.get("state") != "open":
                continue
            service = service_by_key.get((port["target_ip"], port["protocol"], port["port"]), {})
            fofa_payload = _fofa_payload(port)
            domains = sorted(set(service.get("domains") or []) | dns_by_ip.get(port["target_ip"], set()))
            fofa_host = fofa_payload.get("host") or ""
            unit = (
                next((self._unit_for_host(domain, domain_units) for domain in domains if self._unit_for_host(domain, domain_units)), "")
                or ip_units.get(port["target_ip"])
                or dns_unit_by_ip.get(port["target_ip"], "")
            )
            key = (unit, port["target_ip"], port["protocol"], port["port"])
            if key in seen:
                continue
            seen.add(key)
            rows.append(
                {
                    "单位": unit,
                    "IP": port["target_ip"],
                    "域名": ", ".join(domains[:20]),
                    "协议": port["protocol"],
                    "端口": port["port"],
                    "端口状态": port["state"],
                    "服务": service.get("service") or port.get("service") or "",
                    "产品": service.get("product") or port.get("product") or "",
                    "版本": service.get("version") or port.get("version") or "",
                    "资产类型": service.get("asset_kind") or "unknown",
                    "Web URL": service.get("representative_url") or "",
                    "Web标题": service.get("title") or "",
                    "FOFA Host": fofa_host,
                    "FOFA标题": fofa_payload.get("title") or "",
                    "来源": self._port_source(port),
                    "证据类型": self._port_evidence_type(port),
                    "主动扫描确认": "是" if self._has_nmap_evidence(port) else "否",
                    "被动FOFA证据": "是" if self._has_fofa_evidence(port) else "否",
                    "备注": _short(service.get("evidence") or fofa_payload, 800),
                }
            )
        return rows

    def _port_source(self, port: dict) -> str:
        payload = port.get("raw_payload") or {}
        sources = payload.get("sources")
        if isinstance(sources, list) and sources:
            return ",".join(str(item) for item in sources)
        return payload.get("source") or "nmap"

    def _port_evidence_type(self, port: dict) -> str:
        has_nmap = self._has_nmap_evidence(port)
        has_fofa = self._has_fofa_evidence(port)
        if has_nmap and has_fofa:
            return "主动+被动"
        if has_nmap:
            return "主动Nmap"
        if has_fofa:
            return "被动FOFA"
        return "未知"

    def _has_fofa_evidence(self, port: dict) -> bool:
        payload = port.get("raw_payload") or {}
        return payload.get("source") == "fofa" or "fofa" in (payload.get("sources") or []) or isinstance(payload.get("fofa"), dict)

    def _has_nmap_evidence(self, port: dict) -> bool:
        payload = port.get("raw_payload") or {}
        sources = payload.get("sources")
        if isinstance(sources, list):
            return "nmap" in sources
        return payload.get("source") not in {"fofa"} and not isinstance(payload.get("fofa"), dict)

    def _port_target_rows(
        self,
        bundle: dict,
        domain_units: dict[str, str],
        ip_units: dict[str, str],
        dns_unit_by_ip: dict[str, str],
        port_rows: list[dict],
        web_rows: list[dict],
        risk_rows: list[dict],
    ) -> list[dict]:
        nmap_task = bundle.get("nmap_task") or {}
        task_targets = [str(item) for item in (nmap_task.get("targets") or []) if item]
        open_ports_by_ip: dict[str, list[dict]] = {}
        for row in port_rows:
            if row.get("IP"):
                open_ports_by_ip.setdefault(row["IP"], []).append(row)
        web_by_ip: dict[str, list[dict]] = {}
        for row in web_rows:
            if row.get("IP"):
                web_by_ip.setdefault(row["IP"], []).append(row)
        risk_by_ip: dict[str, list[dict]] = {}
        for row in risk_rows:
            asset = row.get("资产") or ""
            ip = asset.split(":", 1)[0] if ":" in asset else ""
            if ip:
                risk_by_ip.setdefault(ip, []).append(row)

        domains_by_ip: dict[str, set[str]] = {}
        for record in bundle.get("dns_records", []):
            if record.get("record_type") in {"A", "AAAA"} and record.get("value") and record.get("fqdn"):
                domains_by_ip.setdefault(record["value"], set()).add(record["fqdn"])

        source_map = self._port_target_source_map(bundle)
        targets = sorted(set(task_targets) | set(open_ports_by_ip) | set(source_map))
        rows = []
        for ip in targets:
            domains = sorted(domains_by_ip.get(ip, set()))
            unit = (
                ip_units.get(ip)
                or dns_unit_by_ip.get(ip)
                or next((self._unit_for_host(domain, domain_units) for domain in domains if self._unit_for_host(domain, domain_units)), "")
            )
            ports = sorted({int(row.get("端口") or 0) for row in open_ports_by_ip.get(ip, []) if row.get("端口")})
            active_ports = sorted({int(row.get("端口") or 0) for row in open_ports_by_ip.get(ip, []) if row.get("主动扫描确认") == "是" and row.get("端口")})
            passive_ports = sorted({int(row.get("端口") or 0) for row in open_ports_by_ip.get(ip, []) if row.get("被动FOFA证据") == "是" and row.get("端口")})
            merged_ports = sorted({int(row.get("端口") or 0) for row in open_ports_by_ip.get(ip, []) if row.get("证据类型") == "主动+被动" and row.get("端口")})
            sources = sorted(source_map.get(ip, set()))
            rows.append(
                {
                    "IP": ip,
                    "归属单位": unit,
                    "目标来源": ", ".join(sources) or "open_port_evidence",
                    "关联域名数": len(domains),
                    "关联域名样例": ", ".join(domains[:20]),
                    "是否进入扫描目标": "是" if ip in task_targets else "否",
                    "开放端口数量": len(ports),
                    "开放端口": ", ".join(str(port) for port in ports),
                    "主动端口数": len(active_ports),
                    "被动端口数": len(passive_ports),
                    "合并端口数": len(merged_ports),
                    "端口证据类型": self._port_target_evidence_type(active_ports, passive_ports, merged_ports),
                    "Web入口数量": len(web_by_ip.get(ip, [])),
                    "高风险数量": sum(1 for row in risk_by_ip.get(ip, []) if row.get("风险等级") == "高"),
                    "中风险数量": sum(1 for row in risk_by_ip.get(ip, []) if row.get("风险等级") == "中"),
                    "复核建议": self._port_target_action(ip in task_targets, ports, sources),
                }
            )
        return sorted(
            rows,
            key=lambda item: (
                0 if item["高风险数量"] else 1,
                0 if item["开放端口数量"] else 1,
                item["归属单位"],
                item["IP"],
            ),
        )

    def _port_target_evidence_type(self, active_ports: list[int], passive_ports: list[int], merged_ports: list[int]) -> str:
        if merged_ports:
            return "主动+被动"
        if active_ports and passive_ports:
            return "主动/被动分散"
        if active_ports:
            return "主动Nmap"
        if passive_ports:
            return "被动FOFA"
        return "无开放端口证据"

    def _port_target_source_map(self, bundle: dict) -> dict[str, set[str]]:
        sources: dict[str, set[str]] = {}
        parked_hosts = {
            record.get("fqdn")
            for record in bundle.get("dns_records", [])
            if record.get("record_type") == "CNAME" and _is_parking_cname(str(record.get("value") or ""))
        }
        for record in bundle.get("dns_records", []):
            if record.get("record_type") not in {"A", "AAAA"} or not record.get("value"):
                continue
            kind = (record.get("raw_payload") or {}).get("kind")
            source = "manual" if kind == "manual_ip" else "dns_public"
            if source == "dns_public" and (record.get("fqdn") in parked_hosts or not _is_real_public_ip(str(record["value"]))):
                continue
            sources.setdefault(record["value"], set()).add(source)
        for asset in bundle.get("assets", []):
            if asset.get("asset_type") == "ip" and asset.get("normalized_identifier"):
                sources.setdefault(asset["normalized_identifier"], set()).add("manual")
        for analysis in bundle.get("ai_analyses", []):
            if analysis.get("analysis_type") != "dns_inference" or analysis.get("status") != "completed":
                continue
            for ip in self._extract_ai_target_ips(analysis.get("summary") or ""):
                sources.setdefault(ip, set()).add("ai")
        nmap_task = bundle.get("nmap_task") or {}
        for ip in nmap_task.get("targets") or []:
            sources.setdefault(str(ip), set())
        return sources

    def _extract_ai_target_ips(self, text: str) -> list[str]:
        block = re.search(r"NMAP_TARGET_IPS\s*[:：]?\s*(.*?)(?:\n\s*\n|$)", text, flags=re.IGNORECASE | re.DOTALL)
        source = block.group(1) if block else text
        values = re.findall(r"\b(?:\d{1,3}\.){3}\d{1,3}\b", source)
        output = []
        for value in values:
            if value not in output:
                output.append(value)
        return output

    def _port_target_action(self, scanned: bool, ports: list[int], sources: list[str]) -> str:
        if ports:
            return "已发现开放端口，继续纳入服务识别、URL识别和风险复核。"
        if not scanned:
            return "存在目标来源但未进入扫描目标，建议确认过滤规则后补扫。"
        if "manual" in sources:
            return "人工补充 IP 暂无开放端口证据，建议保留台账并按周期复测。"
        return "暂无开放端口证据，建议结合 FOFA、历史扫描和业务访谈复核。"

    def _dns_rows(self, dns_records: list[dict], domain_units: dict[str, str]) -> list[dict]:
        rows = []
        seen = set()
        for record in sorted(dns_records, key=lambda item: (item["fqdn"], item["record_type"], item["value"])):
            key = (record["fqdn"], record["record_type"], record["value"])
            if key in seen:
                continue
            seen.add(key)
            rows.append(
                {
                    "单位": self._unit_for_host(record["fqdn"], domain_units),
                    "根域名": record["root_domain"],
                    "主机名": record["fqdn"],
                    "记录类型": record["record_type"],
                    "记录值": record["value"],
                    "TTL": record.get("ttl") or "",
                    "来源": "dns",
                }
            )
        return rows

    def _review_attestations(self, bundle: dict) -> dict[str, dict[str, dict]]:
        attestations: dict[str, dict[str, dict]] = {}
        for record in bundle.get("raw_records", []):
            if record.get("source") != "review_workorder" or record.get("action") != "review_attestation":
                continue
            payload = record.get("response_json") or {}
            category = _text(payload.get("category"))
            item_key = _text(payload.get("item_key"))
            status = _text(payload.get("review_status")).strip()
            if not category or not item_key or status.lower() in {"", "pending", "todo"} or status in {"待确认", "待复核", "未复核"}:
                continue
            attestations.setdefault(category, {})[item_key] = payload
        return attestations

    def _apply_review_attestation(
        self,
        row: dict,
        attestations: dict[str, dict[str, dict]],
        category: str,
        item_key: str,
    ) -> None:
        attestation = attestations.get(category, {}).get(item_key)
        if not attestation:
            return
        row["人工复核状态"] = attestation.get("review_status") or "reviewed"
        row["人工复核备注"] = attestation.get("review_notes") or ""
        row["复核优先级"] = "无"
        if "复核原因" in row:
            row["复核原因"] = "已人工复核：" + _short(attestation.get("review_status"), 120)
        row["建议动作"] = "已导入人工复核结论，后续纳入周期性复测。"

    def _dns_quality_rows(self, bundle: dict, domain_units: dict[str, str], review_attestations: dict[str, dict[str, dict]] | None = None) -> list[dict]:
        roots = sorted(
            {
                (asset.get("normalized_identifier") or "").lower().rstrip(".")
                for asset in bundle["assets"]
                if asset.get("asset_type") == "icp_domain" and asset.get("normalized_identifier")
            }
        )
        subdomains_by_root: dict[str, set[str]] = {root: set() for root in roots}
        records_by_root: dict[str, list[dict]] = {root: [] for root in roots}
        tool_runs_by_root: dict[str, list[dict]] = {root: [] for root in roots}
        for row in bundle.get("subdomains", []):
            root = row.get("root_domain") or self._matching_root(row.get("fqdn"), roots)
            if root in subdomains_by_root and row.get("fqdn"):
                subdomains_by_root[root].add(row["fqdn"])
        for row in bundle.get("dns_records", []):
            root = row.get("root_domain") or self._matching_root(row.get("fqdn"), roots)
            if root in records_by_root:
                records_by_root[root].append(row)
        for row in bundle.get("subdomain_tool_runs", []):
            root = row.get("root_domain")
            if root in tool_runs_by_root:
                tool_runs_by_root[root].append(row)

        ip_domains: dict[str, set[str]] = {}
        for row in bundle.get("dns_records", []):
            if row.get("record_type") in {"A", "AAAA"}:
                ip_domains.setdefault(row.get("value") or "", set()).add(row.get("fqdn") or "")

        rows = []
        for root in roots:
            records = records_by_root[root]
            ips = sorted({row.get("value") for row in records if row.get("record_type") in {"A", "AAAA"} and row.get("value")})
            cnames = sorted({row.get("value") for row in records if row.get("record_type") == "CNAME" and row.get("value")})
            failed_tools = [row for row in tool_runs_by_root[root] if row.get("status") == "failed"]
            completed_tools = [row for row in tool_runs_by_root[root] if row.get("status") == "completed"]
            shared_ip_clues = [
                f"{ip}({len(ip_domains.get(ip, set()))}个域名)"
                for ip in ips
                if len(ip_domains.get(ip, set())) >= 10
            ][:5]
            third_party_cnames = [
                cname
                for cname in cnames
                if self._is_third_party_cname(cname, root)
            ][:8]
            reasons = []
            review_types = []
            if bundle.get("subdomain_task") or tool_runs_by_root[root]:
                if not subdomains_by_root[root]:
                    reasons.append("未发现子域名")
                    review_types.append("no_subdomain")
                if failed_tools:
                    reasons.append("子域名工具失败")
                    review_types.append("tool_failure")
            if not ips:
                reasons.append("未获得公网A/AAAA记录")
                review_types.append("no_public_ip")
            if shared_ip_clues:
                reasons.append("多域名共用IP，需确认是否为真实业务或泛解析/CDN")
                review_types.append("shared_ip")
            if third_party_cnames:
                reasons.append("存在第三方/CDN/停放CNAME")
                review_types.append("third_party_cname")
            priority = "中" if (not ips or (failed_tools and not subdomains_by_root[root])) else ("低" if reasons else "无")
            row = {
                    "复核优先级": priority,
                    "复核类型": "+".join(review_types) if review_types else "none",
                    "根域名": root,
                    "归属单位": domain_units.get(root, ""),
                    "子域名数量": len(subdomains_by_root[root]),
                    "DNS记录数量": len(records),
                    "公网IP数量": len(ips),
                    "CNAME数量": len(cnames),
                    "NS数量": sum(1 for row in records if row.get("record_type") == "NS"),
                    "工具完成": ", ".join(f"{row.get('tool_name')}={row.get('status')}" for row in completed_tools) or "",
                    "工具失败": ", ".join(f"{row.get('tool_name')}={_short(row.get('error_message'), 120)}" for row in failed_tools) or "",
                    "共享IP线索": ", ".join(shared_ip_clues),
                    "第三方CNAME线索": ", ".join(third_party_cnames),
                    "复核原因": "；".join(reasons) if reasons else "无",
                    "建议动作": self._dns_review_action(priority, failed_tools, bool(subdomains_by_root[root]), bool(ips)),
                }
            self._apply_review_attestation(row, review_attestations or {}, "dns", root)
            rows.append(row)
        priority_order = {"中": 0, "低": 1, "无": 2}
        return sorted(rows, key=lambda item: (priority_order.get(item["复核优先级"], 9), item["归属单位"], item["根域名"]))

    def _matching_root(self, fqdn: str | None, roots: list[str]) -> str:
        value = (fqdn or "").lower().rstrip(".")
        return next((root for root in sorted(roots, key=len, reverse=True) if value == root or value.endswith(f".{root}")), "")

    def _is_third_party_cname(self, cname: str, root: str) -> bool:
        value = cname.lower().rstrip(".")
        if value == root or value.endswith(f".{root}"):
            return False
        keywords = (
            "expired.",
            "parking",
            "hichina.com",
            "my3w.com",
            "kunluncan.com",
            "kunlunaq.com",
            "cdn",
            "waf",
            "cloudfront",
            "aliyuncs.com",
            "myqcloud.com",
        )
        return any(keyword in value for keyword in keywords)

    def _dns_review_action(self, priority: str, failed_tools: list[dict], has_subdomains: bool, has_ips: bool) -> str:
        if priority == "无":
            return "保持周期性复测，关注新增子域名、解析变更和证书透明度线索。"
        actions = []
        if failed_tools:
            actions.append("优先重跑失败的子域名工具或检查工具输出/超时配置")
        if not has_subdomains:
            actions.append("补充证书透明度、搜索引擎、历史DNS和人工线索")
        if not has_ips:
            actions.append("复核域名状态、权威DNS和是否仅作为跳转/停放域名")
        actions.append("对共享IP和第三方CNAME确认是否为真实业务入口")
        return "；".join(actions) + "。"

    def _unit_asset_rows(self, bundle: dict, company_names: dict[int, str]) -> list[dict]:
        rows = []
        seen = set()
        for asset in sorted(bundle["assets"], key=lambda item: (item["company_id"], item["asset_type"], item["normalized_identifier"])):
            key = (asset["company_id"], asset["asset_type"], asset["normalized_identifier"])
            if key in seen:
                continue
            seen.add(key)
            rows.append(
                {
                    "单位": company_names.get(asset["company_id"], ""),
                    "资产类型": asset["asset_type"],
                    "资产名称": asset.get("display_name") or "",
                    "资产标识": asset.get("normalized_identifier") or "",
                    "来源": asset.get("source_tool") or "",
                    "原始信息": _short(
                        {
                            "asset": asset.get("raw_payload"),
                            "link_sources": asset.get("source_payloads"),
                        },
                        1000,
                    ),
                }
            )
        return rows

    def _web_rows(
        self,
        bundle: dict,
        domain_units: dict[str, str],
        ip_units: dict[str, str],
        dns_unit_by_ip: dict[str, str],
        review_attestations: dict[str, dict[str, dict]] | None = None,
    ) -> list[dict]:
        service_by_id = {row["id"]: row for row in bundle["service_assets"]}
        rows = []
        seen = set()
        for entry in sorted(bundle["web_entrypoints"], key=lambda item: item["normalized_url"]):
            evidence = entry.get("evidence") or {}
            visual = evidence.get("visual_analysis") or {}
            service = service_by_id.get(entry.get("service_asset_id"), {})
            system_name = visual.get("system_name") or visual.get("website_title") or ""
            confidence = visual.get("confidence") if visual.get("confidence") is not None else ""
            if visual and not system_name:
                fallback_title = self._title_based_system_name(entry, service)
                if fallback_title:
                    system_name = fallback_title
                    confidence_value = self._float_or_none(confidence)
                    confidence = max(confidence_value or 0, 0.45 if self._is_low_value_title(fallback_title) else 0.55)
            fallback_category = self._visual_fallback_category(visual, service)
            fallback_reason = visual.get("screenshot_error") or visual.get("ai_error") or ""
            unit = (
                self._unit_for_host(entry.get("host"), domain_units)
                or ip_units.get(entry.get("target_ip") or "")
                or dns_unit_by_ip.get(entry.get("target_ip") or "", "")
                or (evidence.get("manual_import") or {}).get("unit", "")
            )
            key = entry["normalized_url"]
            if key in seen:
                continue
            seen.add(key)
            row = {
                    "单位": unit,
                    "URL": entry["normalized_url"],
                    "最终URL": entry.get("final_url") or "",
                    "IP": entry.get("target_ip") or "",
                    "端口": entry.get("port") or "",
                    "HTTP状态": entry.get("http_status") or "",
                    "HTML标题": entry.get("title") or "",
                    "AI识别系统": system_name,
                    "网站用途": visual.get("site_purpose") or "",
                    "页面类型": visual.get("page_type") or "",
                    "登录特征": _short(visual.get("login_features"), 600),
                    "业务功能": _short(visual.get("business_functions"), 800),
                    "技术线索": _short(visual.get("visible_technical_clues") or entry.get("tech_stack"), 800),
                    "服务端": entry.get("server") or "",
                    "内容类型": entry.get("content_type") or "",
                    "主机模式": service.get("host_mode") or "",
                    "识别方式": visual.get("analysis_method") or ("screenshot_ai" if visual else ""),
                    "识别置信度": confidence,
                    "降级类型": fallback_category,
                    "降级原因": _short(fallback_reason, 800),
                    "截图": visual.get("screenshot_path") or evidence.get("visual_analysis_screenshot_path") or "",
                    "分析错误": "" if visual else evidence.get("visual_analysis_error") or "",
                }
            self._apply_visual_attestation(row, review_attestations or {}, key)
            rows.append(row)
        return rows

    def _apply_visual_attestation(
        self,
        row: dict,
        review_attestations: dict[str, dict[str, dict]],
        url: str,
    ) -> None:
        attestation = review_attestations.get("visual_identification", {}).get(url)
        if not attestation:
            return
        raw = attestation.get("raw") if isinstance(attestation.get("raw"), dict) else {}
        mapping = {
            "单位": attestation.get("confirmed_owner_unit") or raw.get("confirmed_owner_unit") or raw.get("确认归属单位"),
            "AI识别系统": attestation.get("confirmed_system_name") or raw.get("confirmed_system_name") or raw.get("确认系统名称"),
            "网站用途": attestation.get("confirmed_site_purpose") or raw.get("confirmed_site_purpose") or raw.get("确认网站用途"),
            "页面类型": attestation.get("confirmed_page_type") or raw.get("confirmed_page_type") or raw.get("确认页面类型"),
            "登录特征": attestation.get("confirmed_login_features") or raw.get("confirmed_login_features") or raw.get("确认登录特征"),
            "业务功能": attestation.get("confirmed_business_functions") or raw.get("confirmed_business_functions") or raw.get("确认业务功能"),
        }
        for key, value in mapping.items():
            text = _text(value).strip()
            if text:
                row[key] = text
        row["人工复核状态"] = attestation.get("review_status") or "reviewed"
        row["人工复核备注"] = attestation.get("review_notes") or ""
        row["人工复核人"] = attestation.get("reviewer") or ""
        row["人工复核时间"] = attestation.get("reviewed_at") or ""
        row["人工复核来源"] = ", ".join(_text(item) for item in (attestation.get("source_urls") or []) if _text(item))
        row["识别方式"] = "manual_review"
        row["识别置信度"] = 1.0

    def _title_based_system_name(self, entry: dict, service: dict | None = None) -> str:
        title = _text(entry.get("title")).strip()
        if title:
            return title[:120]
        service = service or {}
        for key in ("title", "app_name", "product", "service", "representative_url"):
            value = _text(service.get(key)).strip()
            if value:
                return value[:120]
        return _text(entry.get("host") or entry.get("normalized_url")).strip()[:120]

    def _visual_fallback_category(self, visual: dict, service: dict | None = None) -> str:
        if not visual or visual.get("analysis_method") != "http_probe_fallback":
            return ""
        service = service or {}
        error = _text(visual.get("screenshot_error") or visual.get("ai_error")).lower()
        if service.get("host_mode") == "passive_fofa":
            return "被动FOFA证据"
        if "download is starting" in error:
            return "下载响应"
        if "hard timeout" in error or "timeout" in error:
            return "截图超时"
        if "err_empty_response" in error or "server disconnected" in error:
            return "空响应/连接断开"
        if "err_connection_closed" in error or "connection closed" in error:
            return "连接关闭"
        if visual.get("ai_error"):
            return "AI分析失败"
        return "HTTP探测降级"

    def _is_low_value_title(self, title: str) -> bool:
        value = title.lower().strip()
        return any(marker in value for marker in LOW_VALUE_WEB_TITLE_MARKERS)

    def _service_audit_rows(
        self,
        bundle: dict,
        domain_units: dict[str, str],
        ip_units: dict[str, str],
        dns_unit_by_ip: dict[str, str],
        review_attestations: dict[str, dict[str, dict]] | None = None,
    ) -> list[dict]:
        probes_by_key: dict[tuple[str, int], list[dict]] = {}
        for probe in bundle.get("web_probe_results", []):
            probes_by_key.setdefault((probe.get("target_ip"), probe.get("port")), []).append(probe)
        entries_by_service: dict[int, list[dict]] = {}
        entries_by_url: dict[str, list[dict]] = {}
        for entry in bundle.get("web_entrypoints", []):
            normalized_url = _normalize_url_for_match(entry.get("normalized_url"))
            if normalized_url:
                entries_by_url.setdefault(normalized_url, []).append(entry)
            if entry.get("service_asset_id"):
                entries_by_service.setdefault(entry["service_asset_id"], []).append(entry)
        rows = []
        for service in sorted(bundle.get("service_assets", []), key=lambda item: (item.get("target_ip") or "", item.get("port") or 0)):
            probes = probes_by_key.get((service.get("target_ip"), service.get("port")), [])
            responded = [row for row in probes if row.get("status") == "responded"]
            domains = sorted(set(service.get("domains") or []))
            unit = (
                next((self._unit_for_host(domain, domain_units) for domain in domains if self._unit_for_host(domain, domain_units)), "")
                or ip_units.get(service.get("target_ip") or "")
                or dns_unit_by_ip.get(service.get("target_ip") or "", "")
            )
            kind = service.get("asset_kind") or "unknown"
            web_like = self._looks_web_like_service(service)
            entries = entries_by_service.get(service.get("id"), [])
            entry_count = len(entries)
            entry_reuse = False
            if not entries:
                representative_url = _normalize_url_for_match(service.get("representative_url"))
                entries = entries_by_url.get(representative_url, []) if representative_url else []
                entry_count = len(entries)
                entry_reuse = bool(entries)
            review_priority = "无"
            if kind != "web" and web_like:
                review_priority = "中"
            elif kind == "web" and entry_count == 0:
                review_priority = "高"
            elif kind == "web" and service.get("host_mode") == "passive_fofa":
                review_priority = "低"
            elif kind == "unknown":
                review_priority = "中"
            row = {
                    "复核优先级": review_priority,
                    "复核类型": self._service_review_type(kind, service, web_like, entry_count),
                    "单位": unit,
                    "IP": service.get("target_ip") or "",
                    "协议": service.get("protocol") or "",
                    "端口": service.get("port") or "",
                    "资产类型": kind,
                    "主机模式": service.get("host_mode") or "",
                    "服务": service.get("service") or "",
                    "产品": service.get("product") or "",
                    "版本": service.get("version") or "",
                    "代表URL": service.get("representative_url") or "",
                    "URL入口数量": entry_count,
                    "URL入口覆盖方式": "代表URL复用覆盖" if entry_reuse else ("直接关联" if entry_count else ""),
                    "关联域名": ", ".join(domains[:20]),
                    "HTTP响应数": len(responded),
                    "HTTP失败数": len(probes) - len(responded),
                    "HTTP状态分布": self._http_status_summary(responded),
                    "标题": service.get("title") or "",
                    "应用线索": service.get("app_name") or "",
                    "分类依据": self._service_classification_reason(kind, service, probes, entry_count),
                    "建议动作": self._service_audit_action(review_priority, kind, web_like, entry_count, service.get("host_mode") or ""),
                }
            self._apply_review_attestation(
                row,
                review_attestations or {},
                "service_classification",
                f"{service.get('target_ip') or ''}:{service.get('port') or ''}",
            )
            rows.append(row)
        priority_order = {"高": 0, "中": 1, "低": 2, "无": 3}
        return sorted(rows, key=lambda item: (priority_order.get(item["复核优先级"], 9), item["单位"], item["IP"], item["端口"]))

    def _service_review_type(self, kind: str, service: dict, web_like: bool, entry_count: int) -> str:
        if kind == "web" and entry_count == 0:
            return "missing_url_entry"
        if kind == "web" and service.get("host_mode") == "passive_fofa":
            return "passive_fofa_review"
        if kind != "web" and web_like:
            return "web_like_non_web"
        if kind == "unknown":
            return "unknown_service"
        return "none"

    def _url_coverage_rows(self, bundle: dict, service_audit_rows: list[dict], review_attestations: dict[str, dict[str, dict]] | None = None) -> list[dict]:
        service_ids = {row.get("id") for row in bundle.get("service_assets", []) if row.get("id")}
        entries_by_service: dict[int, list[dict]] = {}
        orphan_entries = []
        entries_by_url: dict[str, list[dict]] = {}
        for entry in bundle.get("web_entrypoints", []):
            normalized_url = _normalize_url_for_match(entry.get("normalized_url"))
            if normalized_url:
                entries_by_url.setdefault(normalized_url, []).append(entry)
            service_id = entry.get("service_asset_id")
            if service_id in service_ids:
                entries_by_service.setdefault(service_id, []).append(entry)
            else:
                orphan_entries.append(entry)
        rows = []
        service_by_key = {
            (row["IP"], row["端口"]): row
            for row in service_audit_rows
        }
        service_by_id = {row.get("id"): row for row in bundle.get("service_assets", [])}
        for service_id, service in service_by_id.items():
            if service.get("asset_kind") != "web":
                continue
            entries = entries_by_service.get(service_id, [])
            coverage_source = "service_asset/web_probe"
            coverage_conclusion = "已生成URL入口" if entries else "Web服务未生成URL入口"
            if not entries:
                representative_url = _normalize_url_for_match(service.get("representative_url"))
                entries = entries_by_url.get(representative_url, []) if representative_url else []
                if entries:
                    coverage_source = "service_asset/web_probe/representative_url_reuse"
                    coverage_conclusion = "代表URL已由其他服务入口覆盖"
            audit = service_by_key.get((service.get("target_ip") or "", service.get("port") or ""), {})
            row = {
                    "复核优先级": "高" if not entries else "无",
                    "单位": audit.get("单位", ""),
                    "IP": service.get("target_ip") or "",
                    "端口": service.get("port") or "",
                    "主机模式": service.get("host_mode") or "",
                    "代表URL": service.get("representative_url") or "",
                    "URL入口数量": len(entries),
                    "URL入口样例": ", ".join((entry.get("normalized_url") or "") for entry in entries[:10]),
                    "来源": coverage_source,
                    "覆盖结论": coverage_conclusion,
                    "建议动作": "继续视觉识别和风险复核。" if entries else "复核代表URL、Host绑定、HTTP状态过滤和URL入口生成规则。",
                }
            self._apply_review_attestation(row, review_attestations or {}, "url_entrypoint", row.get("URL入口样例") or f"{row.get('IP')}:{row.get('端口')}")
            rows.append(row)
        for entry in orphan_entries:
            audit = service_by_key.get((entry.get("target_ip") or "", entry.get("port") or ""), {})
            row = {
                    "复核优先级": "低",
                    "单位": audit.get("单位", ""),
                    "IP": entry.get("target_ip") or "",
                    "端口": entry.get("port") or "",
                    "主机模式": audit.get("主机模式", ""),
                    "代表URL": "",
                    "URL入口数量": 1,
                    "URL入口样例": entry.get("normalized_url") or "",
                    "来源": (entry.get("evidence") or {}).get("source") or "web_entrypoint",
                    "覆盖结论": "URL入口未关联到服务资产",
                    "建议动作": "复核是否为重复URL、重定向后URL或历史遗留入口；必要时补充服务关联。",
                }
            self._apply_review_attestation(row, review_attestations or {}, "url_entrypoint", row.get("URL入口样例") or f"{row.get('IP')}:{row.get('端口')}")
            rows.append(row)
        priority_order = {"高": 0, "中": 1, "低": 2, "无": 3}
        return sorted(rows, key=lambda item: (priority_order.get(item["复核优先级"], 9), item["单位"], item["IP"], item["端口"], item["URL入口样例"]))

    def _looks_web_like_service(self, service: dict) -> bool:
        text = " ".join(
            _text(service.get(key))
            for key in ("service", "product", "title", "app_name", "representative_url")
        ).lower()
        if service.get("port") in {80, 81, 82, 83, 84, 85, 88, 90, 443, 440, 8000, 8001, 8008, 8080, 8081, 8088, 8089, 8090, 8092, 8443, 8888, 8900, 9000, 9080, 9081, 9443, 9980, 18080, 19080}:
            return True
        return any(marker in text for marker in ("http", "https", "nginx", "apache", "tomcat", "iis", "websphere", "weblogic", "web"))

    def _http_status_summary(self, probes: list[dict]) -> str:
        counts = Counter(str(row.get("http_status")) for row in probes if row.get("http_status") is not None)
        return ", ".join(f"{status}={count}" for status, count in counts.most_common(8))

    def _service_classification_reason(self, kind: str, service: dict, probes: list[dict], entry_count: int) -> str:
        responded = [row for row in probes if row.get("status") == "responded"]
        if kind == "web":
            if service.get("host_mode") == "passive_fofa":
                return f"主动HTTP探测未形成响应，依据FOFA被动Web证据保留为Web入口，URL入口={entry_count}。"
            return f"HTTP探测有响应({len(responded)}个)，主机模式={service.get('host_mode') or 'unknown'}，URL入口={entry_count}。"
        if self._looks_web_like_service(service):
            return f"端口或产品疑似Web，但HTTP探测未形成Web资产；响应={len(responded)}，失败={len(probes) - len(responded)}。"
        return f"服务指纹偏非Web，HTTP探测响应={len(responded)}，失败={len(probes) - len(responded)}。"

    def _service_audit_action(self, priority: str, kind: str, web_like: bool, entry_count: int, host_mode: str = "") -> str:
        if kind == "web" and host_mode == "passive_fofa":
            return "人工确认 FOFA Host 是否为真实业务入口；如需主动验证，补充 Host/访问源后重跑服务识别。"
        if kind == "web" and entry_count == 0:
            return "优先复核URL入口生成规则、允许HTTP状态码、Host绑定和重定向结果。"
        if web_like and kind != "web":
            return "人工复核该端口是否为Web服务；必要时扩大HTTP探测状态码、补充Host或手工URL。"
        if priority == "中":
            return "补充服务指纹或复跑高强度服务识别。"
        return "保持台账，后续按周期复测。"

    def _visual_review_rows(self, web_rows: list[dict], review_attestations: dict[str, dict[str, dict]] | None = None) -> list[dict]:
        rows = []
        for row in web_rows:
            method = row.get("识别方式") or ""
            error = row.get("分析错误") or ""
            confidence = self._float_or_none(row.get("识别置信度"))
            low_value_page = self._is_low_value_visual_row(row)
            reasons = []
            if error:
                reasons.append("截图或AI识别失败")
            if not method:
                reasons.append("缺少视觉识别结果")
            elif method == "http_probe_fallback":
                category = row.get("降级类型") or "HTTP探测降级"
                reasons.append(f"HTTP探测信息降级识别({category})")
            if self._has_mojibake_visual_text(row):
                reasons.append("标题或识别结果疑似编码乱码")
            if confidence is not None and confidence < 0.5:
                reasons.append(
                    f"低价值错误/拦截页面置信度偏低({confidence:.2f})"
                    if low_value_page
                    else f"AI识别置信度偏低({confidence:.2f})"
                )
            if not reasons:
                continue
            title_fallback_ok = method == "http_probe_fallback" and self._has_meaningful_visual_label(row)
            priority = "高" if error or not method else ("低" if low_value_page or title_fallback_ok else ("中" if confidence is not None and confidence < 0.5 else "低"))
            review_type = self._visual_review_type(method, error, confidence, row)
            review_row = {
                    "复核优先级": priority,
                    "复核类型": review_type,
                    "单位": row.get("单位") or "",
                    "URL": row.get("URL") or "",
                    "IP": row.get("IP") or "",
                    "端口": row.get("端口") or "",
                    "HTTP状态": row.get("HTTP状态") or "",
                    "HTML标题": row.get("HTML标题") or "",
                    "识别方式": method,
                    "识别置信度": "" if confidence is None else confidence,
                    "AI识别系统": row.get("AI识别系统") or "",
                    "网站用途": row.get("网站用途") or "",
                    "降级类型": row.get("降级类型") or "",
                    "降级原因": row.get("降级原因") or "",
                    "复核原因": "；".join(reasons),
                    "建议动作": self._visual_review_action(priority, method, error, review_type),
                    "截图": row.get("截图") or "",
                    "分析错误": error,
                }
            self._apply_review_attestation(review_row, review_attestations or {}, "visual_identification", review_row.get("URL") or "")
            rows.append(review_row)
        priority_order = {"高": 0, "中": 1, "低": 2}
        return sorted(rows, key=lambda item: (priority_order.get(item["复核优先级"], 9), item["单位"], item["URL"]))

    def _is_low_value_visual_row(self, row: dict) -> bool:
        status = str(row.get("HTTP状态") or "")
        title = _text(row.get("HTML标题") or row.get("AI识别系统") or "")
        if status in {"400", "401", "403", "404", "500", "502", "503"}:
            return True
        return self._is_low_value_title(title)

    def _has_meaningful_visual_label(self, row: dict) -> bool:
        text = _text(row.get("AI识别系统") or row.get("HTML标题") or "").strip()
        if not text or self._is_low_value_title(text) or self._has_mojibake_visual_text(row):
            return False
        bad_chars = sum(1 for char in text if char == "\ufffd" or "\u0230" <= char <= "\u024f")
        return bad_chars / max(len(text), 1) < 0.2

    def _has_mojibake_visual_text(self, row: dict) -> bool:
        text = _text(row.get("AI识别系统") or "") + _text(row.get("HTML标题") or "")
        if not text:
            return False
        if "\ufffd" in text or "���" in text:
            return True
        weird = sum(1 for char in text if "\u0600" <= char <= "\u06ff" or "\u0230" <= char <= "\u024f")
        letters = sum(1 for char in text if char.isalpha())
        return letters >= 2 and weird / max(letters, 1) > 0.4

    def _float_or_none(self, value: Any) -> float | None:
        if value is None or value == "":
            return None
        try:
            return float(value)
        except (TypeError, ValueError):
            return None

    def _visual_review_action(self, priority: str, method: str, error: str, review_type: str = "") -> str:
        if error or not method:
            return "优先重新截图识别；如页面需要登录态或访问源白名单，补充人工截图和业务用途确认。"
        if review_type == "manual_low_value_page_review":
            return "确认是否为错误页、默认页、拦截页或停放页；若背后存在真实业务入口，补充正确 Host、URL 或登录态截图。"
        if priority == "中":
            return "人工核对截图、标题和业务用途，必要时补充系统名称。"
        return "抽样人工复核降级识别结论，确认是否需要补充截图或登录态验证。"

    def _visual_review_type(self, method: str, error: str, confidence: float | None, row: dict) -> str:
        if error or not method:
            return "automatic_retry"
        if method == "http_probe_fallback":
            return "manual_http_fallback_review"
        if self._is_low_value_visual_row(row):
            return "manual_low_value_page_review"
        if self._has_mojibake_visual_text(row):
            return "manual_text_review"
        if confidence is not None and confidence < 0.5:
            return "manual_low_confidence_review"
        return "manual_review"

    def _risk_rows(self, port_rows: list[dict], web_rows: list[dict]) -> list[dict]:
        rows: list[dict] = []
        seen = set()
        high_ports = {3306, 5432, 6379, 9200, 9300, 11211, 27017, 1521, 1433, 445, 139, 3389, 5900}
        medium_ports = {21, 22, 23, 25, 110, 143, 389, 873, 1099, 8080, 8081, 8089, 8090, 8092, 8888, 9000, 9090}
        for row in port_rows:
            port = int(row["端口"] or 0)
            service_text = " ".join(_text(row.get(key)) for key in ("服务", "产品", "Web标题", "Web URL")).lower()
            if port in high_ports:
                severity = "高"
                risk = "敏感服务暴露"
                advice = "确认是否必须互联网开放；如非必要应收敛访问源、加固认证并纳入漏洞与弱口令核查。"
            elif port in medium_ports:
                severity = "中"
                risk = "管理或高价值服务暴露"
                advice = "确认服务用途和访问控制策略，优先检查默认口令、弱口令、版本漏洞和访问日志。"
            elif "vpn" in service_text or "easyconnect" in service_text:
                severity = "高"
                risk = "远程接入入口暴露"
                advice = "核查远程接入入口的 MFA、账号治理、补丁版本、登录审计和暴力破解防护。"
            else:
                continue
            key = (row.get("单位"), row.get("IP"), port, risk)
            if key in seen:
                continue
            seen.add(key)
            rows.append(
                {
                    "风险ID": f"R-{len(rows) + 1:03d}",
                    "风险等级": severity,
                    "风险分值": self._risk_score(severity, risk, "开放端口"),
                    "风险类型": risk,
                    "单位": row.get("单位"),
                    "资产": f"{row.get('IP')}:{port}",
                    "证据": _short(f"{row.get('服务')} {row.get('产品')} {row.get('版本')} {row.get('Web标题')} {row.get('Web URL')}", 500),
                    "处置建议": advice,
                    "责任建议": self._risk_owner(risk),
                    "验收证据": self._risk_acceptance_evidence(severity, risk),
                    "风险来源": "开放端口",
                }
            )
        for row in web_rows:
            text = " ".join(_text(row.get(key)) for key in ("AI识别系统", "HTML标题", "网站用途", "登录特征", "URL")).lower()
            if not text:
                continue
            if any(keyword in text for keyword in ("easyconnect", "easy connect", "vpn", "ssl vpn", "远程接入", "深信服", "sangfor")):
                severity, risk = "高", "远程接入 Web 入口"
                advice = "重点核查 MFA、补丁版本、弱口令、暴力破解防护和互联网访问控制。"
            elif any(keyword in text for keyword in ("登录", "login", "管理", "admin", "后台", "portal", "门户")):
                severity, risk = "中", "登录/管理入口暴露"
                advice = "核查账号口令策略、验证码/限速、默认账号、权限边界和组件漏洞。"
            else:
                continue
            key = (row.get("单位"), row.get("URL"), risk)
            if key in seen:
                continue
            seen.add(key)
            rows.append(
                {
                    "风险ID": f"R-{len(rows) + 1:03d}",
                    "风险等级": severity,
                    "风险分值": self._risk_score(severity, risk, "Web入口"),
                    "风险类型": risk,
                    "单位": row.get("单位"),
                    "资产": row.get("URL"),
                    "证据": _short(row.get("AI识别系统") or row.get("HTML标题") or row.get("网站用途"), 500),
                    "处置建议": advice,
                    "责任建议": self._risk_owner(risk),
                    "验收证据": self._risk_acceptance_evidence(severity, risk),
                    "风险来源": "Web入口",
                }
            )
        severity_order = {"高": 0, "中": 1, "低": 2}
        return sorted(rows, key=lambda item: (severity_order.get(item["风险等级"], 9), -int(item.get("风险分值") or 0), item["单位"] or "", item["资产"] or ""))

    def _risk_score(self, severity: str, risk_type: str, source: str) -> int:
        base = {"高": 90, "中": 60, "低": 30}.get(severity, 45)
        if "敏感服务" in risk_type:
            base += 6
        if "远程接入" in risk_type:
            base += 5
        if "管理" in risk_type or "登录" in risk_type:
            base += 4
        if source == "Web入口":
            base += 2
        return min(base, 100)

    def _risk_owner(self, risk_type: str) -> str:
        if "远程接入" in risk_type:
            return "系统负责人 / 身份与访问管理负责人 / 网络边界负责人"
        if "敏感服务" in risk_type:
            return "系统负责人 / 数据库或中间件管理员 / 网络边界负责人"
        if "管理" in risk_type or "登录" in risk_type:
            return "系统负责人 / 应用运维负责人 / 账号权限管理员"
        return "资产归属单位 / 系统负责人"

    def _risk_acceptance_evidence(self, severity: str, risk_type: str) -> str:
        if severity == "高":
            return "访问控制策略、强认证/MFA、补丁或版本核查、弱口令核查、复测截图或端口复测结果。"
        if "登录" in risk_type or "管理" in risk_type:
            return "账号策略、验证码/限速、默认账号核查、组件版本核查和复测截图。"
        return "资产归属确认、业务必要性说明、复测记录和监控纳管截图。"

    def _coverage_rows(
        self,
        bundle: dict,
        dns_rows: list[dict],
        dns_quality_rows: list[dict],
        port_rows: list[dict],
        web_rows: list[dict],
        service_audit_rows: list[dict],
        url_coverage_rows: list[dict],
        unit_coverage_rows: list[dict],
        port_target_rows: list[dict],
    ) -> list[dict]:
        company_names = {row["id"]: row["name"] for row in bundle["companies"]}
        company_ids = set(company_names)
        asset_company_ids = {asset["company_id"] for asset in bundle["assets"]}
        companies_without_assets = [row for row in unit_coverage_rows if row.get("覆盖状态") == "无资产线索"]
        high_medium_asset_gaps = [
            row for row in companies_without_assets if row.get("复核优先级") in {"高", "中"}
        ]
        low_asset_gaps = [row for row in companies_without_assets if row.get("复核优先级") == "低"]

        root_domains = {
            asset["normalized_identifier"].lower().rstrip(".")
            for asset in bundle["assets"]
            if asset.get("asset_type") == "icp_domain"
        }
        dns_hosts = {row["主机名"].lower().rstrip(".") for row in dns_rows}
        domains_with_dns = {
            domain
            for domain in root_domains
            if any(host == domain or host.endswith(f".{domain}") for host in dns_hosts)
        }
        domains_without_dns = sorted(root_domains - domains_with_dns)
        enumeration_attempted = bool(bundle.get("subdomain_task") or bundle.get("subdomain_tool_runs"))
        subdomain_roots = {
            row.get("root_domain")
            for row in bundle.get("subdomains", [])
            if row.get("root_domain")
        }
        roots_without_subdomains = sorted(root_domains - subdomain_roots) if enumeration_attempted else []
        failed_tool_runs = [
            row
            for row in bundle.get("subdomain_tool_runs", [])
            if row.get("status") == "failed"
        ]
        dns_review_rows = [row for row in dns_quality_rows if row.get("复核优先级") not in {"", "无"}]
        medium_dns_review_rows = [row for row in dns_review_rows if row.get("复核优先级") == "中"]
        low_dns_review_rows = [row for row in dns_review_rows if row.get("复核优先级") == "低"]

        dns_ips = {row["记录值"] for row in dns_rows if row.get("记录类型") in {"A", "AAAA"}}
        manual_ips = {
            asset["normalized_identifier"]
            for asset in bundle["assets"]
            if asset.get("asset_type") == "ip"
        }
        raw_candidate_ips = dns_ips | manual_ips
        nmap_task = bundle.get("nmap_task") or {}
        task_targets = set()
        if isinstance(nmap_task.get("targets"), list):
            task_targets = {str(target) for target in nmap_task["targets"] if target}
        port_target_ips = {row.get("IP") for row in port_target_rows if row.get("IP")}
        candidate_ips = port_target_ips or task_targets or raw_candidate_ips
        observed_ips = {
            run.get("target_ip")
            for run in bundle.get("nmap_runs", [])
            if run.get("target_ip") and run.get("target_ip") != "__batch__"
        } | {
            port.get("target_ip")
            for port in bundle.get("nmap_ports", [])
            if port.get("target_ip")
        }
        observed_ips.update(task_targets)
        ips_without_port_evidence = sorted(candidate_ips - observed_ips)
        passive_only_port_rows = [row for row in port_rows if row.get("证据类型") == "被动FOFA"]
        merged_evidence_port_rows = [row for row in port_rows if row.get("证据类型") == "主动+被动"]
        active_port_rows = [row for row in port_rows if row.get("主动扫描确认") == "是"]

        identified_web = [row for row in web_rows if row.get("识别方式")]
        fallback_web = [row for row in web_rows if row.get("识别方式") == "http_probe_fallback"]
        missing_visual = [row for row in web_rows if not row.get("识别方式")]

        web_service_rows = [
            row
            for row in bundle.get("service_assets", [])
            if row.get("asset_kind") == "web"
        ]
        service_url_coverage_rows = [
            row for row in url_coverage_rows if str(row.get("来源") or "").startswith("service_asset")
        ]
        web_services_without_entrypoint = [
            row for row in service_url_coverage_rows if row.get("覆盖结论") == "Web服务未生成URL入口"
        ]
        covered_web_services = max(0, len(web_service_rows) - len(web_services_without_entrypoint))
        service_review_rows = [row for row in service_audit_rows if row.get("复核优先级") not in {"", "无"}]
        url_review_rows = [row for row in url_coverage_rows if row.get("复核优先级") not in {"", "无"}]

        return [
            {
                "环节": "企业/备案资产",
                "指标": "单位资产归属覆盖",
                "结果": (
                    f"{len(asset_company_ids)}/{len(company_ids)} 家单位存在资产线索，"
                    f"待补充高/中优先级 {len(high_medium_asset_gaps)} 家，低优先级 {len(low_asset_gaps)} 家"
                ),
                "缺口等级": "中" if high_medium_asset_gaps else ("低" if low_asset_gaps else "无"),
                "缺口样例": ", ".join(row.get("单位") or "" for row in [*high_medium_asset_gaps, *low_asset_gaps][:10]),
                "建议动作": "对无资产单位补充核验备案、官网、公众号、小程序、APP、邮箱和人工线索。",
            },
            {
                "环节": "子域名/DNS",
                "指标": "根域名解析覆盖",
                "结果": f"{len(domains_with_dns)}/{len(root_domains)} 个根域名存在 DNS 证据",
                "缺口等级": "中" if domains_without_dns else "无",
                "缺口样例": ", ".join(domains_without_dns[:10]),
                "建议动作": "对无解析证据的根域名复核域名状态、DNS 服务器和历史解析数据源。",
            },
            {
                "环节": "子域名/DNS",
                "指标": "子域名枚举质量",
                "结果": (
                    f"{len(root_domains - set(roots_without_subdomains))}/{len(root_domains)} 个根域名发现子域名，"
                    f"工具失败 {len(failed_tool_runs)} 个"
                ),
                "缺口等级": "低" if roots_without_subdomains or failed_tool_runs else "无",
                "缺口样例": ", ".join([*roots_without_subdomains[:10], *(f"{row.get('tool_name')}:{row.get('root_domain')}" for row in failed_tool_runs[:10])]),
                "建议动作": "对未发现子域名或工具失败的根域名补充证书透明度、历史 DNS、搜索引擎和人工线索，并按需重跑失败工具。",
            },
            {
                "环节": "子域名/DNS",
                "指标": "DNS 解析复核质量",
                "结果": (
                    f"{len(dns_review_rows)}/{len(dns_quality_rows)} 个根域名需要人工复核，"
                    f"中优先级 {len(medium_dns_review_rows)} 个，低优先级 {len(low_dns_review_rows)} 个"
                ),
                "缺口等级": "中" if medium_dns_review_rows else ("低" if low_dns_review_rows else "无"),
                "缺口样例": ", ".join(row.get("根域名") or "" for row in dns_review_rows[:10]),
                "建议动作": "对无公网 A/AAAA、第三方/停放 CNAME、共享 IP 和工具失败线索进行人工确认；确认停放或无独立业务后在复核工作单留痕。",
            },
            {
                "环节": "端口发现",
                "指标": "候选 IP 端口发现覆盖",
                "结果": f"{len(observed_ips & candidate_ips)}/{len(candidate_ips)} 个候选 IP 存在端口发现任务或开放端口证据",
                "缺口等级": "高" if ips_without_port_evidence else "无",
                "缺口样例": ", ".join(ips_without_port_evidence[:20]),
                "建议动作": "对缺少端口发现证据的公网 IP 补跑主动扫描或 FOFA 被动查询；已扫描但无开放端口的 IP 可作为无开放端口证据保留。",
            },
            {
                "环节": "端口发现",
                "指标": "端口证据来源质量",
                "结果": (
                    f"开放端口 {len(port_rows)} 个，其中主动确认 {len(active_port_rows)} 个，"
                    f"被动FOFA-only {len(passive_only_port_rows)} 个，主动+被动交叉验证 {len(merged_evidence_port_rows)} 个"
                ),
                "缺口等级": "低" if passive_only_port_rows and not active_port_rows else "无",
                "缺口样例": ", ".join(f"{row.get('IP')}:{row.get('端口')}" for row in passive_only_port_rows[:20]),
                "建议动作": "被动FOFA-only端口应作为线索保留；如需交付前主动确认，可启用 nmap 后重跑端口发现，系统会精确扫描已有 FOFA 端口。",
            },
            {
                "环节": "服务识别/URL",
                "指标": "Web 服务入口覆盖",
                "结果": f"{covered_web_services}/{len(web_service_rows)} 个 Web 服务生成 URL 入口",
                "缺口等级": "中" if web_services_without_entrypoint else "无",
                "缺口样例": ", ".join(f"{row.get('IP')}:{row.get('端口')}" for row in web_services_without_entrypoint[:20]),
                "建议动作": "对未生成 URL 入口的 Web 服务复核协议、Host 绑定、HTTPS 证书和探测超时设置。",
            },
            {
                "环节": "服务识别/URL",
                "指标": "服务分类复核",
                "结果": f"{len(service_review_rows)}/{len(service_audit_rows)} 个服务需要复核",
                "缺口等级": "低" if service_review_rows else "无",
                "缺口样例": ", ".join(f"{row.get('IP')}:{row.get('端口')}" for row in service_review_rows[:20]),
                "建议动作": "对疑似 Web 但未形成 Web 资产、未知服务或缺少 URL 的服务进行人工复核或补充 Host/URL 后重跑分类。",
            },
            {
                "环节": "服务识别/URL",
                "指标": "URL 入口关联质量",
                "结果": f"{len(url_review_rows)}/{len(url_coverage_rows)} 个 URL 覆盖项需要复核",
                "缺口等级": "低" if url_review_rows else "无",
                "缺口样例": ", ".join((row.get("URL入口样例") or f"{row.get('IP')}:{row.get('端口')}") for row in url_review_rows[:10]),
                "建议动作": "对未关联服务资产或缺少 URL 的入口复核重定向、去重规则、服务关联和历史遗留入口。",
            },
            {
                "环节": "URL视觉识别",
                "指标": "Web 入口识别覆盖",
                "结果": f"{len(identified_web)}/{len(web_rows)} 个 Web 入口完成识别，其中 HTTP 降级 {len(fallback_web)} 个",
                "缺口等级": "中" if missing_visual else ("低" if fallback_web else "无"),
                "缺口样例": ", ".join((row.get("URL") or "") for row in (missing_visual or fallback_web)[:10]),
                "建议动作": "对未识别或降级识别页面补充人工复核、截图重试或业务登录态验证。",
            },
        ]

    def _unit_coverage_rows(
        self,
        company_names: dict[int, str],
        company_scope: dict[str, dict[str, Any]],
        unit_asset_rows: list[dict],
        dns_rows: list[dict],
        port_rows: list[dict],
        web_rows: list[dict],
        risk_rows: list[dict],
        manual_no_asset_reviews: dict[str, dict[str, Any]] | None = None,
    ) -> list[dict]:
        manual_no_asset_reviews = manual_no_asset_reviews or {}
        rows: dict[str, dict[str, Any]] = {
            name: {
                "单位": name,
                "股权层级": company_scope.get(name, {}).get("股权层级", ""),
                "直接持股": company_scope.get(name, {}).get("直接持股", ""),
                "累计持股": company_scope.get(name, {}).get("累计持股", ""),
                "股权路径": company_scope.get(name, {}).get("股权路径", name),
                "子公司数量": company_scope.get(name, {}).get("子公司数量", 0),
                "资产数量": 0,
                "根域名数量": 0,
                "子域名/DNS主机数量": 0,
                "IP数量": 0,
                "开放端口数量": 0,
                "Web入口数量": 0,
                "高风险数量": 0,
                "中风险数量": 0,
                "复核优先级": "",
                "缺口原因": "",
                "资产类型": set(),
                "代表性资产": [],
                "人工复核状态": "",
                "人工复核依据": "",
                "人工复核来源": "",
                "_ip_values": set(),
                "覆盖状态": "无资产线索",
                "建议动作": "优先补充该单位的备案域名、官网、公众号、小程序、APP、邮箱和公网 IP 线索。",
            }
            for name in sorted(company_names.values())
        }
        extra_units = {
            row.get("单位")
            for source in (unit_asset_rows, dns_rows, port_rows, web_rows, risk_rows)
            for row in source
            if row.get("单位")
        }
        for unit in sorted(extra_units):
            rows.setdefault(
                unit,
                {
                    "单位": unit,
                    "股权层级": company_scope.get(unit, {}).get("股权层级", ""),
                    "直接持股": company_scope.get(unit, {}).get("直接持股", ""),
                    "累计持股": company_scope.get(unit, {}).get("累计持股", ""),
                    "股权路径": company_scope.get(unit, {}).get("股权路径", unit),
                    "子公司数量": company_scope.get(unit, {}).get("子公司数量", 0),
                    "资产数量": 0,
                    "根域名数量": 0,
                    "子域名/DNS主机数量": 0,
                    "IP数量": 0,
                    "开放端口数量": 0,
                    "Web入口数量": 0,
                    "高风险数量": 0,
                    "中风险数量": 0,
                    "复核优先级": "",
                    "缺口原因": "",
                    "资产类型": set(),
                    "代表性资产": [],
                    "人工复核状态": "",
                    "人工复核依据": "",
                    "人工复核来源": "",
                    "_ip_values": set(),
                    "覆盖状态": "无资产线索",
                    "建议动作": "优先补充该单位的备案域名、官网、公众号、小程序、APP、邮箱和公网 IP 线索。",
                },
            )
        for asset in unit_asset_rows:
            unit = asset.get("单位") or ""
            if not unit:
                continue
            row = rows[unit]
            row["资产数量"] += 1
            asset_type = asset.get("资产类型") or ""
            row["资产类型"].add(asset_type)
            if asset_type == "icp_domain":
                row["根域名数量"] += 1
            elif asset_type == "ip":
                row["_ip_values"].add(asset.get("资产标识") or "")
            if len(row["代表性资产"]) < 5:
                row["代表性资产"].append(asset.get("资产标识") or asset.get("资产名称") or "")
        dns_hosts: dict[str, set[str]] = {}
        for dns in dns_rows:
            unit = dns.get("单位") or ""
            if unit in rows:
                dns_hosts.setdefault(unit, set()).add(dns.get("主机名") or "")
                if dns.get("记录类型") in {"A", "AAAA"} and dns.get("记录值"):
                    rows[unit]["_ip_values"].add(dns.get("记录值") or "")
        for unit, hosts in dns_hosts.items():
            rows[unit]["子域名/DNS主机数量"] = len({host for host in hosts if host})
        for port in port_rows:
            unit = port.get("单位") or ""
            if unit in rows:
                rows[unit]["开放端口数量"] += 1
        for web in web_rows:
            unit = web.get("单位") or ""
            if unit in rows:
                rows[unit]["Web入口数量"] += 1
        for risk in risk_rows:
            unit = risk.get("单位") or ""
            if unit not in rows:
                continue
            if risk.get("风险等级") == "高":
                rows[unit]["高风险数量"] += 1
            elif risk.get("风险等级") == "中":
                rows[unit]["中风险数量"] += 1
        output = []
        for row in rows.values():
            row["IP数量"] = len({item for item in row["_ip_values"] if item})
            if row["Web入口数量"] or row["开放端口数量"]:
                row["覆盖状态"] = "已覆盖互联网暴露面"
                row["建议动作"] = "持续复测端口、Web 指纹和风险处置状态。"
            elif row["资产数量"] or row["子域名/DNS主机数量"] or row["IP数量"]:
                row["覆盖状态"] = "有资产线索，待扩大探测"
                row["建议动作"] = "继续补充端口发现、URL 识别和人工核验。"
            elif row["单位"] in manual_no_asset_reviews:
                review = manual_no_asset_reviews[row["单位"]]
                row["覆盖状态"] = "人工确认无独立互联网资产"
                row["人工复核状态"] = review.get("review_status") or "no_assets_found"
                row["人工复核依据"] = review.get("notes") or review.get("review_notes") or "人工复核未发现独立互联网资产"
                row["人工复核来源"] = ", ".join(review.get("source_urls") or [])
                row["建议动作"] = "保留人工复核依据，后续周期复核是否新增独立资产。"
            row["复核优先级"], row["缺口原因"] = self._unit_asset_gap_priority(row)
            row["资产类型"] = ", ".join(sorted(item for item in row["资产类型"] if item))
            row["代表性资产"] = ", ".join(item for item in row["代表性资产"] if item)
            row.pop("_ip_values", None)
            output.append(row)
        return sorted(
            output,
            key=lambda item: (
                0 if item["高风险数量"] else 1,
                0 if item["Web入口数量"] else 1,
                0 if item["资产数量"] else 1,
                {"高": 0, "中": 1, "低": 2, "无": 3, "": 4}.get(item.get("复核优先级"), 4),
                item["单位"],
            ),
        )

    def _unit_asset_gap_priority(self, row: dict) -> tuple[str, str]:
        if row.get("覆盖状态") == "已覆盖互联网暴露面":
            return "无", "已形成端口或 Web 暴露面证据"
        if row.get("覆盖状态") == "有资产线索，待扩大探测":
            return "低", "已有备案/公众号/APP/IP 等线索，待继续扩大探测"
        if row.get("覆盖状态") == "人工确认无独立互联网资产":
            return "无", "人工复核确认未发现独立互联网资产"
        depth = self._int_value(row.get("股权层级"), default=99)
        direct = self._percent_value(row.get("直接持股"))
        cumulative = self._percent_value(row.get("累计持股"))
        if depth <= 1:
            return "高", "根公司或一级核心平台单位暂无资产线索"
        child_count = self._int_value(row.get("子公司数量"), default=0)
        if depth <= 2 and child_count > 0 and (direct >= 50 or cumulative >= 50):
            return "中", "二级控股平台单位暂无资产线索"
        return "低", "深层项目公司或低累计持股单位暂无公开资产线索"

    def _manual_no_asset_reviews(self, bundle: dict) -> dict[str, dict[str, Any]]:
        reviews: dict[str, dict[str, Any]] = {}
        no_asset_statuses = {"no_assets_found", "无资产", "确认无资产", "未发现资产", "confirmed_no_assets"}
        for record in bundle.get("raw_records", []):
            payload = record.get("response_json") if isinstance(record.get("response_json"), dict) else {}
            if record.get("source") == "manual_import" and record.get("action") == "no_assets_found":
                unit = _text(payload.get("unit"))
            elif (
                record.get("source") == "review_workorder"
                and record.get("action") == "review_attestation"
                and payload.get("category") == "asset_supplement"
                and _text(payload.get("review_status")) in no_asset_statuses
            ):
                unit = _text(payload.get("unit") or payload.get("item_key"))
            else:
                continue
            if unit:
                reviews[unit] = payload
        return reviews

    def _int_value(self, value: Any, *, default: int = 0) -> int:
        try:
            return int(value)
        except (TypeError, ValueError):
            return default

    def _percent_value(self, value: Any) -> float:
        if value is None:
            return 0.0
        text = str(value).strip().rstrip("%")
        try:
            return float(text)
        except ValueError:
            return 0.0

    def _remediation_rows(self, risk_rows: list[dict]) -> list[dict]:
        sla = {"高": "7 天内完成确认与收敛", "中": "30 天内完成加固", "低": "90 天内纳入常规治理"}
        accept = {
            "高": "互联网侧访问面已收敛或具备强认证、访问控制、补丁和审计证据。",
            "中": "确认业务必要性，完成账号、口令、版本、访问控制和日志审计加固。",
            "低": "资产归属明确，纳入台账、监控和周期性复测。"
        }
        rows = []
        for index, risk in enumerate(risk_rows, start=1):
            severity = risk.get("风险等级") or "中"
            risk_type = risk.get("风险类型") or ""
            score = risk.get("风险分值") or self._risk_score(severity, risk_type, risk.get("风险来源") or "")
            rows.append(
                {
                    "序号": index,
                    "优先级": severity,
                    "单位": risk.get("单位") or "",
                    "资产": risk.get("资产") or "",
                    "问题": risk_type,
                    "风险分值": score,
                    "整改动作": risk.get("处置建议") or "",
                    "责任建议": risk.get("责任建议") or self._risk_owner(risk_type),
                    "建议时限": sla.get(severity, sla["中"]),
                    "验收标准": accept.get(severity, accept["中"]),
                    "验收证据": risk.get("验收证据") or self._risk_acceptance_evidence(severity, risk_type),
                    "当前状态": "待确认",
                }
            )
        return rows

    def _overview_rows(self, stats: dict) -> list[dict]:
        method_counts = stats.get("Web识别方式分布", {})
        return [
            {"指标": "单位数量", "结果": stats.get("单位数量", 0), "说明": "纳入报告范围的主体单位数量。"},
            {"指标": "有资产线索单位", "结果": stats.get("有资产线索单位", 0), "说明": "存在基础资产、DNS、IP、端口或 Web 线索的单位数量。"},
            {"指标": "有Web入口单位", "结果": stats.get("有Web入口单位", 0), "说明": "存在可访问 Web 入口或降级识别入口的单位数量。"},
            {"指标": "资产数量", "结果": stats.get("资产数量", 0), "说明": "企业采集与手工补充后的资产总量。"},
            {"指标": "DNS记录数量", "结果": stats.get("DNS记录数量", 0), "说明": "域名、子域名及解析记录证据数量。"},
            {"指标": "DNS复核项", "结果": stats.get("DNS复核项", 0), "说明": "需要复核的无子域名、工具失败、无公网解析或第三方解析线索数量。"},
            {"指标": "端口目标数量", "结果": stats.get("端口目标数量", 0), "说明": "候选来源、扫描目标或开放端口证据合并后的公网 IP 数量。"},
            {"指标": "开放端口数量", "结果": stats.get("开放端口数量", 0), "说明": "主动/被动端口发现合并后的开放端口数量。"},
            {"指标": "服务复核项", "结果": stats.get("服务复核项", 0), "说明": "疑似 Web 但未形成 Web 资产、未知服务或 Web 服务缺少 URL 的复核项。"},
            {"指标": "Web入口数量", "结果": stats.get("Web入口数量", 0), "说明": "可用于网站识别和截图取证的 Web 入口数量。"},
            {"指标": "URL入口复核项", "结果": stats.get("URL入口复核项", 0), "说明": "Web 服务未生成 URL 或 URL 未关联服务资产的复核项。"},
            {"指标": "Web识别覆盖率", "结果": stats.get("Web识别覆盖率", "0%"), "说明": "已完成截图、复用或降级识别的 Web 入口比例。"},
            {"指标": "截图AI识别", "结果": method_counts.get("screenshot_ai", 0), "说明": "通过浏览器截图并交由 AI 识别的页面数量。"},
            {"指标": "重复页面复用", "结果": method_counts.get("duplicate_reuse", 0), "说明": "相同页面内容复用已有识别结果的页面数量。"},
            {"指标": "HTTP降级识别", "结果": method_counts.get("http_probe_fallback", 0), "说明": "截图失败但依据 HTTP 探测信息保留的页面数量。"},
            {"指标": "视觉复核项", "结果": stats.get("视觉复核项", 0), "说明": "需要人工复核的截图失败、降级识别或低置信度 Web 入口数量。"},
            {"指标": "覆盖缺口项", "结果": stats.get("覆盖缺口项", 0), "说明": "仍需复核或补扫的资产链路缺口数量。"},
            {"指标": "高风险项", "结果": stats.get("高风险项", 0), "说明": "建议优先确认与收敛的高风险暴露项。"},
            {"指标": "中风险项", "结果": stats.get("中风险项", 0), "说明": "建议纳入近期加固计划的风险项。"},
        ]

    def _risk_summary_rows(self, risk_rows: list[dict]) -> list[dict]:
        summary: dict[tuple[str, str], int] = {}
        for row in risk_rows:
            key = (row.get("风险等级") or "未定级", row.get("风险类型") or "未分类")
            summary[key] = summary.get(key, 0) + 1
        severity_order = {"高": 0, "中": 1, "低": 2, "未定级": 9}
        return [
            {"风险等级": level, "风险类型": risk_type, "数量": count}
            for (level, risk_type), count in sorted(
                summary.items(),
                key=lambda item: (severity_order.get(item[0][0], 9), item[0][1]),
            )
        ]

    def _dashboard_rows(self, stats: dict, risk_summary_rows: list[dict], coverage_rows: list[dict]) -> list[dict]:
        rows: list[dict] = []
        for level in ("高", "中", "低"):
            rows.append({"分组": "风险等级", "指标": level, "数值": stats.get(f"{level}风险项", 0), "说明": "风险等级分布"})
        for asset_type, count in (stats.get("资产类型分布") or {}).items():
            rows.append({"分组": "资产类型", "指标": asset_type, "数值": count, "说明": "基础资产类型分布"})
        for method, count in (stats.get("Web识别方式分布") or {}).items():
            rows.append({"分组": "Web识别方式", "指标": method, "数值": count, "说明": "URL视觉识别来源分布"})
        rows.extend(
            [
                {"分组": "单位覆盖", "指标": "有资产线索单位", "数值": stats.get("有资产线索单位", 0), "说明": "单位覆盖"},
                {"分组": "单位覆盖", "指标": "有Web入口单位", "数值": stats.get("有Web入口单位", 0), "说明": "单位覆盖"},
                {"分组": "单位覆盖", "指标": "无资产线索单位", "数值": max(stats.get("单位数量", 0) - stats.get("有资产线索单位", 0), 0), "说明": "单位覆盖"},
            ]
        )
        for row in coverage_rows:
            if row.get("缺口等级") not in {"", "无"}:
                rows.append({"分组": "覆盖缺口", "指标": f"{row.get('环节')}:{row.get('指标')}", "数值": 1, "说明": row.get("缺口等级")})
        for row in risk_summary_rows[:10]:
            rows.append(
                {
                    "分组": "风险类型Top",
                    "指标": f"{row.get('风险等级')}-{row.get('风险类型')}",
                    "数值": row.get("数量", 0),
                    "说明": "风险类型分布",
                }
            )
        return rows

    def _key_asset_rows(self, risk_rows: list[dict], port_target_rows: list[dict], web_rows: list[dict]) -> list[dict]:
        web_by_url = {row.get("URL"): row for row in web_rows if row.get("URL")}
        port_by_ip = {row.get("IP"): row for row in port_target_rows if row.get("IP")}
        rows = []
        for risk in risk_rows:
            asset = risk.get("资产") or ""
            web = web_by_url.get(asset, {})
            ip = asset.split(":", 1)[0] if ":" in asset else (web.get("IP") or "")
            port_target = port_by_ip.get(ip, {})
            rows.append(
                {
                    "优先级": risk.get("风险等级") or "",
                    "单位": risk.get("单位") or web.get("单位") or port_target.get("归属单位") or "",
                    "资产": asset,
                    "资产类型": "Web入口" if asset.startswith(("http://", "https://")) else "开放端口",
                    "风险分值": risk.get("风险分值") or "",
                    "风险类型": risk.get("风险类型") or "",
                    "系统/标题": web.get("AI识别系统") or web.get("HTML标题") or port_target.get("开放端口") or "",
                    "证据摘要": risk.get("证据") or "",
                    "处置建议": risk.get("处置建议") or "",
                    "责任建议": risk.get("责任建议") or self._risk_owner(risk.get("风险类型") or ""),
                    "关联域名/入口": web.get("URL") or port_target.get("关联域名样例") or "",
                }
            )
        severity_order = {"高": 0, "中": 1, "低": 2}
        return sorted(rows, key=lambda item: (severity_order.get(item["优先级"], 9), -int(item.get("风险分值") or 0), item["单位"], item["资产"]))

    def _navigation_rows(self, workbook_kind: str, sheets: dict[str, list[dict]]) -> list[dict]:
        descriptions = {
            "报告概览": ("先看", "核心指标、覆盖率、风险与复核项汇总。"),
            "管理驾驶舱": ("先看", "风险、资产类型、识别方式和单位覆盖的图表化摘要。"),
            "AI分析审计": ("交付", "报告 AI 分块分析状态、模型、输入规模和可追溯指纹。"),
            "风险统计": ("先看", "按风险等级和类型汇总风险数量。"),
            "重点资产视图": ("先看", "从风险视角汇总最需要优先确认和处置的资产。"),
            "单位覆盖台账": ("复核", "按单位查看资产线索、股权路径、端口和 Web 覆盖情况。"),
            "资产汇总": ("溯源", "企业采集和手工补充形成的基础资产清单。"),
            "覆盖缺口": ("复核", "各流程环节的质量门禁、缺口等级和建议动作。"),
            "风险清单": ("处置", "逐项风险、证据和处置建议。"),
            "整改矩阵": ("处置", "按优先级、时限和验收标准组织整改。"),
            "DNS记录": ("溯源", "域名、子域名、A/AAAA/CNAME/NS/MX/TXT 等解析证据。"),
            "DNS复核清单": ("复核", "无子域名、工具失败、第三方 CNAME、无公网解析等 DNS 复核项。"),
            "端口目标台账": ("复核", "候选 IP 来源、扫描覆盖、开放端口和风险数量。"),
            "开放端口": ("溯源", "主动/被动合并后的开放端口、服务、产品和来源证据。"),
            "服务识别台账": ("复核", "服务分类、HTTP 探测响应、URL 入口数量和分类依据。"),
            "URL入口覆盖": ("复核", "Web 服务是否生成 URL 入口以及入口关联质量。"),
            "非域名资产": ("溯源", "APP、小程序、公众号、邮箱等非域名类资产。"),
            "交付审计文件": ("交付", "说明交付包中的审计 JSON、复核模板和质量摘要。"),
            "重点Web资产": ("先看", "按风险、登录/管理特征和视觉识别质量排序的重点 Web 入口。"),
            "截图证据": ("证据", "重点 Web 页面截图缩略图、截图路径、识别结论和复核建议。"),
            "Web资产详情": ("溯源", "每个 URL 的标题、系统名、用途、截图、识别方式和错误信息。"),
            "视觉复核清单": ("复核", "截图失败、HTTP 降级识别和低置信度页面的复核建议。"),
        }
        rows = [
            {
                "工作表": name,
                "阅读顺序": descriptions.get(name, ("参考", ""))[0],
                "记录数": len(items),
                "用途": descriptions.get(name, ("参考", "资产测绘数据明细。"))[1],
                "建议动作": self._navigation_action(name, workbook_kind),
            }
            for name, items in sheets.items()
        ]
        return rows

    def _navigation_action(self, sheet_name: str, workbook_kind: str) -> str:
        if sheet_name in {"报告概览", "管理驾驶舱", "风险统计", "重点资产视图", "重点Web资产", "截图证据"}:
            return "建议交付评审时优先阅读。"
        if "复核" in sheet_name or sheet_name in {"覆盖缺口", "单位覆盖台账", "端口目标台账", "服务识别台账", "URL入口覆盖"}:
            return "用于补充核验、复测和下一轮整改闭环。"
        if sheet_name in {"风险清单", "整改矩阵"}:
            return "用于派单、整改跟踪和验收。"
        return "作为证据明细留档。"

    def _audit_file_rows(self, task_id: int) -> list[dict]:
        rows = [
            ("quality_summary.txt", "质量门禁摘要", "记录质量状态、覆盖缺口、建议下一步动作。"),
            (f"task_{task_id}_待补充资产模板.yaml", "人工补充模板", "供复核人员优先补充高/中优先级单位的根域名、子域名、IP、URL、APP、小程序、公众号、邮箱等资产。"),
            (f"task_{task_id}_复核工作单.yaml", "复核工作单", "列出 DNS、服务、URL、视觉识别等复核对象和续跑命令。"),
            (f"task_{task_id}_补全计划.txt", "补全计划", "把质量告警转换为按优先级排列的下一轮自动/人工补全动作。"),
            (f"task_{task_id}_补全计划.json", "补全计划数据", "补全计划的机器可读版本，便于后续系统自动续跑或审计。"),
            (f"task_{task_id}_端口目标来源.json", "端口目标来源审计", "说明每个候选 IP 来自 AI、手工补充还是 DNS 公网解析。"),
            (f"task_{task_id}_FOFA失败记录.json", "FOFA 失败记录", "仅在 FOFA 查询发生错误时生成，记录失败 IP 和错误原因。"),
            (f"task_{task_id}_HTTP探测审计.json", "HTTP 探测审计", "记录 HTTP 探测响应/失败、状态码分布和失败样例。"),
            (f"task_{task_id}_服务分类审计.json", "服务分类审计", "记录 Web/non_web 分类数量和疑似 Web 但未响应的复核候选。"),
            (f"task_{task_id}_视觉识别审计.json", "视觉识别审计", "记录截图 AI、重复复用、HTTP 降级、失败和低置信度样例。"),
            (f"task_{task_id}_报告AI分析审计.json", "报告 AI 分析审计", "记录 DNS、端口、Web、总体结论四个 AI 分块的状态、模型和提示规模。"),
            ("manifest.json", "交付清单", "记录交付包内文件大小和 SHA256，用于归档校验。"),
        ]
        return [
            {
                "文件名": filename,
                "类型": kind,
                "用途": purpose,
                "建议动作": "交付归档时保留；复核或二次交付时优先查看。",
            }
            for filename, kind, purpose in rows
        ]

    def _stats(
        self,
        bundle: dict,
        port_rows: list[dict],
        port_target_rows: list[dict],
        web_rows: list[dict],
        dns_quality_rows: list[dict],
        service_audit_rows: list[dict],
        url_coverage_rows: list[dict],
        visual_review_rows: list[dict],
        risk_rows: list[dict],
        coverage_rows: list[dict],
        unit_coverage_rows: list[dict],
    ) -> dict:
        asset_counts = Counter(row["asset_type"] for row in bundle["assets"])
        port_counts = Counter(str(row["端口"]) for row in port_rows)
        web_identified = sum(1 for row in web_rows if row.get("识别方式"))
        web_named = sum(1 for row in web_rows if row["AI识别系统"] or row["网站用途"])
        risk_counts = Counter(row["风险等级"] for row in risk_rows)
        web_methods = Counter(row.get("识别方式") or "unknown" for row in web_rows)
        web_coverage = f"{web_identified}/{len(web_rows)} ({(web_identified / len(web_rows) * 100):.1f}%)" if web_rows else "0/0 (0%)"
        no_asset_units = [row for row in unit_coverage_rows if row.get("覆盖状态") == "无资产线索"]
        priority_asset_gap_units = [row for row in no_asset_units if row.get("复核优先级") in {"高", "中"}]
        low_asset_gap_units = [row for row in no_asset_units if row.get("复核优先级") == "低"]
        return {
            "单位数量": len(bundle["companies"]),
            "股权关系数量": len(bundle["edges"]),
            "资产数量": len(bundle["assets"]),
            "DNS记录数量": len(bundle["dns_records"]),
            "DNS复核项": sum(1 for row in dns_quality_rows if row.get("复核优先级") not in {"", "无"}),
            "端口目标数量": len(port_target_rows),
            "开放端口数量": len(port_rows),
            "服务复核项": sum(1 for row in service_audit_rows if row.get("复核优先级") not in {"", "无"}),
            "Web入口数量": len(web_rows),
            "URL入口复核项": sum(1 for row in url_coverage_rows if row.get("复核优先级") not in {"", "无"}),
            "已完成视觉识别Web入口": web_identified,
            "有系统名称或用途Web入口": web_named,
            "Web识别覆盖率": web_coverage,
            "Web识别方式分布": dict(sorted(web_methods.items())),
            "视觉复核项": len(visual_review_rows),
            "高风险项": risk_counts.get("高", 0),
            "中风险项": risk_counts.get("中", 0),
            "低风险项": risk_counts.get("低", 0),
            "覆盖缺口项": sum(1 for row in coverage_rows if row.get("缺口等级") not in {"", "无"}),
            "有资产线索单位": sum(1 for row in unit_coverage_rows if row.get("资产数量", 0) or row.get("子域名/DNS主机数量", 0) or row.get("IP数量", 0)),
            "有Web入口单位": sum(1 for row in unit_coverage_rows if row.get("Web入口数量", 0)),
            "无资产线索单位": len(no_asset_units),
            "高中优先级无资产单位": len(priority_asset_gap_units),
            "低优先级无资产单位": len(low_asset_gap_units),
            "资产类型分布": dict(sorted(asset_counts.items())),
            "高频端口": dict(port_counts.most_common(10)),
        }

    def _run_analyses(self, task_id: int, bundle: dict, context: dict, rerun_ai: bool = False) -> dict[str, str]:
        sections = {
            "report_dns": {
                "title": "DNS与域名解析分析",
                "payload": {
                    "stats": context["stats"],
                    "coverage_gaps": context["coverage_rows"],
                    "dns_records": context["dns_rows"][: self.config.ai.max_dns_records],
                },
            },
            "report_ports": {
                "title": "端口与服务暴露分析",
                "payload": {
                    "stats": context["stats"],
                    "coverage_gaps": context["coverage_rows"],
                    "open_ports": context["port_rows"][:500],
                },
            },
            "report_web": {
                "title": "Web资产视觉识别分析",
                "payload": {
                    "stats": context["stats"],
                    "coverage_gaps": context["coverage_rows"],
                    "web_assets": context["web_rows"][:300],
                },
            },
        }
        results: dict[str, str] = {}
        for analysis_type, spec in sections.items():
            results[analysis_type] = self._analysis(task_id, analysis_type, spec["title"], spec["payload"], rerun_ai)
        results["report_summary"] = self._analysis(
            task_id,
            "report_summary",
            "总体暴露面结论与处置建议",
            {"stats": context["stats"], "coverage_gaps": context["coverage_rows"], "section_summaries": results, "target": bundle["task"]["target"]},
            rerun_ai,
        )
        return results

    def _analysis(self, task_id: int, analysis_type: str, title: str, payload: dict, rerun_ai: bool) -> str:
        prompt = self._analysis_prompt(title, payload)
        cache_key = self._analysis_cache_key(prompt)
        existing = self.session.exec(
            select(AiAnalysis).where(AiAnalysis.scan_task_id == task_id, AiAnalysis.analysis_type == analysis_type)
        ).first()
        existing_key = self._stored_analysis_cache_key(existing.prompt_json if existing else None)
        if existing and existing.status == "completed" and existing.summary and existing_key == cache_key and not rerun_ai:
            self._log(f"[report] skip cached AI analysis: {analysis_type}")
            return existing.summary
        if existing and existing.status == "completed" and existing.summary and existing_key != cache_key and not rerun_ai:
            self._log(f"[report] refresh stale AI analysis: {analysis_type}")
        prompt["_cache_key"] = cache_key
        if not self.config.ai.enabled:
            summary = self._fallback_analysis(title, payload)
            self._save_analysis(task_id, analysis_type, "skipped", prompt, {"reason": "ai disabled"}, summary)
            return summary
        messages = [
            {
                "role": "system",
                "content": "你是资深互联网资产暴露面分析师。基于给定结构化数据写安全测绘报告分析，不要编造不存在的资产。",
            },
            {"role": "user", "content": json.dumps(prompt, ensure_ascii=False)[: self.config.ai.max_prompt_chars]},
        ]
        try:
            response = chat_completion(self.config.ai, messages, temperature=0.1, max_completion_tokens=1800)
            summary = response.get("choices", [{}])[0].get("message", {}).get("content") or ""
            self._save_analysis(task_id, analysis_type, "completed", prompt, response, summary)
            self._log(f"[report] AI analysis completed: {analysis_type}")
            return summary
        except Exception as exc:
            summary = self._fallback_analysis(title, payload, error=str(exc))
            self._save_analysis(task_id, analysis_type, "failed", prompt, {"error": str(exc)[:1000]}, summary)
            self._log(f"[report] AI analysis failed: {analysis_type} -> {str(exc)[:200]}")
            return summary

    def _analysis_prompt(self, title: str, payload: dict) -> dict:
        return {
            "cache_version": REPORT_ANALYSIS_CACHE_VERSION,
            "analysis_title": title,
            "requirements": (
                "请用中文输出面向安全报告的分析。要求：结论明确、内容丰富但不要虚构；"
                "指出值得关注的资产、潜在风险、可验证证据和优先整改建议。"
            ),
            "data": payload,
        }

    def _stored_analysis_cache_key(self, prompt: dict | None) -> str:
        if not isinstance(prompt, dict):
            return ""
        key = prompt.get("_cache_key")
        return str(key) if key else self._analysis_cache_key(prompt)

    def _analysis_cache_key(self, prompt: dict) -> str:
        clean_prompt = {key: value for key, value in prompt.items() if key != "_cache_key"}
        encoded = json.dumps(clean_prompt, ensure_ascii=False, sort_keys=True, default=str)
        return hashlib.sha256(encoded.encode("utf-8")).hexdigest()

    def _fallback_analysis(self, title: str, payload: dict, error: str | None = None) -> str:
        stats = payload.get("stats") or {}
        lines = [f"{title}："]
        if error:
            lines.append(f"AI 分析未完成，已使用本地统计生成摘要。错误：{error[:300]}")
        lines.append(
            "当前任务共覆盖 "
            f"{stats.get('单位数量', 0)} 家单位、{stats.get('DNS记录数量', 0)} 条 DNS 记录、"
            f"{stats.get('开放端口数量', 0)} 个开放端口、{stats.get('Web入口数量', 0)} 个 Web 入口。"
        )
        if stats.get("高频端口"):
            lines.append(f"高频开放端口包括：{json.dumps(stats['高频端口'], ensure_ascii=False)}。")
        return "\n".join(lines)

    def _save_analysis(self, task_id: int, analysis_type: str, status: str, prompt: dict, response: dict, summary: str) -> None:
        row = self.session.exec(
            select(AiAnalysis).where(AiAnalysis.scan_task_id == task_id, AiAnalysis.analysis_type == analysis_type)
        ).first()
        if not row:
            row = AiAnalysis(scan_task_id=task_id, analysis_type=analysis_type)
        row.status = status
        row.model = self.config.ai.model if self.config.ai.enabled else None
        row.prompt_json = prompt
        row.response_json = response
        row.summary = summary
        row.updated_at = _utcnow()
        self.session.add(row)
        self.session.commit()

    def _write_report_ai_audit(self, task_id: int) -> Path:
        payload = self._report_ai_audit_payload(task_id)
        output_dir = Path("data") / "report" / f"task_{task_id}"
        output_dir.mkdir(parents=True, exist_ok=True)
        path = output_dir / "report_ai_audit.json"
        path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        return path

    def _report_ai_audit_payload(self, task_id: int) -> dict:
        rows = {
            row.analysis_type: row
            for row in self.session.exec(
                select(AiAnalysis).where(AiAnalysis.scan_task_id == task_id)
            ).all()
            if row.analysis_type in REPORT_ANALYSIS_TITLES
        }
        sections = []
        status_counts: dict[str, int] = {}
        for analysis_type, title in REPORT_ANALYSIS_TITLES.items():
            row = rows.get(analysis_type)
            if not row:
                status = "missing"
                status_counts[status] = status_counts.get(status, 0) + 1
                sections.append(
                    {
                        "analysis_type": analysis_type,
                        "title": title,
                        "status": status,
                        "model": None,
                        "mode": "missing",
                        "updated_at": "",
                        "summary_chars": 0,
                        "prompt_chars": 0,
                        "input_fingerprint": "",
                        "data_shape": {},
                        "response_keys": [],
                        "response_id": "",
                        "response_model": "",
                        "usage_keys": [],
                        "prompt_tokens": "",
                        "completion_tokens": "",
                        "total_tokens": "",
                        "error": "",
                        "reason": "",
                    }
                )
                continue
            status_counts[row.status] = status_counts.get(row.status, 0) + 1
            response = row.response_json or {}
            prompt = row.prompt_json or {}
            usage = response.get("usage") if isinstance(response, dict) and isinstance(response.get("usage"), dict) else {}
            prompt_tokens = usage.get("prompt_tokens", usage.get("input_tokens", ""))
            completion_tokens = usage.get("completion_tokens", usage.get("output_tokens", ""))
            sections.append(
                {
                    "analysis_type": analysis_type,
                    "title": title,
                    "status": row.status,
                    "model": row.model,
                    "mode": self._analysis_mode(row),
                    "updated_at": row.updated_at.isoformat() if row.updated_at else "",
                    "summary_chars": len(row.summary or ""),
                    "prompt_chars": len(json.dumps(prompt, ensure_ascii=False)),
                    "input_fingerprint": prompt.get("_cache_key", "") if isinstance(prompt, dict) else "",
                    "data_shape": self._analysis_data_shape(prompt),
                    "response_keys": sorted(response.keys()) if isinstance(response, dict) else [],
                    "response_id": response.get("id", "") if isinstance(response, dict) else "",
                    "response_model": response.get("model", "") if isinstance(response, dict) else "",
                    "usage_keys": sorted(usage.keys()),
                    "prompt_tokens": prompt_tokens,
                    "completion_tokens": completion_tokens,
                    "total_tokens": usage.get("total_tokens", ""),
                    "error": _short(response.get("error"), 800) if isinstance(response, dict) else "",
                    "reason": _short(response.get("reason"), 800) if isinstance(response, dict) else "",
                }
            )
        return {
            "scan_task_id": task_id,
            "generated_at": _utcnow().isoformat(),
            "ai_enabled": self.config.ai.enabled,
            "configured_model": self.config.ai.model,
            "max_prompt_chars": self.config.ai.max_prompt_chars,
            "section_count": len(sections),
            "status_counts": dict(sorted(status_counts.items())),
            "sections": sections,
        }

    def _report_ai_audit_rows(self, task_id: int) -> list[dict]:
        payload = self._report_ai_audit_payload(task_id)
        rows = []
        for section in payload.get("sections", []):
            shape = section.get("data_shape") if isinstance(section.get("data_shape"), dict) else {}
            rows.append(
                {
                    "分析分块": section.get("analysis_type") or "",
                    "标题": section.get("title") or "",
                    "状态": section.get("status") or "",
                    "模式": section.get("mode") or "",
                    "模型": section.get("model") or "",
                    "更新时间": section.get("updated_at") or "",
                    "摘要字数": section.get("summary_chars") or 0,
                    "提示字数": section.get("prompt_chars") or 0,
                    "输入指纹": section.get("input_fingerprint") or "",
                    "输入规模": json.dumps(shape, ensure_ascii=False),
                    "响应字段": ", ".join(section.get("response_keys") or []),
                    "响应ID": section.get("response_id") or "",
                    "响应模型": section.get("response_model") or "",
                    "Usage字段": ", ".join(section.get("usage_keys") or []),
                    "提示Token": section.get("prompt_tokens") or "",
                    "完成Token": section.get("completion_tokens") or "",
                    "总Token": section.get("total_tokens") or "",
                    "错误": section.get("error") or "",
                    "说明": section.get("reason") or "",
                }
            )
        return rows

    def _analysis_mode(self, row: AiAnalysis) -> str:
        if row.status == "completed":
            return "model"
        if row.status == "skipped":
            return "local_fallback_ai_disabled"
        if row.status == "failed":
            return "local_fallback_after_ai_error"
        return row.status

    def _analysis_data_shape(self, prompt: dict) -> dict[str, int | str]:
        payload = prompt.get("data") if isinstance(prompt, dict) and isinstance(prompt.get("data"), dict) else prompt
        if not isinstance(payload, dict):
            return {}
        shape: dict[str, int | str] = {}
        for key, value in payload.items():
            if isinstance(value, list):
                shape[key] = len(value)
            elif isinstance(value, dict):
                shape[key] = len(value)
            else:
                shape[key] = type(value).__name__
        return shape

    def _write_asset_workbook(self, path: Path, context: dict) -> Path:
        sheets = {
            "管理驾驶舱": context["dashboard_rows"],
            "报告概览": context["overview_rows"],
            "AI分析审计": context.get("ai_audit_rows", []),
            "风险统计": context["risk_summary_rows"],
            "重点资产视图": context["key_asset_rows"],
            "单位覆盖台账": context["unit_coverage_rows"],
            "资产汇总": context["unit_asset_rows"],
            "覆盖缺口": context["coverage_rows"],
            "风险清单": context["risk_rows"],
            "整改矩阵": context["remediation_rows"],
            "DNS记录": context["dns_rows"],
            "DNS复核清单": context["dns_quality_rows"],
            "端口目标台账": context["port_target_rows"],
            "开放端口": context["port_rows"],
            "服务识别台账": context["service_audit_rows"],
            "URL入口覆盖": context["url_coverage_rows"],
            "交付审计文件": context["audit_file_rows"],
            "非域名资产": context["named_assets"],
        }
        return self._write_workbook(path, {"阅读导航": self._navigation_rows("asset", sheets), **sheets})

    def _write_web_workbook(self, path: Path, context: dict) -> Path:
        sheets = {
            "重点Web资产": context["key_web_rows"],
            "截图证据": context["screenshot_evidence_rows"],
            "Web资产详情": context["web_rows"],
            "视觉复核清单": context["visual_review_rows"],
        }
        return self._write_workbook(path, {"阅读导航": self._navigation_rows("web", sheets), **sheets})

    def _write_workbook(self, path: Path, sheets: dict[str, list[dict]]) -> Path:
        wb = Workbook()
        wb.remove(wb.active)
        tab_colors = {
            "阅读导航": "17365D",
            "管理驾驶舱": "17365D",
            "报告概览": HEADER_FILL,
            "AI分析审计": "8064A2",
            "风险统计": "C00000",
            "重点资产视图": "C00000",
            "资产汇总": "5B9BD5",
            "覆盖缺口": "ED7D31",
            "风险清单": "C00000",
            "整改矩阵": "70AD47",
            "DNS记录": "4472C4",
            "DNS复核清单": "ED7D31",
            "端口目标台账": "8064A2",
            "开放端口": "7030A0",
            "服务识别台账": "5B9BD5",
            "URL入口覆盖": "70AD47",
            "非域名资产": "A5A5A5",
            "重点Web资产": "C00000",
            "截图证据": "70AD47",
            "Web资产详情": "4472C4",
            "视觉复核清单": "ED7D31",
        }
        for sheet_name, rows in sheets.items():
            ws = wb.create_sheet(sheet_name[:31])
            ws.sheet_view.showGridLines = False
            ws.sheet_properties.tabColor = tab_colors.get(sheet_name, HEADER_FILL)
            if not rows:
                ws.append(["无数据"])
                continue
            headers = self._workbook_headers(rows)
            ws.append(headers)
            for row in rows:
                ws.append([row.get(header, "") for header in headers])
            ws.freeze_panes = "A2"
            ws.auto_filter.ref = ws.dimensions
            self._style_worksheet(ws, headers, sheet_name)
            self._add_sheet_links(ws, headers)
            if sheet_name == "管理驾驶舱":
                self._add_dashboard_charts(ws, rows)
            if sheet_name == "截图证据":
                self._embed_screenshot_thumbnails(ws, rows, headers)
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            wb.save(path)
            return path
        except PermissionError:
            fallback = _timestamped_path(path)
            wb.save(fallback)
            self._log(f"[report] target workbook is locked, wrote fallback -> {fallback}")
            return fallback

    def _workbook_headers(self, rows: list[dict]) -> list[str]:
        headers: list[str] = []
        for row in rows:
            for key in row.keys():
                if key not in headers:
                    headers.append(key)
        return headers

    def _embed_screenshot_thumbnails(self, ws, rows: list[dict], headers: list[str]) -> None:
        if "缩略图" not in headers or "截图文件" not in headers:
            return
        thumb_column = headers.index("缩略图") + 1
        screenshot_column = headers.index("截图文件") + 1
        ws.column_dimensions[get_column_letter(thumb_column)].width = 34
        embedded = 0
        for row_index, row in enumerate(rows, start=2):
            if embedded >= MAX_EXCEL_SCREENSHOT_THUMBNAILS:
                break
            path_text = _text(row.get("截图文件"))
            if not path_text:
                continue
            path = Path(path_text)
            if not path.is_absolute():
                path = Path.cwd() / path
            if not path.exists():
                continue
            try:
                image = XLImage(str(path))
                image.width = 220
                image.height = 124
                ws.add_image(image, f"{get_column_letter(thumb_column)}{row_index}")
                ws.row_dimensions[row_index].height = 98
                ws.cell(row=row_index, column=screenshot_column).hyperlink = str(path)
                ws.cell(row=row_index, column=screenshot_column).style = "Hyperlink"
                embedded += 1
            except Exception:
                ws.cell(row=row_index, column=thumb_column).value = "缩略图嵌入失败"

    def _add_dashboard_charts(self, ws, rows: list[dict]) -> None:
        groups = {
            "风险等级": "风险等级分布",
            "资产类型": "资产类型分布",
            "Web识别方式": "Web识别方式分布",
            "单位覆盖": "单位覆盖情况",
        }
        anchors = {
            "风险等级": "H2",
            "资产类型": "H18",
            "Web识别方式": "P2",
            "单位覆盖": "P18",
        }
        start_row = 2
        for group, title in groups.items():
            items = [(row.get("指标"), row.get("数值")) for row in rows if row.get("分组") == group and row.get("数值") not in {"", None}]
            if not items:
                continue
            ws.cell(row=start_row, column=6).value = title
            ws.cell(row=start_row, column=6).font = Font(name="Microsoft YaHei", bold=True, color="1F4E78")
            ws.cell(row=start_row + 1, column=6).value = "指标"
            ws.cell(row=start_row + 1, column=7).value = "数值"
            for offset, (label, value) in enumerate(items, start=2):
                ws.cell(row=start_row + offset, column=6).value = label
                ws.cell(row=start_row + offset, column=7).value = int(value or 0)
            end_row = start_row + len(items) + 1
            data = Reference(ws, min_col=7, min_row=start_row + 1, max_row=end_row)
            categories = Reference(ws, min_col=6, min_row=start_row + 2, max_row=end_row)
            chart = PieChart() if group in {"风险等级", "单位覆盖"} else BarChart()
            chart.title = title
            chart.add_data(data, titles_from_data=True)
            chart.set_categories(categories)
            chart.height = 7
            chart.width = 12
            if isinstance(chart, BarChart):
                chart.y_axis.title = "数量"
                chart.x_axis.title = group
            ws.add_chart(chart, anchors[group])
            start_row = end_row + 3

    def _add_sheet_links(self, ws, headers: list[str]) -> None:
        if "工作表" not in headers:
            return
        column = headers.index("工作表") + 1
        for row_index in range(2, ws.max_row + 1):
            cell = ws.cell(row=row_index, column=column)
            if not cell.value:
                continue
            cell.hyperlink = f"#'{cell.value}'!A1"
            cell.style = "Hyperlink"

    def _style_worksheet(self, ws, headers: list[str], sheet_name: str) -> None:
        thin = Side(style="thin", color=BORDER_COLOR)
        for cell in ws[1]:
            cell.font = Font(name="Microsoft YaHei", bold=True, color="FFFFFF")
            cell.fill = PatternFill("solid", fgColor=HEADER_FILL)
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = Border(top=thin, right=thin, bottom=thin, left=thin)
        ws.row_dimensions[1].height = 24
        severity_column = self._header_index(headers, "风险等级")
        score_column = self._header_index(headers, "风险分值")
        gap_column = self._header_index(headers, "缺口等级")
        priority_column = self._header_index(headers, "优先级")
        for row_index in range(2, ws.max_row + 1):
            fill = self._row_fill(ws, row_index, severity_column, score_column, gap_column, priority_column)
            for cell in ws[row_index]:
                cell.font = Font(name="Microsoft YaHei", color="000000")
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = Border(top=thin, right=thin, bottom=thin, left=thin)
                if fill:
                    cell.fill = PatternFill("solid", fgColor=fill)
                elif row_index % 2 == 0 and sheet_name in {"报告概览", "覆盖缺口", "风险统计"}:
                    cell.fill = PatternFill("solid", fgColor="F7FBFE")
        for column_index, header in enumerate(headers, start=1):
            width = min(max(len(header) + 4, 12), 48)
            for cell in ws.iter_rows(min_col=column_index, max_col=column_index, min_row=2):
                width = min(max(width, len(_text(cell[0].value)) + 2), 60)
            ws.column_dimensions[get_column_letter(column_index)].width = width

    def _header_index(self, headers: list[str], name: str) -> int | None:
        return headers.index(name) + 1 if name in headers else None

    def _row_fill(self, ws, row_index: int, *columns: int | None) -> str | None:
        severity = ""
        for column in columns:
            if column:
                severity = _text(ws.cell(row=row_index, column=column).value)
                if severity:
                    break
        if severity.isdigit():
            score = int(severity)
            if score >= 90:
                return HIGH_FILL
            if score >= 60:
                return MEDIUM_FILL
            if score > 0:
                return LOW_FILL
        if severity == "高":
            return HIGH_FILL
        if severity == "中":
            return MEDIUM_FILL
        if severity == "低":
            return LOW_FILL
        return None

    def _write_docx(
        self,
        path: Path,
        task: ScanTask,
        context: dict,
        analyses: dict[str, str],
        asset_workbook: Path,
        web_workbook: Path,
    ) -> Path:
        doc = Document()
        self._apply_doc_style(doc, task)
        title = doc.add_heading("互联网数字资产暴露面测绘报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle = doc.add_paragraph(task.target)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        doc.add_paragraph(f"任务编号：{task.id}")
        doc.add_paragraph(f"生成时间：{_utcnow().strftime('%Y-%m-%d %H:%M:%S')}")
        self._add_cover_meta_panel(doc, task, context)
        self._add_cover_metrics(doc, context["stats"])
        self._add_report_directory(doc)

        doc.add_heading("一、报告摘要", level=1)
        self._add_paragraphs(doc, analyses.get("report_summary", ""))
        self._add_executive_findings(doc, context)
        self._add_risk_overview(doc, context)
        doc.add_heading("二、测绘范围与方法", level=1)
        self._add_methodology(doc, context["stats"])
        doc.add_heading("三、资产统计", level=1)
        self._add_stats_table(doc, context["stats"])
        self._add_unit_focus_table(doc, context["unit_coverage_rows"][:12])
        doc.add_heading("四、覆盖缺口分析", level=1)
        self._add_coverage_table(doc, context["coverage_rows"])
        doc.add_heading("五、复核与质量门禁计划", level=1)
        self._add_quality_gate_summary(doc, context)
        self._add_review_action_table(doc, context)
        doc.add_heading("六、重点风险清单", level=1)
        self._add_risk_table(doc, context["risk_rows"][:20])
        doc.add_heading("七、整改优先级矩阵", level=1)
        self._add_remediation_summary(doc, context["remediation_rows"])
        self._add_remediation_table(doc, context["remediation_rows"][:20])
        doc.add_heading("八、DNS 与域名解析分析", level=1)
        self._add_paragraphs(doc, analyses.get("report_dns", ""))
        doc.add_heading("九、端口与服务暴露分析", level=1)
        self._add_paragraphs(doc, analyses.get("report_ports", ""))
        doc.add_heading("十、Web 资产视觉识别分析", level=1)
        self._add_paragraphs(doc, analyses.get("report_web", ""))
        self._add_web_table(doc, self._top_web_rows(context["web_rows"], context["risk_rows"])[:20])
        self._add_screenshot_gallery(doc, context["web_rows"])
        doc.add_heading("十一、附件", level=1)
        doc.add_paragraph(
            f"附件 1：{asset_workbook.name}，包含报告概览、单位覆盖台账、覆盖缺口、风险清单、整改矩阵、"
            "管理驾驶舱、DNS记录、DNS复核清单、端口目标台账、开放端口、服务识别台账、URL入口覆盖、交付审计文件和非域名资产。"
        )
        doc.add_paragraph(
            f"附件 2：{web_workbook.name}，包含重点 Web 资产、截图证据、AI 视觉识别结论、截图路径和视觉复核清单。"
        )
        doc.add_paragraph(
            "交付包同时包含质量摘要、待补充资产模板、复核工作单、端口目标来源、HTTP 探测审计、服务分类审计、"
            "视觉识别审计、报告 AI 分析审计和 manifest 校验清单；具体用途见资产汇总附件中的“交付审计文件”工作表。"
        )
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        path.parent.mkdir(parents=True, exist_ok=True)
        try:
            doc.save(path)
            return path
        except PermissionError:
            fallback = _timestamped_path(path)
            doc.save(fallback)
            self._log(f"[report] target Word report is locked, wrote fallback -> {fallback}")
            return fallback

    def _apply_doc_style(self, doc: Document, task: ScanTask) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = "Microsoft YaHei"
        normal.font.size = Pt(10.5)
        for style_name, size, color in (
            ("Title", 22, RGBColor(31, 78, 121)),
            ("Heading 1", 15, RGBColor(31, 78, 121)),
            ("Heading 2", 12, RGBColor(79, 129, 189)),
        ):
            style = doc.styles[style_name]
            style.font.name = "Microsoft YaHei"
            style.font.size = Pt(size)
            style.font.color.rgb = color
            style.font.bold = True
        section = doc.sections[0]
        header = section.header.paragraphs[0]
        header.text = f"互联网数字资产暴露面测绘报告 | {task.target}"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(127, 127, 127)
        footer = section.footer.paragraphs[0]
        footer.text = "assetmap 自动生成 | 仅用于授权资产测绘与安全治理"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(127, 127, 127)

    def _add_cover_meta_panel(self, doc: Document, task: ScanTask, context: dict) -> None:
        label = doc.add_paragraph()
        label.add_run("报告信息").bold = True
        rows = [
            ("报告对象", task.target),
            ("任务编号", str(task.id)),
            ("报告状态", self._report_status(context["coverage_rows"])),
            (
                "交付范围",
                f"{context['stats'].get('单位数量', 0)} 家单位 / "
                f"{context['stats'].get('端口目标数量', 0)} 个端口目标 / "
                f"{context['stats'].get('Web入口数量', 0)} 个 Web 入口",
            ),
        ]
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "报告信息"
        table.rows[0].cells[1].text = "内容"
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].text = label
            cells[1].text = value
        self._format_table(table, header_fill="17365D")

    def _report_status(self, coverage_rows: list[dict]) -> str:
        if any(row.get("缺口等级") == "高" for row in coverage_rows):
            return "存在高优先级缺口，需补扫后交付"
        if any(row.get("缺口等级") == "中" for row in coverage_rows):
            return "可交付，存在中优先级复核项"
        if any(row.get("缺口等级") == "低" for row in coverage_rows):
            return "可交付，存在低优先级复核项"
        return "可交付，质量门禁通过"

    def _add_cover_metrics(self, doc: Document, stats: dict) -> None:
        rows = [
            ("单位", stats.get("单位数量", 0)),
            ("基础资产", stats.get("资产数量", 0)),
            ("端口目标", stats.get("端口目标数量", 0)),
            ("开放端口", stats.get("开放端口数量", 0)),
            ("Web入口", stats.get("Web入口数量", 0)),
            ("高风险", stats.get("高风险项", 0)),
            ("覆盖缺口", stats.get("覆盖缺口项", 0)),
        ]
        table = doc.add_table(rows=2, cols=len(rows))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for index, (label, value) in enumerate(rows):
            table.rows[0].cells[index].text = label
            table.rows[1].cells[index].text = _text(value)
        self._format_table(table, header_fill=HEADER_FILL, align_center=True)

    def _add_report_directory(self, doc: Document) -> None:
        doc.add_heading("报告目录", level=1)
        rows = [
            ("一、报告摘要", "总体结论、关键发现、风险概览", "先看"),
            ("二、测绘范围与方法", "数据来源、测绘链路、风险判断口径", "先看"),
            ("三、资产统计", "单位、资产、端口、Web入口和覆盖矩阵", "附件1"),
            ("四、覆盖缺口分析", "资产、DNS、端口、服务、URL和视觉识别缺口", "附件1"),
            ("五、复核与质量门禁计划", "交付后复核动作和下一轮补全路径", "工作单"),
            ("六、重点风险清单", "高/中风险暴露面和证据摘要", "附件1"),
            ("七、整改优先级矩阵", "处置节奏、责任建议和验收证据", "附件1"),
            ("八、DNS 与域名解析分析", "根域名、解析质量、CNAME/共享IP线索", "附件1"),
            ("九、端口与服务暴露分析", "开放端口、主动/被动证据和服务指纹", "附件1"),
            ("十、Web 资产视觉识别分析", "系统名称、用途、截图证据和复核清单", "附件2"),
            ("十一、附件", "Excel附件、质量摘要、审计文件和交付清单", "交付包"),
        ]
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = ["章节", "阅读重点", "对应材料"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for section, focus, material in rows:
            cells = table.add_row().cells
            cells[0].text = section
            cells[1].text = focus
            cells[2].text = material
        self._format_table(table)

    def _format_table(self, table, header_fill: str = HEADER_FILL, align_center: bool = False) -> None:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    if align_center or row_index == 0:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Microsoft YaHei"
                        run.font.size = Pt(9)
                if row_index == 0:
                    self._set_cell_shading(cell, header_fill)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                elif row_index % 2 == 0:
                    self._set_cell_shading(cell, "F7FBFE")

    def _set_cell_shading(self, cell, fill: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), fill)

    def _add_paragraphs(self, doc: Document, text: str) -> None:
        lines = (text or "无分析内容。").splitlines()
        index = 0
        while index < len(lines):
            line = lines[index]
            value = line.strip()
            if not value:
                index += 1
                continue
            table_rows, next_index = self._consume_markdown_table(lines, index)
            if table_rows:
                self._add_markdown_table(doc, table_rows)
                index = next_index
                continue
            heading = re.match(r"^(#{1,4})\s+(.+)$", value)
            if heading:
                doc.add_heading(self._plain_markdown(heading.group(2)), level=2)
                index += 1
                continue
            bullet = re.match(r"^[-*]\s+(.+)$", value)
            if bullet:
                doc.add_paragraph(self._plain_markdown(bullet.group(1)), style="List Bullet")
                index += 1
                continue
            ordered = re.match(r"^\d+[.)]\s+(.+)$", value)
            if ordered:
                doc.add_paragraph(self._plain_markdown(ordered.group(1)), style="List Number")
                index += 1
                continue
            doc.add_paragraph(self._plain_markdown(value))
            index += 1

    def _consume_markdown_table(self, lines: list[str], start: int) -> tuple[list[list[str]], int]:
        if start + 1 >= len(lines):
            return [], start
        first = lines[start].strip()
        second = lines[start + 1].strip()
        if not self._is_markdown_table_row(first) or not self._is_markdown_separator_row(second):
            return [], start
        rows = [self._split_markdown_table_row(first)]
        index = start + 2
        while index < len(lines) and self._is_markdown_table_row(lines[index].strip()):
            rows.append(self._split_markdown_table_row(lines[index].strip()))
            index += 1
        column_count = max((len(row) for row in rows), default=0)
        normalized = [
            [self._plain_markdown(cell) for cell in row + [""] * (column_count - len(row))]
            for row in rows
        ]
        return normalized, index

    def _is_markdown_table_row(self, value: str) -> bool:
        return value.startswith("|") and value.count("|") >= 2

    def _is_markdown_separator_row(self, value: str) -> bool:
        if not self._is_markdown_table_row(value):
            return False
        cells = self._split_markdown_table_row(value)
        return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells)

    def _split_markdown_table_row(self, value: str) -> list[str]:
        text = value.strip()
        if text.startswith("|"):
            text = text[1:]
        if text.endswith("|"):
            text = text[:-1]
        return [cell.strip() for cell in text.split("|")]

    def _add_markdown_table(self, doc: Document, rows: list[list[str]]) -> None:
        if not rows:
            return
        table = doc.add_table(rows=1, cols=len(rows[0]))
        table.style = "Table Grid"
        for index, cell_value in enumerate(rows[0]):
            table.rows[0].cells[index].text = cell_value
        for row in rows[1:]:
            cells = table.add_row().cells
            for index, cell_value in enumerate(row):
                cells[index].text = cell_value
        self._format_table(table)

    def _plain_markdown(self, value: str) -> str:
        text = value.strip()
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"__([^_]+)__", r"\1", text)
        text = text.replace("`", "")
        return text

    def _add_stats_table(self, doc: Document, stats: dict) -> None:
        rows = [
            ("单位数量", stats.get("单位数量", 0)),
            ("有资产线索单位", stats.get("有资产线索单位", 0)),
            ("有Web入口单位", stats.get("有Web入口单位", 0)),
            ("股权关系数量", stats.get("股权关系数量", 0)),
            ("资产数量", stats.get("资产数量", 0)),
            ("DNS记录数量", stats.get("DNS记录数量", 0)),
            ("端口目标数量", stats.get("端口目标数量", 0)),
            ("开放端口数量", stats.get("开放端口数量", 0)),
            ("服务复核项", stats.get("服务复核项", 0)),
            ("Web入口数量", stats.get("Web入口数量", 0)),
            ("URL入口复核项", stats.get("URL入口复核项", 0)),
            ("已完成视觉识别Web入口", stats.get("已完成视觉识别Web入口", 0)),
            ("有系统名称或用途Web入口", stats.get("有系统名称或用途Web入口", 0)),
            ("Web识别覆盖率", stats.get("Web识别覆盖率", "0%")),
            ("视觉复核项", stats.get("视觉复核项", 0)),
            ("覆盖缺口项", stats.get("覆盖缺口项", 0)),
            ("高风险项", stats.get("高风险项", 0)),
            ("中风险项", stats.get("中风险项", 0)),
            ("低风险项", stats.get("低风险项", 0)),
            ("资产类型分布", json.dumps(stats.get("资产类型分布", {}), ensure_ascii=False)),
            ("Web识别方式分布", json.dumps(stats.get("Web识别方式分布", {}), ensure_ascii=False)),
            ("高频端口", json.dumps(stats.get("高频端口", {}), ensure_ascii=False)),
        ]
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        table.rows[0].cells[0].text = "指标"
        table.rows[0].cells[1].text = "数值"
        for name, value in rows:
            cells = table.add_row().cells
            cells[0].text = _text(name)
            cells[1].text = _text(value)
        self._format_table(table)

    def _add_unit_focus_table(self, doc: Document, rows: list[dict]) -> None:
        if not rows:
            return
        doc.add_heading("重点单位覆盖矩阵", level=2)
        table = doc.add_table(rows=1, cols=8)
        table.style = "Table Grid"
        headers = ["单位", "层级", "根域名", "IP", "开放端口", "Web入口", "高/中风险", "覆盖状态"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = _short(row.get("单位"), 80)
            cells[1].text = _short(row.get("股权层级"), 20)
            cells[2].text = _text(row.get("根域名数量") or 0)
            cells[3].text = _text(row.get("IP数量") or 0)
            cells[4].text = _text(row.get("开放端口数量") or 0)
            cells[5].text = _text(row.get("Web入口数量") or 0)
            cells[6].text = f"{row.get('高风险数量', 0)}/{row.get('中风险数量', 0)}"
            cells[7].text = _short(row.get("覆盖状态"), 80)
        self._format_table(table)

    def _add_methodology(self, doc: Document, stats: dict) -> None:
        paragraphs = [
            "本报告基于企业股权/备案资产采集、手工补充资产、子域名枚举、DNS 解析、被动/主动端口发现、服务识别、Web 页面截图与多模态 AI 识别结果生成。",
            "测绘对象覆盖目标企业及其控股链路中的相关单位，资产类型包括备案域名、子域名、公网 IP、开放端口、Web 入口、APP、小程序、微信公众号/服务号与邮箱线索。",
            "风险判断采用本地规则与 AI 分块分析结合的方式：本地规则用于识别远程接入、敏感服务、管理入口等高价值暴露面；AI 用于对 DNS、端口、Web 视觉识别结果进行上下文总结。",
            f"本次数据规模：单位 {stats.get('单位数量', 0)} 家，DNS 记录 {stats.get('DNS记录数量', 0)} 条，开放端口 {stats.get('开放端口数量', 0)} 个，Web 入口 {stats.get('Web入口数量', 0)} 个。",
            "报告中的风险等级用于整改优先级排序，不替代漏洞验证结论；高风险项建议优先完成业务必要性确认、访问控制收敛、弱口令与版本漏洞核查。",
        ]
        for item in paragraphs:
            doc.add_paragraph(item)

    def _add_executive_findings(self, doc: Document, context: dict) -> None:
        doc.add_heading("关键发现", level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        headers = ["类别", "结论", "建议动作"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        stats = context["stats"]
        findings = [
            (
                "资产规模",
                f"覆盖 {stats.get('单位数量', 0)} 家单位、{stats.get('资产数量', 0)} 个基础资产、{stats.get('Web入口数量', 0)} 个 Web 入口。",
                f"其中 {stats.get('有资产线索单位', 0)} 家单位已有资产线索，{stats.get('有Web入口单位', 0)} 家单位已有 Web 入口；保持企业、域名、IP、端口和 Web 入口台账同步更新。",
            ),
            (
                "暴露风险",
                f"识别高风险 {stats.get('高风险项', 0)} 项、中风险 {stats.get('中风险项', 0)} 项。",
                "优先处理远程接入、敏感服务和管理入口暴露。",
            ),
            (
                "视觉识别",
                f"Web 识别覆盖率 {stats.get('Web识别覆盖率', '0%')}，识别方式：{json.dumps(stats.get('Web识别方式分布', {}), ensure_ascii=False)}。",
                "对降级识别页面安排人工复核，必要时补充截图或登录态验证。",
            ),
            (
                "覆盖缺口",
                f"当前仍有 {stats.get('覆盖缺口项', 0)} 类资产链路缺口需要复核。",
                "优先补齐无资产单位、无 DNS 证据域名、未覆盖端口证据 IP 和降级识别页面。",
            ),
        ]
        for category, conclusion, action in findings:
            cells = table.add_row().cells
            cells[0].text = category
            cells[1].text = conclusion
            cells[2].text = action
        self._format_table(table)

    def _add_risk_overview(self, doc: Document, context: dict) -> None:
        doc.add_heading("风险概览", level=2)
        risk_summary = context["risk_summary_rows"]
        summary_by_level: dict[str, list[str]] = {}
        count_by_level: dict[str, int] = {}
        for row in risk_summary:
            level = row.get("风险等级") or "未定级"
            count_by_level[level] = count_by_level.get(level, 0) + int(row.get("数量") or 0)
            summary_by_level.setdefault(level, []).append(f"{row.get('风险类型')}({row.get('数量')})")
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["风险等级", "数量", "主要类型", "建议处置节奏"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for level, cadence in (
            ("高", "7 天内确认并收敛"),
            ("中", "30 天内完成加固"),
            ("低", "90 天内纳入周期治理"),
        ):
            cells = table.add_row().cells
            cells[0].text = level
            cells[1].text = str(count_by_level.get(level, 0))
            cells[2].text = "、".join(summary_by_level.get(level, [])[:4]) or "无"
            cells[3].text = cadence
        self._format_table(table)

    def _add_coverage_table(self, doc: Document, rows: list[dict]) -> None:
        if not rows:
            doc.add_paragraph("未生成覆盖缺口分析。")
            return
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["环节", "指标", "结果", "缺口等级", "建议动作"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = _short(row.get("环节"), 50)
            cells[1].text = _short(row.get("指标"), 80)
            cells[2].text = _short(row.get("结果"), 100)
            cells[3].text = _short(row.get("缺口等级"), 20)
            cells[4].text = _short(row.get("建议动作"), 220)
        self._format_table(table)

    def _add_quality_gate_summary(self, doc: Document, context: dict) -> None:
        stats = context["stats"]
        asset_gap_priority = "中" if stats.get("高中优先级无资产单位", 0) else ("低" if stats.get("低优先级无资产单位", 0) else "无")
        rows = [
            (
                "企业/备案资产",
                f"{stats.get('有资产线索单位', 0)}/{stats.get('单位数量', 0)} 家单位已有资产线索；高/中优先级待补充 {stats.get('高中优先级无资产单位', 0)} 家",
                asset_gap_priority,
                "优先补充高/中优先级单位的备案、官网、公众号、小程序、APP、邮箱和公网 IP。",
            ),
            ("子域名/DNS", f"DNS复核项 {stats.get('DNS复核项', 0)} 个", "低", "优先处理工具失败、无公网解析和第三方/停放 CNAME。"),
            ("服务识别/URL", f"服务复核项 {stats.get('服务复核项', 0)} 个，URL入口复核项 {stats.get('URL入口复核项', 0)} 个", "低", "复核疑似 Web 端口、URL 关联和 Host 绑定。"),
            ("URL视觉识别", f"视觉复核项 {stats.get('视觉复核项', 0)} 个", "低", "复核降级识别、低置信度和截图失败页面。"),
        ]
        doc.add_paragraph("以下复核计划来自质量门禁和附件台账，用于指导交付后的补充核验与下一轮复测。")
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["环节", "当前状态", "优先级", "建议动作"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for stage, state, priority, action in rows:
            cells = table.add_row().cells
            cells[0].text = stage
            cells[1].text = state
            cells[2].text = priority
            cells[3].text = action
        self._format_table(table)

    def _add_review_action_table(self, doc: Document, context: dict) -> None:
        rows = [
            self._review_item("单位资产补充", context["unit_coverage_rows"], "单位", "缺口原因", "建议动作", priorities={"高", "中"}),
            self._review_item("DNS复核", context["dns_quality_rows"], "根域名", "复核原因", "建议动作"),
            self._review_item("服务识别复核", context["service_audit_rows"], "IP", "分类依据", "建议动作", port_key="端口"),
            self._review_item("URL入口复核", context["url_coverage_rows"], "URL入口样例", "覆盖结论", "建议动作"),
            self._review_item("视觉识别复核", context["visual_review_rows"], "URL", "复核原因", "建议动作"),
        ]
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["复核类型", "数量", "典型对象", "主要原因", "建议动作"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = row["复核类型"]
            cells[1].text = str(row["数量"])
            cells[2].text = _short(row["典型对象"], 160)
            cells[3].text = _short(row["主要原因"], 180)
            cells[4].text = _short(row["建议动作"], 220)
        self._format_table(table)

    def _review_item(
        self,
        label: str,
        rows: list[dict],
        object_key: str,
        reason_key: str,
        action_key: str,
        *,
        port_key: str | None = None,
        priorities: set[str] | None = None,
    ) -> dict:
        if priorities is not None:
            candidates = [row for row in rows if row.get("复核优先级") in priorities]
        else:
            candidates = [row for row in rows if row.get("复核优先级") not in {"", "无"}]
        samples = []
        for row in candidates[:5]:
            value = row.get(object_key) or ""
            if port_key and row.get(port_key):
                value = f"{value}:{row.get(port_key)}"
            samples.append(_text(value))
        return {
            "复核类型": label,
            "数量": len(candidates),
            "典型对象": ", ".join(item for item in samples if item) or "无",
            "主要原因": _text(candidates[0].get(reason_key)) if candidates else "无",
            "建议动作": _text(candidates[0].get(action_key)) if candidates else "保持周期性复测。",
        }

    def _add_risk_table(self, doc: Document, rows: list[dict]) -> None:
        if not rows:
            doc.add_paragraph("未根据当前规则识别到明确的重点风险项。")
            return
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["等级", "分值", "风险类型", "单位", "资产", "处置建议"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = _short(row.get("风险等级"), 20)
            cells[1].text = _short(row.get("风险分值"), 20)
            cells[2].text = _short(row.get("风险类型"), 80)
            cells[3].text = _short(row.get("单位"), 80)
            cells[4].text = _short(row.get("资产"), 120)
            cells[5].text = _short(row.get("处置建议"), 220)
        self._format_table(table)

    def _add_remediation_table(self, doc: Document, rows: list[dict]) -> None:
        if not rows:
            doc.add_paragraph("未生成整改项。")
            return
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"
        headers = ["优先级", "分值", "单位", "资产", "问题", "建议时限", "责任建议"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = _short(row.get("优先级"), 20)
            cells[1].text = _short(row.get("风险分值"), 20)
            cells[2].text = _short(row.get("单位"), 80)
            cells[3].text = _short(row.get("资产"), 110)
            cells[4].text = _short(row.get("问题"), 80)
            cells[5].text = _short(row.get("建议时限"), 80)
            cells[6].text = _short(row.get("责任建议"), 160)
        self._format_table(table)

    def _add_remediation_summary(self, doc: Document, rows: list[dict]) -> None:
        if not rows:
            doc.add_paragraph("当前没有进入整改矩阵的风险项。")
            return
        counts = Counter(row.get("优先级") or "未定级" for row in rows)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        headers = ["优先级", "数量", "建议时限", "验收重点"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        guidance = {
            "高": ("7 天内完成确认与收敛", "业务必要性、访问控制、强认证、补丁和审计证据"),
            "中": ("30 天内完成加固", "账号口令、版本漏洞、访问范围和日志审计"),
            "低": ("90 天内纳入周期治理", "资产归属、台账更新和周期复测"),
        }
        for level in ("高", "中", "低"):
            cells = table.add_row().cells
            cells[0].text = level
            cells[1].text = str(counts.get(level, 0))
            cells[2].text = guidance[level][0]
            cells[3].text = guidance[level][1]
        self._format_table(table)

    def _top_web_rows(self, web_rows: list[dict], risk_rows: list[dict]) -> list[dict]:
        risk_assets = {row.get("资产") for row in risk_rows if row.get("资产")}
        keywords = ("vpn", "easyconnect", "easy connect", "远程接入", "登录", "login", "管理", "admin", "后台", "oa", "portal", "门户")

        def score(row: dict) -> tuple[int, str]:
            text = " ".join(_text(row.get(key)) for key in ("URL", "AI识别系统", "HTML标题", "网站用途", "登录特征")).lower()
            value = 0
            if row.get("URL") in risk_assets:
                value += 100
            if any(keyword in text for keyword in keywords):
                value += 30
            if row.get("AI识别系统") or row.get("网站用途"):
                value += 10
            if row.get("识别方式") == "screenshot_ai":
                value += 5
            return (-value, row.get("URL") or "")

        return sorted(web_rows, key=score)

    def _screenshot_evidence_rows(self, web_rows: list[dict], risk_rows: list[dict]) -> list[dict]:
        rows = []
        for row in self._top_web_rows(web_rows, risk_rows)[:50]:
            screenshot = row.get("截图") or ""
            screenshot_exists = bool(screenshot and Path(str(screenshot)).exists())
            if screenshot_exists:
                status = "有截图"
                missing_reason = ""
            elif screenshot:
                status = "截图文件缺失"
                missing_reason = f"数据库记录了截图路径，但本地文件不存在：{screenshot}"
            elif row.get("分析错误"):
                status = "需复核"
                missing_reason = row.get("分析错误") or ""
            else:
                status = "无截图"
                missing_reason = row.get("降级原因") or "该页面使用重复页面复用或 HTTP 探测降级识别，未生成独立截图。"
            rows.append(
                {
                    "缩略图": "",
                    "单位": row.get("单位"),
                    "URL": row.get("URL"),
                    "系统/标题": row.get("AI识别系统") or row.get("HTML标题"),
                    "网站用途": row.get("网站用途"),
                    "页面类型": row.get("页面类型"),
                    "识别方式": row.get("识别方式"),
                    "识别置信度": row.get("识别置信度"),
                    "截图状态": status,
                    "截图文件": screenshot,
                    "截图缺失原因": missing_reason,
                    "复核建议": missing_reason or self._visual_review_action("低", row.get("识别方式") or "", ""),
                }
            )
        return rows

    def _add_web_table(self, doc: Document, rows: list[dict]) -> None:
        if not rows:
            return
        doc.add_heading("重点 Web 资产摘录", level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        headers = ["单位", "URL", "系统/标题", "用途", "截图/错误"]
        for index, header in enumerate(headers):
            table.rows[0].cells[index].text = header
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = _short(row.get("单位"), 80)
            cells[1].text = _short(row.get("URL"), 120)
            cells[2].text = _short(row.get("AI识别系统") or row.get("HTML标题"), 120)
            cells[3].text = _short(row.get("网站用途"), 180)
            cells[4].text = _short(row.get("截图") or row.get("分析错误"), 180)
        self._format_table(table)

    def _add_screenshot_gallery(self, doc: Document, rows: list[dict]) -> None:
        screenshots = []
        for row in rows:
            path_text = row.get("截图")
            if not path_text:
                continue
            path = Path(path_text)
            if not path.is_absolute():
                path = Path.cwd() / path
            if path.exists():
                screenshots.append((row, path))
        if not screenshots:
            return
        doc.add_heading("重点 Web 页面截图", level=2)
        for row, path in screenshots[:8]:
            caption = doc.add_paragraph()
            caption.add_run(_short(row.get("AI识别系统") or row.get("HTML标题") or "Web资产", 80)).bold = True
            caption.add_run(f"  {row.get('URL') or ''}")
            try:
                doc.add_picture(str(path), width=Cm(15.5))
            except Exception:
                doc.add_paragraph(f"截图无法嵌入：{path}")
