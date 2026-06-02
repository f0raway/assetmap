from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from sqlmodel import Session

from assetmap.config import AppConfig
from assetmap.models import ScanTask
from assetmap.services.exporter import ExportService
from assetmap.services.report import ReportService, _safe_name
from assetmap.services.status import PipelineStatusService


EXPECTED_ASSET_SHEETS = {
    "阅读导航",
    "管理驾驶舱",
    "报告概览",
    "AI分析审计",
    "风险统计",
    "重点资产视图",
    "单位覆盖台账",
    "资产汇总",
    "覆盖缺口",
    "风险清单",
    "整改矩阵",
    "DNS记录",
    "DNS复核清单",
    "端口目标台账",
    "开放端口",
    "服务识别台账",
    "URL入口覆盖",
    "交付审计文件",
    "非域名资产",
}
EXPECTED_WEB_SHEETS = {"阅读导航", "重点Web资产", "截图证据", "Web资产详情", "视觉复核清单"}


@dataclass
class QualityResult:
    status: str
    failures: list[str]
    warnings: list[str]
    lines: list[str]


class DeliveryQualityService:
    def __init__(self, session: Session, config: AppConfig) -> None:
        self.session = session
        self.config = config

    def check(self, task_id: int, output_dir: Path | str = "reports") -> QualityResult:
        task = self.session.get(ScanTask, task_id)
        if not task:
            raise ValueError(f"Task not found: {task_id}")
        failures: list[str] = []
        warnings: list[str] = []

        pipeline = PipelineStatusService(self.session).get(task_id)
        incomplete = [name for name, status, _ in pipeline.stages if status != "completed"]
        if incomplete:
            failures.append(f"流程未完成: {', '.join(incomplete)}")

        bundle = ExportService(self.session)._bundle(task_id)
        context = ReportService(self.session, self.config)._context(bundle)
        stats = context["stats"]
        coverage_rows = context["coverage_rows"]
        visual_failed = self._visual_failed(bundle)
        if visual_failed:
            failures.append(f"仍有 {visual_failed} 个 Web 入口视觉识别失败")
        if stats.get("Web入口数量", 0) and stats.get("已完成视觉识别Web入口", 0) < stats.get("Web入口数量", 0):
            failures.append(f"Web 识别未全覆盖: {stats.get('Web识别覆盖率')}")

        high_gaps = [row for row in coverage_rows if row.get("缺口等级") == "高"]
        medium_gaps = [row for row in coverage_rows if row.get("缺口等级") == "中"]
        low_gaps = [row for row in coverage_rows if row.get("缺口等级") == "低"]
        if high_gaps:
            failures.append("存在高等级覆盖缺口: " + ", ".join(_unique(row["环节"] for row in high_gaps)))
        if medium_gaps:
            warnings.append("存在中等级覆盖缺口: " + ", ".join(_unique(row["环节"] for row in medium_gaps)))
        if low_gaps:
            warnings.append("存在低等级覆盖缺口: " + ", ".join(_unique(row["环节"] for row in low_gaps)))

        paths = self._report_paths(task, output_dir)
        self._check_artifacts(paths, failures, warnings)

        status = "FAIL" if failures else ("WARN" if warnings else "PASS")
        lines = self._lines(
            task,
            status,
            stats,
            coverage_rows,
            context["visual_review_rows"],
            paths,
            failures,
            warnings,
        )
        return QualityResult(status=status, failures=failures, warnings=warnings, lines=lines)

    def _report_paths(self, task: ScanTask, output_dir: Path | str) -> dict[str, Path]:
        root = Path(output_dir) / f"task_{task.id}_{_safe_name(task.target)}"
        report = root / f"task_{task.id}_互联网资产暴露面测绘报告.docx"
        asset_workbook = root / f"task_{task.id}_资产汇总.xlsx"
        web_workbook = root / f"task_{task.id}_Web资产详情.xlsx"
        return {
            "Word报告": self._latest_artifact(report),
            "资产汇总附件": self._latest_artifact(asset_workbook),
            "Web资产详情附件": self._latest_artifact(web_workbook),
        }

    def _latest_artifact(self, preferred: Path) -> Path:
        candidates = [preferred] if preferred.exists() else []
        candidates.extend(preferred.parent.glob(f"{preferred.stem}_*{preferred.suffix}"))
        if not candidates:
            return preferred
        return max(candidates, key=lambda path: path.stat().st_mtime)

    def _check_artifacts(self, paths: dict[str, Path], failures: list[str], warnings: list[str]) -> None:
        for label, path in paths.items():
            if not path.exists():
                failures.append(f"{label}不存在: {path}")
            elif path.stat().st_size <= 0:
                failures.append(f"{label}为空文件: {path}")
        report_path = paths["Word报告"]
        if report_path.exists() and report_path.stat().st_size > 0:
            try:
                doc = Document(report_path)
                if len(doc.tables) < 6:
                    warnings.append(f"Word报告表格数量偏少: {len(doc.tables)}")
                if not doc.sections[0].header.paragraphs[0].text.strip():
                    warnings.append("Word报告缺少页眉")
                if not doc.sections[0].footer.paragraphs[0].text.strip():
                    warnings.append("Word报告缺少页脚")
            except Exception as exc:
                failures.append(f"Word报告无法打开: {str(exc)[:200]}")
        asset_path = paths["资产汇总附件"]
        if asset_path.exists() and asset_path.stat().st_size > 0:
            self._check_workbook(asset_path, EXPECTED_ASSET_SHEETS, "资产汇总附件", failures, warnings)
        web_path = paths["Web资产详情附件"]
        if web_path.exists() and web_path.stat().st_size > 0:
            self._check_workbook(web_path, EXPECTED_WEB_SHEETS, "Web资产详情附件", failures, warnings)

    def _check_workbook(
        self,
        path: Path,
        expected_sheets: set[str],
        label: str,
        failures: list[str],
        warnings: list[str],
    ) -> None:
        try:
            workbook = load_workbook(path, data_only=True)
        except Exception as exc:
            failures.append(f"{label}无法打开: {str(exc)[:200]}")
            return
        missing = sorted(expected_sheets - set(workbook.sheetnames))
        if missing:
            failures.append(f"{label}缺少Sheet: {', '.join(missing)}")
            return
        if label == "资产汇总附件":
            self._check_asset_workbook_contract(workbook, failures, warnings)
        elif label == "Web资产详情附件":
            self._check_web_workbook_contract(workbook, failures, warnings)

    def _check_asset_workbook_contract(self, workbook, failures: list[str], warnings: list[str]) -> None:
        dashboard = workbook["管理驾驶舱"]
        if len(getattr(dashboard, "_charts", [])) < 2:
            warnings.append("资产汇总附件管理驾驶舱图表数量偏少")
        self._require_headers(
            workbook["风险清单"],
            {"风险分值", "责任建议", "验收证据"},
            "资产汇总附件/风险清单",
            failures,
        )
        self._require_headers(
            workbook["重点资产视图"],
            {"风险分值", "责任建议"},
            "资产汇总附件/重点资产视图",
            failures,
        )
        self._require_headers(
            workbook["整改矩阵"],
            {"风险分值", "责任建议", "验收证据"},
            "资产汇总附件/整改矩阵",
            failures,
        )

    def _check_web_workbook_contract(self, workbook, failures: list[str], warnings: list[str]) -> None:
        sheet = workbook["截图证据"]
        self._require_headers(sheet, {"缩略图", "截图文件", "截图状态"}, "Web资产详情附件/截图证据", failures)
        headers = [cell.value for cell in sheet[1]]
        status_column = headers.index("截图状态") + 1 if "截图状态" in headers else None
        rows_with_screenshot = 0
        if status_column:
            rows_with_screenshot = sum(
                1
                for row_index in range(2, sheet.max_row + 1)
                if sheet.cell(row=row_index, column=status_column).value == "有截图"
            )
        if rows_with_screenshot and not getattr(sheet, "_images", []):
            warnings.append("Web资产详情附件截图证据存在截图路径但未嵌入缩略图")

    def _require_headers(self, sheet, required: set[str], label: str, failures: list[str]) -> None:
        headers = {cell.value for cell in sheet[1] if cell.value}
        missing = sorted(required - headers)
        if missing:
            failures.append(f"{label}缺少字段: {', '.join(missing)}")

    def _visual_failed(self, bundle: dict[str, Any]) -> int:
        return sum(
            1
            for row in bundle["web_entrypoints"]
            if (row.get("evidence") or {}).get("visual_analysis_error")
            and not (row.get("evidence") or {}).get("visual_analysis")
        )

    def _lines(
        self,
        task: ScanTask,
        status: str,
        stats: dict,
        coverage_rows: list[dict],
        visual_review_rows: list[dict],
        paths: dict[str, Path],
        failures: list[str],
        warnings: list[str],
    ) -> list[str]:
        lines = [
            f"Quality: {status}",
            f"Task: {task.id}",
            f"Target: {task.target}",
            "",
            "Key metrics:",
            f"- Companies: {stats.get('单位数量', 0)}",
            f"- Base assets: {stats.get('资产数量', 0)}",
            f"- Port target IPs: {stats.get('端口目标数量', 0)}",
            f"- Open ports: {stats.get('开放端口数量', 0)}",
            f"- Web entrypoints: {stats.get('Web入口数量', 0)}",
            f"- Web visual coverage: {stats.get('Web识别覆盖率', '0/0 (0%)')}",
            f"- High risks: {stats.get('高风险项', 0)}",
            f"- Medium risks: {stats.get('中风险项', 0)}",
            "",
            "Coverage gates:",
        ]
        for row in coverage_rows:
            lines.append(f"- {row.get('环节')}: {row.get('缺口等级')} ({row.get('结果')})")
        lines.extend(["", "Artifacts:"])
        for label, path in paths.items():
            state = "ok" if path.exists() and path.stat().st_size > 0 else "missing"
            lines.append(f"- {label}: {state} -> {path}")
        if failures:
            lines.extend(["", "Failures:"])
            lines.extend(f"- {item}" for item in failures)
        if warnings:
            lines.extend(["", "Warnings:"])
            lines.extend(f"- {item}" for item in warnings)
        next_actions = self._next_actions(task, coverage_rows, visual_review_rows, failures, warnings)
        if next_actions:
            lines.extend(["", "Suggested next actions:"])
            lines.extend(f"- {item}" for item in next_actions)
        return lines

    def _next_actions(
        self,
        task: ScanTask,
        coverage_rows: list[dict],
        visual_review_rows: list[dict],
        failures: list[str],
        warnings: list[str],
    ) -> list[str]:
        actions: list[str] = []
        manual_review_needed = False
        if failures:
            actions.append(f"先执行 assetmap run {task.id} 补齐未完成流程，再重新执行 quality-check。")
        for row in coverage_rows:
            if row.get("环节") == "企业/备案资产" and row.get("缺口等级") in {"高", "中", "低"}:
                actions.append(
                    f"生成缺口补充模板：assetmap asset-gap-template {task.id} --priority high-medium --include-partial --force --output data/manual_assets.task_{task.id}.gaps.yaml"
                )
                actions.append(
                    f"补充模板后执行：assetmap run {task.id} --manual-file data/manual_assets.task_{task.id}.gaps.yaml"
                )
                break
        dns_gaps = [
            row
            for row in coverage_rows
            if row.get("环节") == "子域名/DNS" and row.get("缺口等级") in {"高", "中", "低"}
        ]
        if any(row.get("指标") in {"根域名解析覆盖", "子域名枚举质量"} for row in dns_gaps):
            actions.append(
                f"重跑子域名枚举和 DNS，并刷新后续流程：assetmap run {task.id} --from-stage subdomains --rerun-subdomain-tools"
            )
        if any(row.get("指标") == "DNS 解析复核质量" for row in dns_gaps):
            actions.append("复核 DNS复核清单中的无公网解析、第三方/停放 CNAME 和共享 IP 线索；确认停放或无独立业务后在复核工作单留痕。")
            manual_review_needed = True
        if any(row.get("环节") == "端口发现" and row.get("指标") == "端口证据来源质量" and row.get("缺口等级") in {"高", "中", "低"} for row in coverage_rows):
            actions.append(
                f"端口台账存在仅被动FOFA证据；如需主动验证，可执行 assetmap nmap-scan {task.id} --sources nmap,fofa --rerun，"
                "系统会优先精确验证已有 FOFA 端口；也可将 config.yaml 的 port_scan.sources_enabled 固化为 nmap+fofa。"
            )
        service_gaps = [
            row
            for row in coverage_rows
            if row.get("环节") == "服务识别/URL" and row.get("缺口等级") in {"高", "中", "低"}
        ]
        if any(row.get("指标") in {"Web 服务入口覆盖", "URL 入口关联质量"} for row in service_gaps):
            actions.append(
                f"复核服务识别和 URL 入口后重跑后续流程：assetmap run {task.id} --from-stage classify --rerun-classify"
            )
        if any(row.get("指标") == "服务分类复核" for row in service_gaps):
            actions.append(
                "复核服务识别台账中的 passive_fofa、疑似 Web 和未知服务项；仅当需要重新探测 Host、服务指纹或 URL 入口时再重跑 classify。"
            )
            manual_review_needed = True
        if any(row.get("环节") == "URL视觉识别" and row.get("缺口等级") in {"高", "中", "低"} for row in coverage_rows):
            if self._has_visual_retry_items(visual_review_rows):
                actions.append(f"重试截图失败或缺失识别页面：assetmap url-discover {task.id} --retry-failed")
            else:
                actions.append("对视觉复核清单中的降级识别、低置信度、低价值错误/拦截页或疑似乱码页面进行人工核对；当前无自动重试项。")
                manual_review_needed = True
        if manual_review_needed:
            actions.append(
                f"复核工作单填写 review_status 后导入结论：assetmap import-review {task.id} --file data/review_workorder.task_{task.id}.yaml"
            )
        return actions

    def _has_visual_retry_items(self, visual_review_rows: list[dict]) -> bool:
        return any(
            row.get("复核类型") == "automatic_retry"
            or row.get("分析错误")
            or not row.get("识别方式")
            for row in visual_review_rows
        )


def _unique(values) -> list[str]:
    output = []
    for value in values:
        if value and value not in output:
            output.append(value)
    return output
