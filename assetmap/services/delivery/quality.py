from __future__ import annotations

import shutil
import subprocess
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from sqlmodel import Session

from assetmap.config import AppConfig
from assetmap.models import ScanTask
from assetmap.services.delivery.exporter import ExportService
from assetmap.services.delivery.report import ReportService, _safe_name
from assetmap.services.operations.status import PipelineStatusService


EXPECTED_ASSET_SHEETS = {
    "阅读导航",
    "管理驾驶舱",
    "报告概览",
    "AI分析审计",
    "风险统计",
    "重点资产视图",
    "单位覆盖台账",
    "资产汇总",
    "覆盖缺口",
    "风险清单",
    "整改矩阵",
    "DNS记录",
    "DNS复核清单",
    "端口目标台账",
    "开放端口",
    "服务识别台账",
    "URL入口覆盖",
    "交付审计文件",
    "非域名资产",
}
EXPECTED_WEB_SHEETS = {"阅读导航", "重点Web资产", "HTML证据", "Web资产详情", "页面识别复核清单"}


@dataclass
class QualityResult:
    status: str
    failures: list[str]
    warnings: list[str]
    lines: list[str]


class DeliveryQualityService:
    def __init__(self, session: Session, config: AppConfig) -> None:
        self.session = session
        self.config = config

    def check(self, task_id: int, output_dir: Path | str = "reports") -> QualityResult:
        task = self.session.get(ScanTask, task_id)
        if not task:
            raise ValueError(f"Task not found: {task_id}")
        failures: list[str] = []
        warnings: list[str] = []

        pipeline = PipelineStatusService(self.session).get(task_id)
        incomplete, pipeline_warnings = self._pipeline_issues(pipeline.stages)
        if incomplete:
            failures.append(f"流程未完成: {', '.join(incomplete)}")
        warnings.extend(pipeline_warnings)

        bundle = ExportService(self.session)._bundle(task_id)
        context = ReportService(self.session, self.config)._context(bundle)
        stats = context["stats"]
        coverage_rows = context["coverage_rows"]
        visual_failed = self._visual_failed(bundle)
        if visual_failed:
            failures.append(f"仍有 {visual_failed} 个 Web 入口页面识别失败")
        if stats.get("Web入口数量", 0) and stats.get("已完成页面识别Web入口", 0) < stats.get("Web入口数量", 0):
            failures.append(f"Web 识别未全覆盖: {stats.get('Web识别覆盖率')}")

        high_gaps = [row for row in coverage_rows if row.get("缺口等级") == "高"]
        medium_gaps = [row for row in coverage_rows if row.get("缺口等级") == "中"]
        low_gaps = [row for row in coverage_rows if row.get("缺口等级") == "低"]
        if high_gaps:
            failures.append("存在高等级覆盖缺口: " + ", ".join(_unique(row["环节"] for row in high_gaps)))
        if medium_gaps:
            warnings.append("存在中等级覆盖缺口: " + ", ".join(_unique(row["环节"] for row in medium_gaps)))
        if low_gaps:
            warnings.append("存在低等级覆盖缺口: " + ", ".join(_unique(row["环节"] for row in low_gaps)))

        paths = self._report_paths(task, output_dir)
        self._check_artifacts(paths, failures, warnings)

        status = "FAIL" if failures else ("WARN" if warnings else "PASS")
        lines = self._lines(
            task,
            status,
            stats,
            coverage_rows,
            context["visual_review_rows"],
            paths,
            failures,
            warnings,
        )
        return QualityResult(status=status, failures=failures, warnings=warnings, lines=lines)

    def _pipeline_issues(self, stages: list[tuple[str, str, str]]) -> tuple[list[str], list[str]]:
        """Separate resumable partial coverage from hard delivery blockers."""
        incomplete: list[str] = []
        warnings: list[str] = []
        for name, status, detail in stages:
            if status == "completed":
                continue
            if status == "completed_with_errors" and name in {"subdomains", "port-scan"}:
                warnings.append(f"{name}阶段存在失败子任务，交付数据可能不完整: {detail}")
                continue
            incomplete.append(name)
        return incomplete, warnings

    def _report_paths(self, task: ScanTask, output_dir: Path | str) -> dict[str, Path]:
        root = Path(output_dir) / f"task_{task.id}_{_safe_name(task.target)}"
        report = root / f"task_{task.id}_互联网资产暴露面测绘报告.docx"
        asset_workbook = root / f"task_{task.id}_资产汇总.xlsx"
        web_workbook = root / f"task_{task.id}_Web资产详情.xlsx"
        return {
            "Word报告": self._latest_artifact(report),
            "资产汇总附件": self._latest_artifact(asset_workbook),
            "Web资产详情附件": self._latest_artifact(web_workbook),
        }

    def _latest_artifact(self, preferred: Path) -> Path:
        candidates = [preferred] if preferred.exists() else []
        candidates.extend(preferred.parent.glob(f"{preferred.stem}_*{preferred.suffix}"))
        if not candidates:
            return preferred
        return max(candidates, key=lambda path: path.stat().st_mtime)

    def _check_artifacts(self, paths: dict[str, Path], failures: list[str], warnings: list[str]) -> None:
        for label, path in paths.items():
            if not path.exists():
                failures.append(f"{label}不存在: {path}")
            elif path.stat().st_size <= 0:
                failures.append(f"{label}为空文件: {path}")
        report_path = paths["Word报告"]
        if report_path.exists() and report_path.stat().st_size > 0:
            try:
                doc = Document(report_path)
                if len(doc.tables) < 6:
                    warnings.append(f"Word报告表格数量偏少: {len(doc.tables)}")
                if not doc.sections[0].header.paragraphs[0].text.strip():
                    warnings.append("Word报告缺少页眉")
                if not doc.sections[0].footer.paragraphs[0].text.strip():
                    warnings.append("Word报告缺少页脚")
                self._check_docx_renderability(report_path, warnings)
            except Exception as exc:
                failures.append(f"Word报告无法打开: {str(exc)[:200]}")
        asset_path = paths["资产汇总附件"]
        if asset_path.exists() and asset_path.stat().st_size > 0:
            self._check_workbook(asset_path, EXPECTED_ASSET_SHEETS, "资产汇总附件", failures, warnings)
        web_path = paths["Web资产详情附件"]
        if web_path.exists() and web_path.stat().st_size > 0:
            self._check_workbook(web_path, EXPECTED_WEB_SHEETS, "Web资产详情附件", failures, warnings)

    def _check_docx_renderability(self, report_path: Path, warnings: list[str]) -> None:
        """Confirm that the generated Word document can be rendered to PDF.

        Opening OOXML only proves that the ZIP/XML is readable.  LibreOffice
        conversion catches a separate class of customer-facing defects such as
        unsupported objects, damaged fonts, and invalid layout markup.  This is
        intentionally a warning when LibreOffice is unavailable so ordinary
        data processing is not blocked; formal delivery should use --strict.
        """
        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            warnings.append("未执行 Word 渲染检查：未检测到 LibreOffice；正式交付前请在可渲染环境复核版式")
            return
        try:
            with tempfile.TemporaryDirectory(prefix="assetmap-docx-render-") as temp_dir:
                output_dir = Path(temp_dir) / "output"
                profile_dir = Path(temp_dir) / "profile"
                output_dir.mkdir()
                profile_dir.mkdir()
                command = [
                    soffice,
                    f"-env:UserInstallation={profile_dir.resolve().as_uri()}",
                    "--headless",
                    "--norestore",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    str(output_dir),
                    str(report_path),
                ]
                result = subprocess.run(command, capture_output=True, text=True, timeout=90, check=False)
                pdf_path = output_dir / f"{report_path.stem}.pdf"
                if result.returncode != 0 or not pdf_path.exists() or pdf_path.stat().st_size <= 0:
                    detail = (result.stderr or result.stdout or "未知转换错误").strip().replace("\n", " ")
                    warnings.append(f"Word 渲染检查未通过：{detail[:200]}")
        except (OSError, subprocess.SubprocessError) as exc:
            warnings.append(f"Word 渲染检查未完成：{str(exc)[:200]}")

    def _check_workbook(
        self,
        path: Path,
        expected_sheets: set[str],
        label: str,
        failures: list[str],
        warnings: list[str],
    ) -> None:
        try:
            workbook = load_workbook(path, data_only=True)
        except Exception as exc:
            failures.append(f"{label}无法打开: {str(exc)[:200]}")
            return
        missing = sorted(expected_sheets - set(workbook.sheetnames))
        if missing:
            failures.append(f"{label}缺少Sheet: {', '.join(missing)}")
            return
        if label == "资产汇总附件":
            self._check_asset_workbook_contract(workbook, failures, warnings)
        elif label == "Web资产详情附件":
            self._check_web_workbook_contract(workbook, failures, warnings)

    def _check_asset_workbook_contract(self, workbook, failures: list[str], warnings: list[str]) -> None:
        dashboard = workbook["管理驾驶舱"]
        if len(getattr(dashboard, "_charts", [])) < 2:
            warnings.append("资产汇总附件管理驾驶舱图表数量偏少")
        self._require_headers(
            workbook["风险清单"],
            {"风险分值", "责任建议", "验收证据"},
            "资产汇总附件/风险清单",
            failures,
        )
        self._require_headers(
            workbook["重点资产视图"],
            {"风险分值", "责任建议"},
            "资产汇总附件/重点资产视图",
            failures,
        )
        self._require_headers(
            workbook["整改矩阵"],
            {"风险分值", "责任建议", "验收证据"},
            "资产汇总附件/整改矩阵",
            failures,
        )

    def _check_web_workbook_contract(self, workbook, failures: list[str], warnings: list[str]) -> None:
        sheet = workbook["HTML证据"]
        self._require_headers(sheet, {"HTML文件", "HTML状态"}, "Web资产详情附件/HTML证据", failures)

    def _require_headers(self, sheet, required: set[str], label: str, failures: list[str]) -> None:
        headers = {cell.value for cell in sheet[1] if cell.value}
        missing = sorted(required - headers)
        if missing:
            failures.append(f"{label}缺少字段: {', '.join(missing)}")

    def _visual_failed(self, bundle: dict[str, Any]) -> int:
        return sum(
            1
            for row in bundle["web_entrypoints"]
            if (row.get("evidence") or {}).get("visual_analysis_error")
            and not (row.get("evidence") or {}).get("visual_analysis")
        )

    def _lines(
        self,
        task: ScanTask,
        status: str,
        stats: dict,
        coverage_rows: list[dict],
        visual_review_rows: list[dict],
        paths: dict[str, Path],
        failures: list[str],
        warnings: list[str],
    ) -> list[str]:
        lines = [
            f"Quality: {status}",
            f"Task: {task.id}",
            f"Target: {task.target}",
            "",
            "Key metrics:",
            f"- Companies: {stats.get('单位数量', 0)}",
            f"- Base assets: {stats.get('资产数量', 0)}",
            f"- Port target IPs: {stats.get('端口目标数量', 0)}",
            f"- Open ports: {stats.get('开放端口数量', 0)}",
            f"- Web entrypoints: {stats.get('Web入口数量', 0)}",
            f"- Web visual coverage: {stats.get('Web识别覆盖率', '0/0 (0%)')}",
            f"- High risks: {stats.get('高风险项', 0)}",
            f"- Medium risks: {stats.get('中风险项', 0)}",
            "",
            "Coverage gates:",
        ]
        for row in coverage_rows:
            lines.append(f"- {row.get('环节')}: {row.get('缺口等级')} ({row.get('结果')})")
        lines.extend(["", "Artifacts:"])
        for label, path in paths.items():
            state = "ok" if path.exists() and path.stat().st_size > 0 else "missing"
            lines.append(f"- {label}: {state} -> {path}")
        if failures:
            lines.extend(["", "Failures:"])
            lines.extend(f"- {item}" for item in failures)
        if warnings:
            lines.extend(["", "Warnings:"])
            lines.extend(f"- {item}" for item in warnings)
        next_actions = self._next_actions(task, coverage_rows, visual_review_rows, failures, warnings)
        if next_actions:
            lines.extend(["", "Suggested next actions:"])
            lines.extend(f"- {item}" for item in next_actions)
        return lines

    def _next_actions(
        self,
        task: ScanTask,
        coverage_rows: list[dict],
        visual_review_rows: list[dict],
        failures: list[str],
        warnings: list[str],
    ) -> list[str]:
        actions: list[str] = []
        manual_review_needed = False
        if failures:
            actions.append(f"先执行 assetmap run {task.id} 补齐未完成流程，再重新执行 quality-check。")
        for row in coverage_rows:
            if row.get("环节") == "企业/备案资产" and row.get("缺口等级") in {"高", "中", "低"}:
                actions.append(
                    f"生成缺口补充模板：assetmap asset-gap-template {task.id} --priority high-medium --include-partial --force --output data/manual_assets.task_{task.id}.gaps.yaml"
                )
                actions.append(
                    f"补充模板后执行：assetmap run {task.id} --manual-file data/manual_assets.task_{task.id}.gaps.yaml"
                )
                break
        dns_gaps = [
            row
            for row in coverage_rows
            if row.get("环节") == "子域名/DNS" and row.get("缺口等级") in {"高", "中", "低"}
        ]
        if any(row.get("指标") in {"根域名解析覆盖", "子域名枚举质量"} for row in dns_gaps):
            actions.append(
                f"重跑子域名枚举和 DNS，并刷新后续流程：assetmap run {task.id} --from-stage subdomains --rerun-subdomain-tools"
            )
        if any(row.get("指标") == "DNS 解析复核质量" for row in dns_gaps):
            actions.append("复核 DNS复核清单中的无公网解析、第三方/停放 CNAME 和共享 IP 线索；确认停放或无独立业务后在复核工作单留痕。")
            manual_review_needed = True
        if any(row.get("环节") == "端口发现" and row.get("指标") == "端口证据来源质量" and row.get("缺口等级") in {"高", "中", "低"} for row in coverage_rows):
            actions.append(
                f"端口台账存在仅被动FOFA证据；可执行 assetmap nmap-scan {task.id} --rerun，"
                "系统会串行精确验证已有 FOFA 端口并合并证据。"
            )
        service_gaps = [
            row
            for row in coverage_rows
            if row.get("环节") == "服务识别/URL" and row.get("缺口等级") in {"高", "中", "低"}
        ]
        if any(row.get("指标") in {"Web 服务入口覆盖", "URL 入口关联质量"} for row in service_gaps):
            actions.append(
                f"复核服务识别和 URL 入口后重跑后续流程：assetmap run {task.id} --from-stage classify --rerun-classify"
            )
        if any(row.get("指标") == "服务分类复核" for row in service_gaps):
            actions.append(
                "复核服务识别台账中的 passive_fofa、疑似 Web 和未知服务项；仅当需要重新探测 Host、服务指纹或 URL 入口时再重跑 classify。"
            )
            manual_review_needed = True
        if any(row.get("环节") == "URL页面识别" and row.get("缺口等级") in {"高", "中", "低"} for row in coverage_rows):
            if self._has_visual_retry_items(visual_review_rows):
                actions.append(f"重试页面渲染或识别失败入口：assetmap url-discover {task.id} --retry-failed")
            else:
                actions.append("对页面识别复核清单中的降级识别、低置信度、低价值错误/拦截页或疑似乱码页面进行人工核对；当前无自动重试项。")
                manual_review_needed = True
        if manual_review_needed:
            actions.append(
                f"复核工作单填写 review_status 后导入结论：assetmap import-review {task.id} --file data/review_workorder.task_{task.id}.yaml"
            )
        return actions

    def _has_visual_retry_items(self, visual_review_rows: list[dict]) -> bool:
        return any(
            row.get("复核类型") == "automatic_retry"
            or row.get("分析错误")
            or not row.get("识别方式")
            for row in visual_review_rows
        )


def _unique(values) -> list[str]:
    output = []
    for value in values:
        if value and value not in output:
            output.append(value)
    return output
